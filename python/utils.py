"""
utils.py — Shared helpers untuk semua paket penomoran halaman.
Semua fungsi XML/docx, deteksi zona, dan section formatting ada di sini.
"""
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt, Cm
from copy import deepcopy
import re


# =========================================================
# DETECTION — module-level (tidak butuh doc/font)
# =========================================================

ROMAN_START_KEYWORDS = [
    "kata pengantar", "prakata", "ucapan terima kasih",
    "abstrak",
    "lembar pengesahan", "lembar persetujuan", "halaman pengesahan",
    "pengesahan", "persetujuan",
    "lembar pernyataan", "pernyataan keaslian",
    "halaman pernyataan",
    "rekomendasi",
    "halaman persembahan", "persembahan", "motto",
    "ringkasan",
    "daftar isi",
    "daftar tabel", "daftar gambar", "daftar grafik", "daftar rumus",
    "daftar singkatan", "daftar lambang", "daftar notasi", "daftar simbol",
    "daftar istilah", "daftar arti lambang", "daftar arti simbol",
    "abstract", "summary", "executive summary",
    "preface", "foreword", "acknowledgment", "acknowledgements",
    "approval page", "approval sheet", "declaration", "originality statement",
    "dedication",
    "table of contents", "list of contents", "contents",
    "list of tables", "list of figures", "list of appendices",
    "list of abbreviations", "list of symbols", "list of notations",
    "nomenclature",
]

_ROMAN_PAT = r'm{0,4}(cm|cd|d?c{0,3})(xc|xl|l?x{0,3})(ix|iv|v?i{0,3})'
BAB_HEAD_RE = re.compile(
    rf'^\s*(bab|chapter)\s*[-.]?\s+({_ROMAN_PAT}|\d+)\.?\b(.*?)$',
    re.IGNORECASE | re.DOTALL
)


def is_roman_start(text):
    lower = text.strip().lower()
    return any(lower == k or lower.startswith(k) for k in ROMAN_START_KEYWORDS)


def _has_toc_field(p_elem):
    """Cek apakah paragraf mengandung field TOC otomatis Word ({TOC} field)."""
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for instr in p_elem.iter('{%s}instrText' % W):
        if instr.text and 'TOC' in instr.text.upper():
            return True
    return False


def is_bab_heading(text):
    text = text.strip()
    if not text:
        return False
    if BAB_HEAD_RE.match(text):
        return True
    if re.match(
        r'^\s*(daftar\s+pust?aka|referensi|references?|reference\s+list|bibliography|'
        r'bibliographies|works?\s+cited|literature\s+cited)\s*$', text, re.IGNORECASE
    ):
        return True
    if re.match(r'^\s*(lampiran|appendix|appendices|attachment)(\s+.*)?$', text, re.IGNORECASE):
        return True
    return False


def is_false_bab(para):
    text  = para.text.strip()
    style = para.style.name.lower() if para.style else ""
    # Entry daftar isi: "BAB I\tJUDUL\t1" — ada tab + nomor halaman di akhir
    if is_toc_entry(text):
        return True
    m = BAB_HEAD_RE.match(text)
    if m:
        sisa_raw = m.group(6) or ""
        sisa = sisa_raw.strip()
        # Heading 1 dengan teks BAB → pasti BAB nyata, skip length check.
        # "BAB II PENGARUH JENIS MINUMAN TERHADAP KONDISI GIGI (SIMULASI...)" boleh panjang.
        _is_h1 = bool(re.match(r'^heading\s*1$', style))
        if not _is_h1 and not sisa_raw.startswith('\n'):
            if len(sisa) > 60:
                return True
            if len(text.split()) > 8:
                return True
        non_heading = ('body', 'list', 'quote', 'caption', 'web')
        if sisa and any(s in style for s in non_heading):
            return True
    for pattern in [
        r'^\s*(daftar\s+pust?aka|referensi|references?|reference\s+list|bibliography|'
        r'bibliographies|works?\s+cited|literature\s+cited)\s+\S',
        r'^\s*(lampiran|appendix|appendices|attachment)\s+[^\n]{60,}',
    ]:
        if re.match(pattern, text, re.IGNORECASE):
            return True
    if re.match(r'^\s*(lampiran|appendix|appendices|attachment)', text, re.IGNORECASE):
        if (re.search(r'\t\s*\d+\s*$', text) or
                re.search(r'\.{3,}.*\d+\s*$', text) or
                re.search(r'\s{4,}\d+\s*$', text)):
            return True
    # LAMPIRAN / DAFTAR PUSTAKA sebagai sub-heading (Heading 2+) bukan section mandiri
    _endpoint_pat = (
        r'^\s*(lampiran|appendix|appendices|attachment|'
        r'daftar\s+pust?aka|referensi|references?|bibliography)'
    )
    if re.match(_endpoint_pat, text, re.IGNORECASE):
        if re.search(r'heading\s*[2-9]', style, re.IGNORECASE):
            return True
    return False


def is_toc_heading(text):
    lower = text.strip().lower()
    return any(k in lower for k in [
        "daftar isi", "daftar tabel", "daftar gambar", "daftar lampiran",
        "table of contents", "list of contents", "list of tables",
        "list of figures", "list of appendices",
    ])


def is_toc_entry(text):
    # Multi-line (soft return): TOC entry hanya jika ada pola nomor halaman di salah satu baris.
    # "BAB I\nPENDAHULUAN" tidak punya nomor → bukan TOC entry.
    if '\n' in text:
        for line in text.split('\n'):
            if re.search(r'\t\s*[\divxlcdmIVXLCDM]+\s*$', line):
                return True
            if re.search(r'\.{3,}.*\d+\s*$', line):
                return True
        return False
    if re.search(r'\t\s*[\divxlcdmIVXLCDM]+\s*$', text):
        return True
    if re.search(r'\.{3,}.*\d+\s*$', text):
        return True
    if re.search(r'\s{4,}\d+\s*$', text):
        return True
    return False


# =========================================================
# DocProcessor — menyimpan doc + font settings, semua metode
# =========================================================

class DocProcessor:

    def __init__(self, doc, font_name, font_size):
        self.doc       = doc
        self.font_name = font_name
        self.font_size = font_size  # int (pt)

    # ── Low-level helpers ─────────────────────────────────

    def add_page_number(self, paragraph):
        run = paragraph.add_run()
        run.font.name = self.font_name
        run.font.size = Pt(self.font_size)
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = " PAGE "
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    @staticmethod
    def _first_para(part):
        if not part.paragraphs:
            part._element.append(OxmlElement('w:p'))
        return part.paragraphs[0]

    @staticmethod
    def clear_paragraph(paragraph):
        p = paragraph._p
        for child in list(p):
            p.remove(child)

    @staticmethod
    def _set_pn_spacing(p):
        pf = p.paragraph_format
        pf.space_after        = Pt(0)
        pf.space_before       = Pt(0)
        pf.line_spacing_rule  = WD_LINE_SPACING.SINGLE

    @staticmethod
    def _get_h_align(pos):
        if 'Tengah' in pos: return WD_ALIGN_PARAGRAPH.CENTER
        if 'Kanan'  in pos: return WD_ALIGN_PARAGRAPH.RIGHT
        return WD_ALIGN_PARAGRAPH.LEFT

    @staticmethod
    def _is_top(pos):
        return 'Atas' in pos

    def _place_num_in_part(self, part, align):
        part.is_linked_to_previous = False
        for p in part.paragraphs:
            self.clear_paragraph(p)
        p = self._first_para(part)
        self.clear_paragraph(p)
        p.alignment = align
        self.add_page_number(p)
        self._set_pn_spacing(p)

    # ── Cover advance helper ──────────────────────────────

    @staticmethod
    def advance_roman_start(doc, roman_start_p, num_cover):
        """
        Geser roman_start_p ke paragraf pada halaman (num_cover + 1).
        Returns (new_roman_start_p, use_exact):
          use_exact=True  → caller harus set exact_roman_start=True di insert_breaks,
                            agar _find_section_start tidak memindahkan roman_start_p
                            ke awal section (yang akan mengacaukan multi-cover).
          use_exact=False → biarkan insert_breaks berjalan normal.

        Kasus:
          breaks_before >= num_cover        : advance ke halaman num_cover+1.
          breaks_before == num_cover - 1    : dokumen sudah punya jumlah section cover
                                              yang tepat → gunakan roman_start_p apa adanya
                                              tapi set use_exact=True.
          breaks_before < num_cover - 1 / 0: tidak bisa diperbaiki tanpa rendering.
        """
        W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        body_els = list(doc.element.body)

        if num_cover <= 1:
            # Untuk num_cover=1: exit early hanya jika bb ≤ 1 (tidak ada extra cover pages).
            # Jika bb > 1, ada duplikat cover — lanjutkan ke advance logic agar cover
            # section hanya mencakup 1 halaman saja.
            def _hb_early(el):
                if any(br.get("{%s}type" % W) == 'page' for br in el.iter("{%s}br" % W)):
                    return True
                pPr = el.find("{%s}pPr" % W)
                return pPr is not None and pPr.find("{%s}sectPr" % W) is not None
            try:
                _rsp_i = body_els.index(roman_start_p)
            except ValueError:
                return roman_start_p, False
            _bb = sum(1 for el in body_els[:_rsp_i] if _hb_early(el))
            if _bb <= num_cover:
                return roman_start_p, False
            # _bb > num_cover: lanjut ke advance logic di bawah

        def _has_break(el):
            if any(br.get("{%s}type" % W) == 'page' for br in el.iter("{%s}br" % W)):
                return True
            pPr = el.find("{%s}pPr" % W)
            return pPr is not None and pPr.find("{%s}sectPr" % W) is not None

        # Hitung page break sebelum roman_start_p
        try:
            rsp_idx = body_els.index(roman_start_p)
        except ValueError:
            return roman_start_p, False

        breaks_before = sum(1 for el in body_els[:rsp_idx] if _has_break(el))

        def _txt(el):
            return ''.join(t.text or '' for t in el.iter('{%s}t' % W)).strip()

        def _el_tag(el):
            return el.tag.split('}')[-1] if '}' in el.tag else el.tag

        def _is_empty_sectPr(el):
            if _el_tag(el) != 'p':
                return False
            if _txt(el):
                return False
            pPr = el.find('{%s}pPr' % W)
            return pPr is not None and pPr.find('{%s}sectPr' % W) is not None

        def _remove_pgbr(el):
            for br in list(el.findall('.//{%s}br' % W)):
                if br.get('{%s}type' % W) == 'page':
                    p = br.getparent()
                    if p is not None:
                        p.remove(br)

        # Advance: hanya jika ada cukup break eksplisit (satu per halaman cover).
        if breaks_before >= num_cover:
            target_pg  = num_cover + 1
            current_pg = 1
            candidate  = None
            for el in body_els:
                if current_pg >= target_pg:
                    candidate = el
                    break
                if _has_break(el):
                    current_pg += 1

            if candidate is not None:
                # Lewati paragraf kosong ber-sectPr (cegah blank page extra).
                # Juga hapus pgBr dari paragraf tepat sebelum paragraf kosong itu.
                cand_idx = next((i for i, e in enumerate(body_els) if e is candidate), len(body_els))
                while cand_idx < len(body_els) and _is_empty_sectPr(body_els[cand_idx]):
                    if cand_idx > 0:
                        _remove_pgbr(body_els[cand_idx - 1])
                    cand_idx += 1
                candidate = body_els[cand_idx] if cand_idx < len(body_els) else candidate

                # Cek apakah ada is_roman_start di antara roman_start_p dan candidate
                cand_idx2 = next((i for i, e in enumerate(body_els) if e is candidate), len(body_els))
                for el in body_els[rsp_idx + 1:cand_idx2]:
                    _t = _el_tag(el)
                    if _t == 'p':
                        _t2 = _txt(el)
                        if _t2 and is_roman_start(_t2):
                            return el, True
                # num_cover=1: gunakan use_exact=False agar allow_empty_boundary tidak aktif
                # saat insert_breaks memproses BAB boundaries (mencegah sectPr kosong ekstra).
                return candidate, (num_cover > 1)

        # breaks_before == num_cover - 1: sudah ada (num_cover-1) break sebelum roman_start_p.
        # Step 1: cek page-header yang diulang (kasus Docx 2 — mundurkan rsp ke judul).
        # Step 2: scan maju dari rsp untuk page break eksplisit — kalau ketemu, roman zone
        #         sebenarnya dimulai SETELAH page break itu (halaman num_cover masih cover).
        if breaks_before == num_cover - 1:
            new_rsp   = DocProcessor._find_roman_page_start(body_els, rsp_idx, W)
            moved_back = new_rsp is not body_els[rsp_idx]

            if not moved_back:
                # Fix A: jika break terakhir sebelum rsp adalah sectPr di paragraf KOSONG,
                # boundary sudah ada dengan benar → tidak perlu advance.
                # (sectPr di paragraf ber-konten = batas cover1/cover2, bukan batas covers/roman)
                for _j in range(rsp_idx - 1, -1, -1):
                    if _has_break(body_els[_j]):
                        _pPr = body_els[_j].find('{%s}pPr' % W)
                        if (_pPr is not None and _pPr.find('{%s}sectPr' % W) is not None
                                and not _txt(body_els[_j])):
                            return roman_start_p, True
                        # Fix C (Docx 9): pgBr kosong sebelum rsp dengan banyak konten sebelumnya
                        # → cover sudah memenuhi ≥ num_cover halaman → pgBr adalah batas covers/roman.
                        if (any(br.get('{%s}type' % W) == 'page'
                                for br in body_els[_j].iter('{%s}br' % W))
                                and not _txt(body_els[_j])):
                            _nonempty = sum(
                                1 for _k in range(_j)
                                if _el_tag(body_els[_k]) == 'p' and _txt(body_els[_k])
                            )
                            if _nonempty >= 20:
                                return roman_start_p, True
                        break  # last break adalah pgBr atau sectPr ber-konten → lanjut

                # Hapus pgBr dari paragraf kosong tepat sebelum rsp (cegah blank page
                # dalam cover section ketika cover penuh mengisi halaman).
                if rsp_idx > 0:
                    prev_el = body_els[rsp_idx - 1]
                    if (_el_tag(prev_el) == 'p' and not _txt(prev_el) and
                            any(br.get('{%s}type' % W) == 'page'
                                for br in prev_el.iter('{%s}br' % W))):
                        _remove_pgbr(prev_el)

                # Scan maju untuk page break eksplisit.
                pgbr_found = False
                for j in range(rsp_idx, min(rsp_idx + 150, len(body_els))):
                    el  = body_els[j]
                    tag = _el_tag(el)
                    if tag != 'p':
                        continue
                    pPr = el.find('{%s}pPr' % W)
                    if pPr is not None and pPr.find('{%s}sectPr' % W) is not None:
                        break  # section boundary — stop
                    if any(br.get('{%s}type' % W) == 'page' for br in el.iter('{%s}br' % W)):
                        for k in range(j + 1, len(body_els)):
                            nxt  = body_els[k]
                            ntag = _el_tag(nxt)
                            if ntag == 'p':
                                new_rsp = nxt
                                pgbr_found = True
                                break
                        break

                # Fallback: jika tidak ada pgBr, cari roman keyword berikutnya
                # (menangani kasus sectPr asli di-dokumen yang sudah membatasi cover1).
                if not pgbr_found:
                    rsp_txt = _txt(body_els[rsp_idx]).lower()
                    for j in range(rsp_idx + 1, min(rsp_idx + 50, len(body_els))):
                        nxt = body_els[j]
                        if _el_tag(nxt) != 'p':
                            continue
                        pPr = nxt.find('{%s}pPr' % W)
                        if pPr is not None and pPr.find('{%s}sectPr' % W) is not None:
                            break
                        ntxt = _txt(nxt).strip()
                        if ntxt and is_roman_start(ntxt) and ntxt.lower() != rsp_txt:
                            new_rsp = nxt
                            break

            return new_rsp, True

        # breaks_before < num_cover - 1 atau 0: coba forward scan sampai ≤50 elemen.
        # pgBr dalam roman zone → jangan advance (tetap di rsp). Dua kondisi:
        #   (B1) element setelah pgBr adalah roman keyword (Docx 3, 18).
        #   (B2) ada roman keyword ANTARA rsp dan pgBr (Docx 6 — DAFTAR ISI manual).
        for j in range(rsp_idx, min(rsp_idx + 50, len(body_els))):
            el  = body_els[j]
            tag = _el_tag(el)
            if tag != 'p':
                continue
            pPr = el.find('{%s}pPr' % W)
            if pPr is not None and pPr.find('{%s}sectPr' % W) is not None:
                break  # section boundary — stop
            if any(br.get('{%s}type' % W) == 'page' for br in el.iter('{%s}br' % W)):
                # Fix B2: cek apakah ada roman keyword antara rsp dan pgBr ini.
                # Pengecualian: jika rsp sendiri adalah 'daftar isi', konten antara
                # rsp dan pgBr adalah entry TOC, bukan heading section — tidak trigger.
                _rsp_lower = _txt(body_els[rsp_idx]).lower()
                if not _rsp_lower.startswith('daftar isi'):
                    for _bi in range(rsp_idx + 1, j):
                        _bt = body_els[_bi]
                        if _el_tag(_bt) != 'p':
                            continue
                        _bt_txt = _txt(_bt)
                        if _bt_txt and is_roman_start(_bt_txt):
                            return roman_start_p, True  # pgBr di dalam roman zone
                for k in range(j + 1, len(body_els)):
                    nxt  = body_els[k]
                    ntag = _el_tag(nxt)
                    if ntag == 'p':
                        _nxt_txt = _txt(nxt)
                        if _nxt_txt and is_roman_start(_nxt_txt):
                            # Fix B1: pgBr dalam roman zone → jangan advance
                            return roman_start_p, True
                        return nxt, True
                break
        return roman_start_p, False

    @staticmethod
    def _find_roman_page_start(body_els, rsp_idx, W):
        """
        Cari awal halaman yang berisi roman_start_p (body_els[rsp_idx]).

        Pola umum skripsi Indonesia: setiap halaman front matter diawali blok judul
        dokumen yang DIULANG dari halaman sebelumnya (cover repeat, author repeat, dsb.).
        Fungsi ini mendeteksi blok yang diulang dan menyertakannya ke roman zone agar
        tidak terbentuk halaman ekstra dalam cover section.

        Algoritma:
          1. Cari boundary terakhir (sectPr/page-break) sebelum roman_start_p.
          2. Kumpulkan teks yang sudah muncul SEBELUM boundary itu (= texts_before).
          3. Mundur dari roman_start_p: selama paragraf non-kosong teksnya ada di
             texts_before, sertakan dalam roman zone (kandidat page start).
          4. Berhenti saat teks tidak dikenal atau ketemu boundary baru.
        """
        def _get_txt(el):
            return ''.join(t.text or '' for t in el.iter('{%s}t' % W)).strip()

        def _has_boundary(el):
            pPr = el.find("{%s}pPr" % W)
            if pPr is not None and pPr.find("{%s}sectPr" % W) is not None:
                return True
            return any(br.get("{%s}type" % W) == 'page' for br in el.iter("{%s}br" % W))

        roman_start_p = body_els[rsp_idx]

        # Langkah 1: cari boundary terakhir sebelum rsp_idx
        boundary_idx = -1
        for j in range(rsp_idx - 1, -1, -1):
            el = body_els[j]
            if (el.tag.split('}')[-1] if '}' in el.tag else el.tag) != 'p':
                continue
            if _has_boundary(el):
                boundary_idx = j
                break

        if boundary_idx < 0:
            return roman_start_p  # tidak ada boundary → tidak bisa deteksi

        # Langkah 2: kumpulkan teks yang pernah muncul sebelum boundary
        texts_before = set()
        for j in range(boundary_idx + 1):
            el = body_els[j]
            if (el.tag.split('}')[-1] if '}' in el.tag else el.tag) != 'p':
                continue
            txt = _get_txt(el)
            if txt:
                texts_before.add(txt[:50])

        if not texts_before:
            return roman_start_p

        # Langkah 3: mundur dari rsp_idx mencari blok "page header" yang diulang
        candidate = roman_start_p
        for j in range(rsp_idx - 1, max(rsp_idx - 30, -1), -1):
            el  = body_els[j]
            tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
            if tag != 'p':
                continue
            if _has_boundary(el):
                break  # batas section/page — stop
            txt = _get_txt(el)
            if not txt:
                continue  # lewati kosong
            if txt[:50] in texts_before and len(txt.strip()) > 5:
                candidate = el  # teks ini diulang → bagian page header, sertakan
            else:
                break  # teks baru/unik → bukan bagian page header, stop
        return candidate

    # ── Header / Footer helpers ───────────────────────────

    def purge_all_headers_footers(self):
        # Hapus evenAndOddHeaders dari document settings agar semua halaman
        # pakai footer yang sama (dokumen user kadang punya setting ini aktif
        # sehingga even_page_footer lama — sering berisi angka hardcoded —
        # muncul di semua halaman genap output).
        _W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        settings_el = self.doc.settings.element
        for el in list(settings_el.findall(f'{{{_W}}}evenAndOddHeaders')):
            settings_el.remove(el)

        for section in self.doc.sections:
            # Buang referensi even header/footer dari sectPr secara langsung
            # (tanpa mengakses via python-docx agar tidak membuat relationship baru).
            sectPr = section._sectPr
            for ref in list(sectPr.findall(f'{{{_W}}}footerReference')):
                if ref.get(f'{{{_W}}}type') == 'even':
                    sectPr.remove(ref)
            for ref in list(sectPr.findall(f'{{{_W}}}headerReference')):
                if ref.get(f'{{{_W}}}type') == 'even':
                    sectPr.remove(ref)

            # Hanya akses first_page_header/footer jika memang dipakai di section ini.
            # Mengakses part yang tidak dipakai membuat python-docx membuat relationship
            # w:headerReference type="first" yang kemudian jadi yatim piatu saat
            # fmt_roman/fmt_bab_continuation set different_first_page_header_footer=False,
            # sehingga Word melaporkan "unreadable content".
            parts = [section.header, section.footer]
            if section.different_first_page_header_footer:
                parts += [section.first_page_header, section.first_page_footer]
            for part in parts:
                part.is_linked_to_previous = False
                elem = part._element
                for sdt in list(elem.findall(qn('w:sdt'))):
                    elem.remove(sdt)
                paras = list(elem.findall(qn('w:p')))
                if paras:
                    for child in list(paras[0]):
                        paras[0].remove(child)
                    for extra in paras[1:]:
                        elem.remove(extra)
                else:
                    elem.append(OxmlElement('w:p'))

    def sanitize_margins(self):
        """Perbaiki margin yang rusak (bottom < 1cm) agar footer tidak tenggelam.
        Hanya memperbaiki section portrait; landscape dibiarkan.
        Saat bottom diperbaiki, pgMar.footer (footer distance) juga dinormalisasi
        ke 1cm agar footer jelas terlihat di dalam area bottom margin."""
        _W       = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        _MIN     = 567   # 1 cm dalam twips — threshold 'rusak' (< 1cm)
        _FIX     = 1134  # 2 cm dalam twips — nilai perbaikan bottom
        _FTR_STD = 567   # 1 cm dalam twips — footer distance (gap 1cm dari tepi)
        for sec in self.doc.sections:
            sectPr = sec._sectPr
            pgSz   = sectPr.find(f'{{{_W}}}pgSz')
            if pgSz is not None and pgSz.get(f'{{{_W}}}orient') == 'landscape':
                continue
            pgMar = sectPr.find(f'{{{_W}}}pgMar')
            if pgMar is None:
                continue
            # Hanya perbaiki bottom — right dibiarkan agar layout konten tidak berubah.
            try:
                bot = int(pgMar.get(f'{{{_W}}}bottom', _FIX + 1))
            except (ValueError, TypeError):
                bot = _FIX + 1
            if bot < _MIN:
                pgMar.set(f'{{{_W}}}bottom', str(_FIX))
                # Normalisasi footer distance agar posisi footer sesuai bottom baru.
                pgMar.set(f'{{{_W}}}footer', str(_FTR_STD))

    @staticmethod
    def clear_header(section):
        h = section.header
        h.is_linked_to_previous = False
        for p in h.paragraphs:
            DocProcessor.clear_paragraph(p)

    @staticmethod
    def clear_footer(section):
        f = section.footer
        f.is_linked_to_previous = False
        for p in f.paragraphs:
            DocProcessor.clear_paragraph(p)

    @staticmethod
    def set_page_number_format(section, fmt='decimal', start=None):
        sectPr = section._sectPr
        for old in sectPr.findall(qn('w:pgNumType')):
            sectPr.remove(old)
        pgNumType = OxmlElement('w:pgNumType')
        pgNumType.set(qn('w:fmt'), fmt)
        if start is not None:
            pgNumType.set(qn('w:start'), str(start))
        sectPr.append(pgNumType)

    # ── XML section-break helpers ─────────────────────────

    @staticmethod
    def _p_text(p_elem):
        # Kumpulkan teks hanya dari run biasa, lewati subtree drawing/shape
        _SKIP = {qn('w:drawing'), qn('w:pict')}
        result = []
        def _collect(el):
            if el.tag in _SKIP:
                return
            if el.tag == qn('w:t'):
                result.append(el.text or '')
            for child in el:
                _collect(child)
        _collect(p_elem)
        return ''.join(result).strip()

    @staticmethod
    def _p_has_content(p_elem):
        if DocProcessor._p_text(p_elem):
            return True
        if p_elem.find('.//' + qn('w:drawing')) is not None:
            return True
        if p_elem.find('.//' + qn('w:pict')) is not None:
            return True
        if p_elem.find('.//' + qn('w:fldChar')) is not None:
            return True
        return False

    @staticmethod
    def _has_sectPr(p_elem):
        pPr = p_elem.find(qn('w:pPr'))
        return pPr is not None and pPr.find(qn('w:sectPr')) is not None

    def _make_sectPr(self):
        new_sectPr = deepcopy(self.doc.sections[-1]._sectPr)
        for child in list(new_sectPr):
            tag = child.tag
            if tag.endswith('headerReference') or tag.endswith('footerReference'):
                new_sectPr.remove(child)
        type_elem = new_sectPr.find(qn('w:type'))
        if type_elem is None:
            type_elem = OxmlElement('w:type')
            new_sectPr.append(type_elem)
        type_elem.set(qn('w:val'), 'nextPage')
        # Normalisasi landscape → portrait agar section break baru tidak mewarisi
        # orientasi landscape dari template (body-level sectPr dokumen mungkin landscape).
        pgSz = new_sectPr.find(qn('w:pgSz'))
        if pgSz is not None and pgSz.get(qn('w:orient'), '') == 'landscape':
            w = pgSz.get(qn('w:w'), '11910')
            h = pgSz.get(qn('w:h'), '16840')
            pgSz.set(qn('w:w'), h)
            pgSz.set(qn('w:h'), w)
            del pgSz.attrib[qn('w:orient')]
        return new_sectPr

    def _attach_sectPr(self, p_elem):
        pPr = p_elem.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            p_elem.insert(0, pPr)
        pPr.append(self._make_sectPr())

    def _purge_continuous_sectPr_in_bab_zone(self, first_bab_p):
        body     = self.doc.element.body
        children = list(body)
        try:
            start_idx = children.index(first_bab_p)
        except ValueError:
            return
        for elem in children[start_idx:]:
            if not elem.tag.endswith('}p'):
                continue
            pPr = elem.find(qn('w:pPr'))
            if pPr is None:
                continue
            sectPr = pPr.find(qn('w:sectPr'))
            if sectPr is None:
                continue
            type_e = sectPr.find(qn('w:type'))
            if type_e is None or type_e.get(qn('w:val')) != 'continuous':
                continue
            pgsz = sectPr.find(qn('w:pgSz'))
            if pgsz is not None and pgsz.get(qn('w:orient')) == 'landscape':
                continue
            pPr.remove(sectPr)

    def _strip_empty_paras_before_bab(self, target_p):
        body     = self.doc.element.body
        children = list(body)
        try:
            tgt_idx = children.index(target_p)
        except ValueError:
            return
        to_remove = []
        for j in range(tgt_idx - 1, -1, -1):
            elem = children[j]
            if not elem.tag.endswith('}p'):
                break
            pPr = elem.find(qn('w:pPr'))
            if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
                break
            if self._p_has_content(elem):
                break
            to_remove.append(elem)
        for elem in to_remove:
            body.remove(elem)

    def _remove_page_breaks_before(self, target_p):
        body     = self.doc.element.body
        children = list(body)
        idx      = children.index(target_p)
        for br in target_p.findall('.//' + qn('w:br')):
            if br.get(qn('w:type')) == 'page':
                parent = br.getparent()
                if parent is not None:
                    parent.remove(br)
        for j in range(idx - 1, -1, -1):
            elem = children[j]
            if not elem.tag.endswith('}p'):
                break
            if self._p_has_content(elem):
                break
            for br in elem.findall('.//' + qn('w:br')):
                if br.get(qn('w:type')) == 'page':
                    parent = br.getparent()
                    if parent is not None:
                        parent.remove(br)

    def insert_break_before_xml(self, target_p, allow_empty_boundary=False):
        self._remove_page_breaks_before(target_p)
        body     = self.doc.element.body
        children = list(body)
        tgt_idx  = children.index(target_p)
        j = tgt_idx - 1
        while j >= 0:
            elem = children[j]
            tag  = elem.tag
            if tag.endswith('}p'):
                pPr = elem.find(qn('w:pPr'))
                if pPr is not None:
                    sectPr = pPr.find(qn('w:sectPr'))
                    if sectPr is not None:
                        if sectPr.find(qn('w:pgSz')) is not None:
                            # Kalau sectPr ini pada paragraf berisi konten (artinya
                            # kita sendiri yang pasang untuk batas roman zone) dan ada
                            # paragraf kosong antara sini dan target, pakai paragraf
                            # kosong itu sebagai batas BAB section baru.
                            if allow_empty_boundary and self._p_has_content(elem):
                                for k in range(j + 1, tgt_idx):
                                    nxt = children[k]
                                    if (nxt.tag.endswith('}p') and
                                            not self._p_has_content(nxt) and
                                            not self._has_sectPr(nxt)):
                                        self._attach_sectPr(nxt)
                                        return
                            return
                        else:
                            pPr.remove(sectPr)
                            if self._p_has_content(elem):
                                self._attach_sectPr(elem)
                                return
                            j -= 1
                            continue
                if self._p_has_content(elem):
                    self._attach_sectPr(elem)
                    return
            elif tag.endswith('}tbl') or tag.endswith('}sdt'):
                new_p   = OxmlElement('w:p')
                new_pPr = OxmlElement('w:pPr')
                new_p.append(new_pPr)
                new_pPr.append(self._make_sectPr())
                body.insert(tgt_idx, new_p)
                return
            j -= 1

    def _find_section_start(self, target_p):
        body     = self.doc.element.body
        children = list(body)
        try:
            idx = children.index(target_p)
        except ValueError:
            return target_p
        for j in range(idx - 1, -1, -1):
            elem = children[j]
            if elem.tag.endswith('}p') and self._has_sectPr(elem):
                for k in range(j + 1, idx + 1):
                    c = children[k]
                    if c.tag.endswith('}p'):
                        return c
                break
        return target_p

    @staticmethod
    def _has_page_break_before(paras, from_idx, to_idx):
        for para in paras[from_idx:to_idx]:
            if DocProcessor._has_sectPr(para._p):
                return True
            for br in para._p.findall('.//' + qn('w:br')):
                if br.get(qn('w:type')) == 'page':
                    return True
        return False

    @staticmethod
    def _para_has_page_break_before(para):
        pPr = para._p.find(qn('w:pPr'))
        if pPr is None:
            return False
        pb = pPr.find(qn('w:pageBreakBefore'))
        if pb is None:
            return False
        return pb.get(qn('w:val'), 'true') not in ('false', '0')

    # ── Phase pipeline ────────────────────────────────────

    def scan_zones(self):
        """Phase 1: Scan paragraphs. Returns (roman_start_p, bab_p_list)."""
        roman_start_p      = None
        bab_p_list         = []
        lampiran_found     = False
        last_bab_para_idx  = None
        found_numbered_bab = False
        # bab_num_key → {'idx': int, 'has_break': bool, 'is_heading': bool}
        # Memungkinkan penggantian entri lama (TOC palsu) dengan entri nyata
        seen_bab_info      = {}
        inside_toc         = False
        toc_start_idx      = -1   # index paragraf saat inside_toc terakhir di-set True
        all_paras          = list(self.doc.paragraphs)

        for para_idx, para in enumerate(all_paras):
            text  = para.text.strip()
            lower = text.lower()

            if is_toc_heading(lower) or _has_toc_field(para._p):
                if not inside_toc:
                    toc_start_idx = para_idx  # hanya catat saat pertama masuk TOC
                inside_toc = True
                if roman_start_p is None:
                    roman_start_p = para._p
                continue

            if inside_toc:
                # [Fix A] Keluar TOC jika ada section break / page break SETELAH
                # toc_start_idx dalam window 15 paragraf ke belakang.
                # Window 15 agar page break yang beberapa paragraf sebelum BAB
                # (misal di paragraf kosong) tetap terdeteksi.
                prev_has_break = self._has_page_break_before(
                    all_paras,
                    max(toc_start_idx + 1, para_idx - 15),
                    para_idx
                )
                if prev_has_break or self._para_has_page_break_before(para):
                    inside_toc = False
                    # Lanjut ke pemrosesan normal (tidak continue)
                elif not text or is_toc_entry(text) or (
                    # [Fix E] Entri sub-bab bernomor seperti "2.1 Sistem", "3.5triangulasi"
                    # adalah bagian dari TOC — HANYA jika bukan Heading style.
                    # Guard heading: "1.2 Rumusan masalah" (Heading 2) adalah konten nyata, bukan TOC.
                    re.match(r'^\d+\.\d', text) and
                    not re.search(r'heading', para.style.name.lower() if para.style else '')
                ):
                    continue
                else:
                    # [Fix D] Paragraf ini sendiri adalah BAB heading tanpa
                    # page break / Heading style → kemungkinan entri TOC tanpa
                    # nomor halaman → tetap dalam TOC.
                    _cur_style = (para.style.name.lower() if para.style else "")
                    if (is_bab_heading(text)
                            and not re.search(r'heading', _cur_style)
                            and not self._para_has_page_break_before(para)):
                        continue

                    # Cek lookahead: kalau dalam 8 paragraf ke depan masih ada
                    # entri TOC (punya toc style atau nomor halaman), tetap di TOC.
                    still_in_toc = False
                    for _lk in range(para_idx + 1, min(para_idx + 9, len(all_paras))):
                        _lp = all_paras[_lk]
                        _ls = (_lp.style.name.lower() if _lp.style else "")
                        _lt = _lp.text.strip()
                        if 'toc' in _ls:
                            still_in_toc = True
                            break
                        if _lt and is_toc_entry(_lt):
                            still_in_toc = True
                            break
                        # [Fix E] Sub-bab bernomor (e.g. "2.1", "4.2.1") = sinyal masih di TOC
                        # Hanya jika bukan Heading style — Heading 2 "1.2 Judul" adalah konten nyata
                        if _lt and re.match(r'^\d+\.\d', _lt) and not re.search(r'heading', _ls):
                            still_in_toc = True
                            break
                        if _lt and is_bab_heading(_lt):
                            # [Fix B] Hanya keluar TOC jika BAB heading ini punya
                            # indikator nyata (Heading style atau page break).
                            _lp_heading = bool(re.search(r'heading', _ls))
                            _lp_has_brk = self._para_has_page_break_before(_lp)
                            if _lp_heading or _lp_has_brk:
                                break  # BAB heading nyata → keluar TOC
                            still_in_toc = True
                            break
                    if still_in_toc:
                        continue
                    inside_toc = False

            if not text:
                continue

            if roman_start_p is None and is_roman_start(text):
                roman_start_p = para._p

            if is_bab_heading(text) and not is_false_bab(para):
                is_numbered_bab = BAB_HEAD_RE.match(text) is not None
                if not is_numbered_bab and not found_numbered_bab:
                    continue
                if re.match(r'^\s*lampiran', text, re.IGNORECASE):
                    if lampiran_found:
                        continue
                    lampiran_found = True

                _style_lower  = (para.style.name.lower() if para.style else "")
                _is_heading   = bool(re.search(r'heading', _style_lower))
                _has_brk_para = self._para_has_page_break_before(para)

                if is_numbered_bab:
                    found_numbered_bab = True
                    m_num = BAB_HEAD_RE.match(text)
                    bab_num_key = m_num.group(2).strip().lower() if m_num else None
                    if bab_num_key and bab_num_key in seen_bab_info:
                        # [Fix C] Ganti entri lama jika yang baru lebih terpercaya
                        # (punya page break / Heading style, sedangkan lama tidak).
                        # Gunakan _has_page_break_before (cek sectPr + pgbr dalam window)
                        # agar section break sebelum BAB nyata ikut terdeteksi.
                        old = seen_bab_info[bab_num_key]
                        if last_bab_para_idx is not None:
                            _win_c = max(last_bab_para_idx + 1, para_idx - 15)
                            _full_brk = (self._has_page_break_before(all_paras, _win_c, para_idx)
                                         or _has_brk_para)
                        else:
                            _full_brk = _has_brk_para
                        if (_full_brk or _is_heading) and not (old['has_break'] or old['is_heading']):
                            bab_p_list[old['idx']] = para._p
                            old['has_break']  = _full_brk
                            old['is_heading'] = _is_heading
                            last_bab_para_idx = para_idx  # update untuk window berikutnya
                        continue  # duplikat sudah ditangani
                else:
                    # [Fix C-ext] Deduplikasi DAFTAR PUSTAKA / LAMPIRAN sama seperti BAB bernomor.
                    # Mencegah entri TOC palsu (tanpa page break) bertahan saat entri asli ditemukan.
                    _ep = re.match(
                        r'^\s*(daftar\s*pust?aka|lampiran|appendix|referensi|references?|bibliography)',
                        text, re.IGNORECASE
                    )
                    bab_num_key = re.sub(r'\s+', '', _ep.group(1).lower()) if _ep else None
                    if bab_num_key and bab_num_key in seen_bab_info:
                        old = seen_bab_info[bab_num_key]
                        if last_bab_para_idx is not None:
                            _win_c = max(last_bab_para_idx + 1, para_idx - 15)
                            _full_brk = (self._has_page_break_before(all_paras, _win_c, para_idx)
                                         or _has_brk_para)
                        else:
                            _full_brk = _has_brk_para
                        if (_full_brk or _is_heading) and not (old['has_break'] or old['is_heading']):
                            bab_p_list[old['idx']] = para._p
                            old['has_break']  = _full_brk
                            old['is_heading'] = _is_heading
                            last_bab_para_idx = para_idx
                        continue
                if last_bab_para_idx is not None:
                    _window = max(last_bab_para_idx + 1, para_idx - 15)
                    has_break = self._has_page_break_before(all_paras, _window, para_idx)
                    if not has_break:
                        has_break = _has_brk_para
                    # "Heading X" style + page break → langsung dipercaya.
                    # Non-heading style + page break → tetap cek forward_count karena sectPr
                    # di body text (misal Sistematika Penulisan) bisa memicu false positive.
                    _trusted = has_break and _is_heading
                    if not _trusted:
                        # Lampiran & Daftar Pustaka dibebaskan dari forward_count
                        _is_endpoint = bool(re.match(
                            r'^\s*(lampiran|daftar\s+pust?aka)', text, re.IGNORECASE
                        ))
                        if not _is_endpoint:
                            forward_count = 0
                            for _k in range(para_idx + 1, len(all_paras)):
                                _nt = all_paras[_k].text.strip()
                                if is_bab_heading(_nt) and not is_false_bab(all_paras[_k]):
                                    break
                                if _nt:
                                    forward_count += 1
                            # Threshold lebih rendah (2) jika ada page break, karena
                            # Sistematika entries hanya punya 1 kalimat per BAB.
                            _threshold = 2 if has_break else 5
                            if forward_count < _threshold:
                                continue
                bab_p_list.append(para._p)
                if bab_num_key:
                    seen_bab_info[bab_num_key] = {
                        'idx':       len(bab_p_list) - 1,
                        'has_break': _has_brk_para,
                        'is_heading': _is_heading,
                    }
                last_bab_para_idx = para_idx

            # Fallback: Heading 1 langsung setelah section break, bukan roman-zone heading.
            # Menangani dokumen di mana BAB I/II/III tidak punya prefix "BAB" pada teksnya
            # (autonumber dari Word list numbering — teks hanya "PENDAHULUAN" dll).
            # Guard found_numbered_bab dihapus agar BAB II, III, dst. yang juga autonumber
            # tetap terdeteksi. Pengaman utama: wajib Heading 1 + sectPr pada paragraf
            # sebelumnya sehingga sub-bab biasa tidak ikut terdeteksi.
            elif (roman_start_p is not None and text and not is_roman_start(text) and
                  not lampiran_found):
                _style_lower = para.style.name.lower() if para.style else ""
                if re.match(r'^heading\s*1$', _style_lower):
                    _prev_has_sect = (
                        para_idx > 0 and
                        self._has_sectPr(all_paras[para_idx - 1]._p)
                    )
                    if _prev_has_sect:
                        bab_p_list.append(para._p)
                        found_numbered_bab = True
                        last_bab_para_idx = para_idx

        # Fallback: cari roman_start_p di dalam sel tabel jika belum ditemukan.
        # Beberapa template cover universitas memakai tabel untuk layout,
        # sehingga keyword seperti "KATA PENGANTAR" ada di w:tc bukan w:p biasa.
        if roman_start_p is None:
            W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            body_ch = list(self.doc.element.body)
            for elem in body_ch:
                if not elem.tag.endswith('}tbl'):
                    continue
                found_in_tbl = False
                for tc in elem.iter('{%s}tc' % W):
                    for p_el in tc.findall('{%s}p' % W):
                        cell_text = ''.join(
                            t.text or '' for t in p_el.iter('{%s}t' % W)
                        ).strip()
                        if is_roman_start(cell_text):
                            found_in_tbl = True
                            break
                    if found_in_tbl:
                        break
                if found_in_tbl:
                    # Gunakan paragraf pertama setelah tabel ini sebagai roman_start_p
                    tbl_idx = body_ch.index(elem)
                    for k in range(tbl_idx + 1, len(body_ch)):
                        if body_ch[k].tag.endswith('}p'):
                            roman_start_p = body_ch[k]
                            break
                    break

        # Pastikan urutan bab_p_list sesuai posisi di dokumen.
        # Fix C (replacement) bisa membuat entry tidak berurutan.
        _body = list(self.doc.element.body)
        def _body_pos(p):
            try: return _body.index(p)
            except ValueError: return len(_body)
        bab_p_list.sort(key=_body_pos)

        return roman_start_p, bab_p_list

    def insert_breaks(self, roman_start_p, bab_p_list, exact_roman_start=False):
        """Phase 1.5 + 2: Purge continuous sectPr, insert zone breaks.
        Returns updated roman_start_p (may differ after _find_section_start).
        exact_roman_start=True: skip _find_section_start (untuk multi-cover advance)."""
        if bab_p_list:
            self._purge_continuous_sectPr_in_bab_zone(bab_p_list[0])

        if roman_start_p is not None:
            if not exact_roman_start:
                roman_start_p = self._find_section_start(roman_start_p)
            # Jika section boundary sudah ada tepat sebelum roman_start_p,
            # tidak perlu insert break (agar page break di roman zone tidak dihapus).
            body_ch = list(self.doc.element.body)
            _rsp_idx = next((i for i, e in enumerate(body_ch) if e is roman_start_p), -1)
            _roman_already_bounded = (
                _rsp_idx > 0 and
                body_ch[_rsp_idx - 1].tag.endswith('}p') and
                self._has_sectPr(body_ch[_rsp_idx - 1])
            )
            if not _roman_already_bounded:
                self.insert_break_before_xml(roman_start_p)
                self._strip_empty_paras_before_bab(roman_start_p)

        for bab_p in bab_p_list:
            self.insert_break_before_xml(bab_p, allow_empty_boundary=exact_roman_start)
            self._strip_empty_paras_before_bab(bab_p)

        return roman_start_p  # dikembalikan karena bisa berubah

    def build_section_map(self, roman_start_p, bab_p_list):
        """Phase 3: Build section boundary map.
        Returns (roman_sec, bab_sec_list, n_sections)."""
        breaks = []
        for para in self.doc.paragraphs:
            if self._has_sectPr(para._p):
                breaks.append(para._p)

        para_list   = list(self.doc.paragraphs)
        para_p_list = [p._p for p in para_list]
        break_indices = [para_p_list.index(bp) for bp in breaks]
        boundaries    = [0] + [bi + 1 for bi in break_indices]
        n_sections    = len(boundaries)

        def para_to_sec(p_elem):
            try:
                idx = para_p_list.index(p_elem)
            except ValueError:
                return n_sections - 1
            for s in range(n_sections):
                start = boundaries[s]
                end   = boundaries[s + 1] if s + 1 < n_sections else len(para_p_list)
                if start <= idx < end:
                    return s
            return n_sections - 1

        roman_sec    = para_to_sec(roman_start_p) if roman_start_p is not None else None
        bab_sec_list = [para_to_sec(p) for p in bab_p_list]
        return roman_sec, bab_sec_list, n_sections

    # ── Section formatters (dipakai oleh paket2 & paket3) ─

    def fmt_cover(self, section, first_cover=False, show_pos=None, cover_start=1,
                  visible_pos=None):
        """
        show_pos    : (align, top) saat hidden_cov='Tidak' — semua halaman cover tampil.
        visible_pos : (align, top) saat hidden_cov='Ya' — posisi nomor di halaman cover 2+.
                      Jika None, default CENTER BAWAH.
        """
        _vis_align = WD_ALIGN_PARAGRAPH.CENTER
        _vis_top   = False
        if visible_pos:
            _vis_align, _vis_top = visible_pos
        elif show_pos:
            _vis_align, _vis_top = show_pos

        self.clear_header(section)
        self.clear_footer(section)
        if first_cover:
            self.set_page_number_format(section, 'lowerRoman', cover_start)
            if show_pos:
                # hidden_cov='Tidak': tampilkan nomor di semua halaman cover
                section.different_first_page_header_footer = False
                align, top = show_pos
                if top:
                    self._place_num_in_part(section.header, align)
                else:
                    self._place_num_in_part(section.footer, align)
            else:
                # hidden_cov='Ya': sembunyikan cover 1 via first-page footer.
                # Regular footer hanya diisi jika visible_pos eksplisit diberikan.
                # Jika visible_pos=None, regular footer kosong → semua halaman di
                # section ini (termasuk cover 2 yang overflow konten) tersembunyi.
                section.different_first_page_header_footer = True
                fph = section.first_page_header
                fph.is_linked_to_previous = False
                for p in fph.paragraphs:
                    self.clear_paragraph(p)
                fpf = section.first_page_footer
                fpf.is_linked_to_previous = False
                for p in fpf.paragraphs:
                    self.clear_paragraph(p)
                if visible_pos:
                    if _vis_top:
                        self._place_num_in_part(section.header, _vis_align)
                    else:
                        self._place_num_in_part(section.footer, _vis_align)
                # else: regular footer kosong → cover 2+ tersembunyi
        else:
            section.different_first_page_header_footer = False
            self.set_page_number_format(section, 'lowerRoman')
            # Sampul ke-2 dst
            if show_pos:
                # hidden_cov='Tidak': tampilkan nomor di semua halaman cover
                align, top = show_pos
                if top:
                    self._place_num_in_part(section.header, align)
                else:
                    self._place_num_in_part(section.footer, align)
            elif visible_pos:
                # hidden_cov='Ya' dengan posisi eksplisit: tampilkan nomor
                if _vis_top:
                    self._place_num_in_part(section.header, _vis_align)
                else:
                    self._place_num_in_part(section.footer, _vis_align)
            # else: hidden_cov='Ya', visible_pos=None → footer kosong = nomor tersembunyi

    def fmt_roman(self, section, start=None):
        section.different_first_page_header_footer = False
        self.clear_header(section)
        f = section.footer
        f.is_linked_to_previous = False
        p = self._first_para(f)
        self.clear_paragraph(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.add_page_number(p)
        self._set_pn_spacing(p)
        self.set_page_number_format(section, 'lowerRoman', start)

    def fmt_bab_first(self, section, reset_to_1=False):
        section.different_first_page_header_footer = True
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(1.25)
        self.clear_header(section)
        self.clear_footer(section)

        fph = section.first_page_header
        fph.is_linked_to_previous = False
        for p in fph.paragraphs:
            self.clear_paragraph(p)

        fpf = section.first_page_footer
        fpf.is_linked_to_previous = False
        for p in fpf.paragraphs:
            self.clear_paragraph(p)
        p = self._first_para(fpf)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.add_page_number(p)
        self._set_pn_spacing(p)

        h = section.header
        h.is_linked_to_previous = False
        p = self._first_para(h)
        self.clear_paragraph(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self.add_page_number(p)
        self._set_pn_spacing(p)

        if reset_to_1:
            self.set_page_number_format(section, 'decimal', 1)
        else:
            self.set_page_number_format(section, 'decimal')

    def fmt_bab_continuation(self, section):
        section.different_first_page_header_footer = False
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(1.25)
        self.clear_header(section)
        self.clear_footer(section)
        h = section.header
        h.is_linked_to_previous = False
        p = self._first_para(h)
        self.clear_paragraph(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self.add_page_number(p)
        self._set_pn_spacing(p)
        self.set_page_number_format(section, 'decimal')

    def fmt_uniform_section(self, section, align, top, fmt='decimal', start=None):
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(1.25)
        section.different_first_page_header_footer = False
        self.clear_header(section)
        self.clear_footer(section)
        if top:
            self._place_num_in_part(section.header, align)
        else:
            self._place_num_in_part(section.footer, align)
        self.set_page_number_format(section, fmt, start)
