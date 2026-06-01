import sys
sys.path.insert(0, r'd:\Freelaces\Server\htdocs\adk\python')
from docx import Document
from utils import DocProcessor, is_bab_heading, is_false_bab, is_roman_start, is_toc_heading, is_toc_entry, _has_toc_field
import re

path = r'd:\Freelaces\Server\htdocs\adk\test_files\paket3\cover 1 dimulai dari ii\Docx 19.docx'
doc  = Document(path)
BAB_HEAD_RE = re.compile(r'^\s*(bab|chapter)\s+([IVXLCDM]+|\d+)\b', re.IGNORECASE)
proc        = DocProcessor(doc, 'Times New Roman', 12)
all_paras   = list(doc.paragraphs)

# ── 1. Apa yang memicu inside_toc sebelum para[65] ─────────────────
print("=== [1] PEMICU inside_toc SEBELUM para[65] ===\n")
for i in range(0, 66):
    p    = all_paras[i]
    text = p.text.strip()
    lower = text.lower()
    if is_toc_heading(lower) or _has_toc_field(p._p):
        print(f"  para[{i:3d}] TOC_TRIGGER  text={repr(text[:60])}")
        print(f"           style={repr(p.style.name)}")

# ── 2. Mengapa inside_toc tidak keluar sebelum para[65] ────────────
print()
print("=== [2] PARAGRAF 50-65 DETAIL ===\n")
inside_toc    = False
toc_start_idx = -1
for i in range(0, 66):
    p    = all_paras[i]
    text = p.text.strip()
    lower = text.lower()
    if is_toc_heading(lower) or _has_toc_field(p._p):
        if not inside_toc:
            toc_start_idx = i
        inside_toc = True
        if i >= 50:
            print(f"  [{i:3d}] SET_TOC  text={repr(text[:60])}")
        continue
    if inside_toc and i >= 50:
        prev_break = proc._has_page_break_before(all_paras, max(toc_start_idx+1, i-15), i)
        brk_para   = proc._para_has_page_break_before(p)
        is_entry   = is_toc_entry(text) if text else False
        is_sub     = bool(re.match(r'^\d+\.\d', text) and not re.search(r'heading', p.style.name.lower() if p.style else ''))
        is_bab     = is_bab_heading(text) and not is_false_bab(p)
        print(f"  [{i:3d}] inside_toc=True  prev_break={prev_break}  brk_para={brk_para}  "
              f"is_entry={is_entry}  is_sub={is_sub}  is_bab={is_bab}  "
              f"style={repr(p.style.name[:15])}  text={repr(text[:40])}")

# ── 3. Full trace BAB IV ────────────────────────────────────────────
print()
print("=== [3] FULL TRACE AREA BAB IV (para 240-260) ===\n")

# Re-simulasi penuh sampai para 260 untuk tahu state saat BAB IV ditemukan
inside_toc2       = False
toc_start_idx2    = -1
found_numbered2   = False
seen_bab2         = {}
bab_list2         = []
last_bab_idx2     = None

for para_idx, para in enumerate(all_paras[:261]):
    text  = para.text.strip()
    lower = text.lower()

    if is_toc_heading(lower) or _has_toc_field(para._p):
        if not inside_toc2:
            toc_start_idx2 = para_idx
        inside_toc2 = True
        continue

    if inside_toc2:
        prev_break = proc._has_page_break_before(all_paras, max(toc_start_idx2+1, para_idx-15), para_idx)
        brk_para   = proc._para_has_page_break_before(para)
        if prev_break or brk_para:
            inside_toc2 = False
        elif not text or is_toc_entry(text) or (
            re.match(r'^\d+\.\d', text) and
            not re.search(r'heading', para.style.name.lower() if para.style else '')
        ):
            continue
        else:
            _cs = (para.style.name.lower() if para.style else "")
            if is_bab_heading(text) and not re.search(r'heading', _cs) and not brk_para:
                continue
            still = False
            for _lk in range(para_idx+1, min(para_idx+9, len(all_paras))):
                _lp = all_paras[_lk]; _ls = (_lp.style.name.lower() if _lp.style else "")
                _lt = _lp.text.strip()
                if 'toc' in _ls or is_toc_entry(_lt): still = True; break
                if _lt and re.match(r'^\d+\d', _lt) and not re.search(r'heading', _ls): still = True; break
            if still: continue
            inside_toc2 = False

    if not text: continue

    if is_bab_heading(text) and not is_false_bab(para):
        is_num = BAB_HEAD_RE.match(text) is not None
        if not is_num and not found_numbered2: continue
        _sl   = (para.style.name.lower() if para.style else "")
        _head = bool(re.search(r'heading', _sl))
        _brk  = proc._para_has_page_break_before(para)

        if is_num:
            found_numbered2 = True
            m    = BAB_HEAD_RE.match(text)
            key  = m.group(2).strip().lower() if m else None

            if key and key in seen_bab2:
                old = seen_bab2[key]
                if last_bab_idx2 is not None:
                    _wc = max(last_bab_idx2+1, para_idx-15)
                    _fb = proc._has_page_break_before(all_paras, _wc, para_idx) or _brk
                else:
                    _fb = _brk
                replace = (_fb or _head) and not (old['has_break'] or old['is_heading'])
                if para_idx >= 240:
                    print(f"  [{para_idx:3d}] DUPLICATE key={key}  replace={replace}  "
                          f"full_brk={_fb}  is_head={_head}  old={old}")
                if replace:
                    bab_list2[old['idx']] = para._p
                    old['has_break'] = _fb; old['is_heading'] = _head
                    last_bab_idx2 = para_idx
                continue

            # Baru — cek window break
            if last_bab_idx2 is not None:
                _wc  = max(last_bab_idx2+1, para_idx-15)
                _fb  = proc._has_page_break_before(all_paras, _wc, para_idx) or _brk
            else:
                _fb  = _brk

            if para_idx >= 240:
                print(f"  [{para_idx:3d}] NEW_BAB key={key}  full_brk={_fb}  is_head={_head}  text={repr(text[:40])}")
            seen_bab2[key] = {'idx': len(bab_list2), 'has_break': _fb, 'is_heading': _head}
            bab_list2.append(para._p)
            last_bab_idx2 = para_idx

print()
print(f"bab_list2: {[doc.paragraphs[all_paras.index(doc.paragraphs[i])] .text[:30] if hasattr(doc.paragraphs[i], 'text') else '?' for i in range(len(bab_list2))]}")
# Simpler: just get text from XML
from docx.oxml.ns import qn as _qn
def _txt(p_elem):
    return ''.join(t.text or '' for t in p_elem.iter(_qn('w:t'))).strip()[:30]
print(f"bab_list2 texts: {[_txt(p) for p in bab_list2]}")
