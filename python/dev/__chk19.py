import sys
sys.path.insert(0, r'd:\Freelaces\Server\htdocs\adk\python')
from docx import Document
from docx.oxml.ns import qn

def read_page_nums(docx_path):
    doc = Document(docx_path)
    result = []
    for sec in doc.sections:
        pn    = sec._sectPr.find(qn('w:pgNumType'))
        fmt   = pn.get(qn('w:fmt'),   'decimal') if pn is not None else 'decimal'
        start = pn.get(qn('w:start'), None)      if pn is not None else None
        result.append({'fmt': fmt, 'start': start})
    return result

path = r'd:\Freelaces\Server\htdocs\adk\test_files\paket3\cover 1 dimulai dari ii\hasil\Docx 19.docx'

# Page numbering per section
nums = read_page_nums(path)
print("=== PAGE NUMBERING ===")
for i, n in enumerate(nums):
    fmt   = n.get('fmt', 'NONE')
    start = n.get('start', 'null')
    print(f"Section {i:2d}: fmt={fmt:14s}  start={start}")

# Section count vs detected sections
print()
print("=== SECTION COUNT ===")
doc = Document(path)
print(f"doc.sections count: {len(doc.sections)}")

# First content of each section
print()
print("=== SECTIONS FIRST CONTENT ===")
paras = list(doc.paragraphs)
from utils import DocProcessor
breaks = [i for i, p in enumerate(paras) if DocProcessor._has_sectPr(p._p)]
bounds = [0] + [b + 1 for b in breaks]
for si, start in enumerate(bounds):
    end = bounds[si + 1] if si + 1 < len(bounds) else len(paras)
    first = ''
    for i in range(start, min(start + 15, end)):
        t = paras[i].text.strip()
        if t:
            first = repr(t[:60])
            break
    total_paras = end - start
    print(f"Section {si:2d} ({total_paras:3d} paras): {first}")
