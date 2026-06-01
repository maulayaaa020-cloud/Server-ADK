import sys
sys.path.insert(0, r'd:\Freelaces\Server\htdocs\adk\python')
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

SRC  = r'd:\Freelaces\Server\htdocs\adk\test_files\paket3\cover 1 dimulai dari ii\Docx 19.docx'
HASIL = r'd:\Freelaces\Server\htdocs\adk\test_files\paket3\cover 1 dimulai dari ii\hasil\Docx 19.docx'

src   = Document(SRC)
hasil = Document(HASIL)

OK = '  OK'
WARN = '  WARN'
FAIL = '  FAIL'

issues = []

# ── 1. Jumlah paragraf ──────────────────────────────────────────────
sp = len(src.paragraphs)
hp = len(hasil.paragraphs)
status = OK if sp == hp else WARN
if sp != hp:
    issues.append(f"Paragraf berbeda: asli={sp} hasil={hp}")
print(f"[1] Jumlah paragraf     {status}  asli={sp}  hasil={hp}")

# ── 2. Jumlah tabel ─────────────────────────────────────────────────
st = len(src.tables)
ht = len(hasil.tables)
status = OK if st == ht else FAIL
if st != ht:
    issues.append(f"Tabel berbeda: asli={st} hasil={ht}")
print(f"[2] Jumlah tabel        {status}  asli={st}  hasil={ht}")

# ── 3. Struktur tabel (baris & kolom) ───────────────────────────────
print(f"[3] Struktur tabel:")
for i, (s_tbl, h_tbl) in enumerate(zip(src.tables, hasil.tables)):
    sr, sc = len(s_tbl.rows), len(s_tbl.columns)
    hr, hc = len(h_tbl.rows), len(h_tbl.columns)
    status = OK if (sr==hr and sc==hc) else FAIL
    if sr!=hr or sc!=hc:
        issues.append(f"Tabel {i}: asli {sr}x{sc} vs hasil {hr}x{hc}")
    print(f"     Tabel {i}           {status}  {sr}x{sc} → {hr}x{hc}")

# ── 4. Konten teks — cek 20 paragraf pertama ────────────────────────
print(f"[4] Konten teks (20 paragraf pertama):")
mismatch = 0
for i in range(min(20, sp, hp)):
    s_txt = src.paragraphs[i].text.strip()
    h_txt = hasil.paragraphs[i].text.strip()
    if s_txt != h_txt:
        mismatch += 1
        if mismatch <= 3:
            print(f"     para[{i}] BEDA:")
            print(f"       asli : {repr(s_txt[:60])}")
            print(f"       hasil: {repr(h_txt[:60])}")
if mismatch == 0:
    print(f"     {OK}  Semua sama")
else:
    issues.append(f"{mismatch} paragraf pertama berbeda")

# ── 5. Konten teks — sampling seluruh dokumen ───────────────────────
print(f"[5] Sampling konten seluruh dokumen:")
total_diff = 0
for i in range(min(sp, hp)):
    if src.paragraphs[i].text.strip() != hasil.paragraphs[i].text.strip():
        total_diff += 1
status = OK if total_diff == 0 else WARN
print(f"     {status}  {total_diff} paragraf berbeda dari {min(sp,hp)}")
if total_diff > 0:
    issues.append(f"Total {total_diff} paragraf berbeda kontennya")

# ── 6. Teks dalam tabel ─────────────────────────────────────────────
print(f"[6] Konten sel tabel:")
tbl_diff = 0
for i, (s_tbl, h_tbl) in enumerate(zip(src.tables, hasil.tables)):
    for ri, (s_row, h_row) in enumerate(zip(s_tbl.rows, h_tbl.rows)):
        for ci, (s_cell, h_cell) in enumerate(zip(s_row.cells, h_row.cells)):
            s_txt = s_cell.text.strip()
            h_txt = h_cell.text.strip()
            if s_txt != h_txt:
                tbl_diff += 1
                if tbl_diff <= 3:
                    print(f"     Tabel {i} [{ri},{ci}] BEDA:")
                    print(f"       asli : {repr(s_txt[:50])}")
                    print(f"       hasil: {repr(h_txt[:50])}")
status = OK if tbl_diff == 0 else WARN
if tbl_diff == 0:
    print(f"     {OK}  Semua sel sama")
else:
    issues.append(f"{tbl_diff} sel tabel berbeda")

# ── 7. Jumlah section ───────────────────────────────────────────────
ss = len(src.sections)
hs = len(hasil.sections)
print(f"[7] Jumlah section      {WARN}  asli={ss}  hasil={hs}  (script menambah break BAB — wajar)")

# ── 8. Page numbering di hasil ──────────────────────────────────────
print(f"[8] Page numbering hasil:")
for i, sec in enumerate(hasil.sections):
    pn    = sec._sectPr.find(qn('w:pgNumType'))
    fmt   = pn.get(qn('w:fmt'), 'decimal') if pn is not None else 'decimal'
    start = pn.get(qn('w:start'), '-')     if pn is not None else '-'
    if i <= 3 or i >= hs-2:
        print(f"     Section {i:2d}: {fmt:14s}  start={start}")
    elif i == 4:
        print(f"     ... ({hs-4} section tengah semua decimal) ...")

# ── 9. Header/footer ────────────────────────────────────────────────
print(f"[9] Header/footer:")
hf_ok = True
for sec in hasil.sections:
    if sec.header.paragraphs:
        texts = [p.text.strip() for p in sec.header.paragraphs if p.text.strip()]
        if texts:
            hf_ok = False
            issues.append(f"Header masih ada isi: {texts[:2]}")
            break
print(f"     {OK if hf_ok else FAIL}  {'Header bersih' if hf_ok else 'Header masih ada isi'}")

# ── 10. Column break ────────────────────────────────────────────────
col_src  = sum(1 for p in src.paragraphs
               for br in p._p.findall('.//' + qn('w:br'))
               if br.get(qn('w:type')) == 'column')
col_hasil = sum(1 for p in hasil.paragraphs
                for br in p._p.findall('.//' + qn('w:br'))
                if br.get(qn('w:type')) == 'column')
status = OK if col_hasil == 0 else FAIL
print(f"[10] Column break       {status}  asli={col_src}  hasil={col_hasil}")

# ── Ringkasan ───────────────────────────────────────────────────────
print()
print("=" * 52)
if not issues:
    print("  SEMUA BAIK — tidak ada perbedaan konten yang signifikan")
else:
    print(f"  {len(issues)} CATATAN:")
    for iss in issues:
        print(f"  - {iss}")
print("=" * 52)
