import sys, os
sys.stdout.reconfigure(encoding='utf-8')
sys.path.insert(0, os.path.normpath(os.path.join(os.path.dirname(__file__), '..')))
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

SRC = r'C:\Users\farizal\Downloads\SKRIPSI GEMA RAMADHAN fiks.docx'
doc = Document(SRC)
all_paras = list(doc.paragraphs)

print('=== Paragraphs near "lampiran" (para.text) ===')
for i, para in enumerate(all_paras):
    txt = para.text.strip()
    if 'lampiran' in txt.lower() or 'kuesioner' in txt.lower():
        pPr = para._p.find(f'{{{W}}}pPr')
        sp  = pPr.find(f'{{{W}}}sectPr') if pPr is not None else None
        pgbr = any(br.get(f'{{{W}}}type') == 'page' for br in para._p.iter(f'{{{W}}}br'))
        print(f'  [{i}] style={para.style.name!r} pgBr={pgbr} sect={sp is not None}')
        print(f'       text={repr(txt[:80])}')

print('\n=== Text boxes / drawings (w:txbx) containing "lampiran" ===')
body = list(doc.element.body)
for i, el in enumerate(body):
    # Scan for text boxes in any element
    for txbx in el.iter(f'{{{W}}}txbx'):
        # Collect all text in text box
        txt = ''.join(t.text or '' for t in txbx.iter(f'{{{W}}}t')).strip()
        if 'lampiran' in txt.lower():
            tag = el.tag.split('}')[-1]
            print(f'  body[{i}] <{tag}> txbx text: {repr(txt[:100])}')

print('\n=== Runs with "lampiran" in w:t (including in drawings) ===')
for i, el in enumerate(body):
    for t_el in el.iter(f'{{{W}}}t'):
        txt = (t_el.text or '').strip()
        if 'lampiran' in txt.lower():
            # Walk up to find nearest p
            parent_tags = []
            p = t_el.getparent()
            while p is not None:
                parent_tags.append(p.tag.split('}')[-1])
                p = p.getparent()
            tag = el.tag.split('}')[-1]
            print(f'  body[{i}] <{tag}> text={repr(txt[:60])} parents={parent_tags[:6]}')

print('\n=== Sections in document ===')
for i, sec in enumerate(doc.sections):
    pn  = sec._sectPr.find(qn('w:pgNumType'))
    fmt = pn.get(qn('w:fmt')) if pn is not None else None
    st  = pn.get(qn('w:start')) if pn is not None else None
    print(f'  sec[{i}] fmt={fmt} start={st}')

print('\n=== scan_zones result ===')
from utils import DocProcessor
d2   = Document(SRC)
proc = DocProcessor(d2, 'Times New Roman', 12)
proc.purge_all_headers_footers()
roman_start_p, bab_p_list = proc.scan_zones()
body2 = list(d2.element.body)
print(f'  roman_start_p index: {body2.index(roman_start_p) if roman_start_p is not None else None}')
print(f'  bab_p_list count: {len(bab_p_list)}')
for j, p in enumerate(bab_p_list):
    txt = ''.join(t.text or '' for t in p.iter(f'{{{W}}}t')).strip()
    idx = body2.index(p) if p in body2 else -1
    print(f'  bab[{j}] body[{idx}] text={repr(txt[:60])}')
