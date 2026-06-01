import sys, os
sys.stdout.reconfigure(encoding='utf-8')
sys.path.insert(0, os.path.normpath(os.path.join(os.path.dirname(__file__), '..')))
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
from utils import DocProcessor
import paket4

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

SRC  = r'C:\Users\farizal\Downloads\SKRIPSI GEMA RAMADHAN fiks.docx'
OUT  = r'C:\Users\farizal\Downloads\GEMA_c4_semb_lamprn.docx'

# ── Full paket4 run (semb_lamprn=Ya) ──────────────────────────────────────
doc  = Document(SRC)
proc = DocProcessor(doc, 'Times New Roman', 12)
proc.purge_all_headers_footers()
roman_start_p, bab_p_list = proc.scan_zones()
from utils import DocProcessor as DP
new_rsp, use_exact = DP.advance_roman_start(doc, roman_start_p, 1)
roman_start_p = proc.insert_breaks(roman_start_p, bab_p_list, exact_roman_start=use_exact)
roman_sec, bab_sec_list, n_sections = proc.build_section_map(roman_start_p, bab_p_list)
print(f'roman_sec={roman_sec}  n_sections={n_sections}')
print(f'bab_sec_list: {bab_sec_list}')

# Apply paket4 with semb_lamprn=Ya
paket4.apply(
    proc, roman_sec, bab_sec_list, n_sections, 'Ya',
    pos_romawi='Tengah Bawah', pos_bab='Tengah Bawah', pos_isi_bab='Kanan Atas',
    dimulai_dari='i', semb_dafus='Tidak', semb_lamprn='Ya',
    bab_p_list=bab_p_list
)
doc.save(OUT)
print(f'\nSaved to {OUT}')

# ── Inspect output sections ────────────────────────────────────────────────
doc2 = Document(OUT)
print('\n=== Output sections ===')
for i, sec in enumerate(doc2.sections):
    pn   = sec._sectPr.find(qn('w:pgNumType'))
    fmt  = pn.get(qn('w:fmt'))   if pn is not None else None
    st   = pn.get(qn('w:start')) if pn is not None else None
    # footer text
    try:
        ftxt = sec.footer.paragraphs[0].text.strip()[:30]
        flink = sec.footer.is_linked_to_previous
    except:
        ftxt = '?'; flink = '?'
    try:
        htxt = sec.header.paragraphs[0].text.strip()[:30]
        hlink = sec.header.is_linked_to_previous
    except:
        htxt = '?'; hlink = '?'
    print(f'  sec[{i}] fmt={fmt} start={st} '
          f'hdr={repr(htxt)} hlink={hlink} '
          f'ftr={repr(ftxt)} flink={flink}')

# ── Body structure near lampiran ──────────────────────────────────────────
body = list(doc2.element.body)
print('\n=== Body [1040:1075] of OUTPUT ===')
for i in range(1040, min(1075, len(body))):
    el  = body[i]
    tag = el.tag.split('}')[-1]
    if tag != 'p':
        print(f'  [{i}] <{tag}>')
        continue
    txt  = ''.join(t.text or '' for t in el.iter(f'{{{W}}}t')).strip()[:50]
    pPr  = el.find(f'{{{W}}}pPr')
    sp   = pPr.find(f'{{{W}}}sectPr') if pPr is not None else None
    pgbr = any(br.get(f'{{{W}}}type') == 'page' for br in el.iter(f'{{{W}}}br'))
    flags = []
    if sp is not None:
        t2 = sp.find(f'{{{W}}}type')
        flags.append(f'SECT({t2.get(qn("w:val")) if t2 is not None else "nextPage"})')
    if pgbr:
        flags.append('pgBr')
    print(f'  [{i}] {repr(txt) if txt else "(empty)"} {flags}')
