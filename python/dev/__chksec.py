"""Check section structure of output files"""
import os, sys
sys.path.insert(0, os.path.normpath(os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '..')))
from docx import Document
from docx.oxml.ns import qn

ROOT  = os.path.normpath(os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '..'))
HASIL = os.path.join(ROOT, 'test_files', 'paket3', 'cover 2 dimulai dari iii', 'hasil')
W     = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

TARGETS = ['Docx 6_v2.docx', 'Docx 7_v2.docx']

for fname in TARGETS:
    path = os.path.join(HASIL, fname)
    if not os.path.exists(path):
        print(f"{fname}: NOT FOUND"); continue
    doc = Document(path)
    body_els = list(doc.element.body)

    def _has_break(el):
        if any(br.get("{%s}type" % W) == 'page' for br in el.iter("{%s}br" % W)):
            return True
        pPr = el.find("{%s}pPr" % W)
        return pPr is not None and pPr.find("{%s}sectPr" % W) is not None

    def txt(el):
        return ''.join(t.text or '' for t in el.iter('{%s}t' % W)).strip()[:50]

    def get_pn(sec):
        pn = sec._sectPr.find(qn('w:pgNumType'))
        if pn is None: return 'decimal:None'
        fmt   = pn.get(qn('w:fmt'), 'decimal')
        start = pn.get(qn('w:start'), None)
        return f"{fmt}:{start}"

    def get_footer_info(sec):
        info = []
        if sec.different_first_page_header_footer:
            info.append('diff_first')
        try:
            fpf_text = ''.join(p.text for p in sec.first_page_footer.paragraphs).strip()
            info.append(f"fpf={fpf_text!r}")
        except: pass
        try:
            ft_text = ''.join(p.text for p in sec.footer.paragraphs).strip()
            info.append(f"ft={ft_text!r}")
        except: pass
        return ', '.join(info)

    print(f"\n{'='*60}")
    print(f"FILE: {fname}")
    print(f"Sections: {len(doc.sections)}")
    for i, sec in enumerate(doc.sections):
        pn_str = get_pn(sec)
        ft_str = get_footer_info(sec)
        print(f"  sec[{i}] {pn_str}  | {ft_str}")

    # Show body structure around section breaks
    print()
    print("Body structure (breaks and nearby content):")
    for i, el in enumerate(body_els):
        tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
        pPr = el.find("{%s}pPr" % W)
        has_spr  = (pPr is not None and pPr.find("{%s}sectPr" % W) is not None)
        has_pgbr = any(br.get("{%s}type" % W) == 'page' for br in el.iter("{%s}br" % W))
        if has_spr or has_pgbr or (i < 5) or any(abs(i-j) < 2 for j,e2 in enumerate(body_els) if e2.find("{%s}pPr" % W) is not None and e2.find("{%s}pPr" % W).find("{%s}sectPr" % W) is not None):
            mark = '[sectPr]' if has_spr else ('[pgBr]' if has_pgbr else '       ')
            print(f"  [{i:3d}] {tag:<5}  {mark}  {txt(el)!r}")
