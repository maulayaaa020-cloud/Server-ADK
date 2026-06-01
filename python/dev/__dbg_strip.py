import sys
sys.path.insert(0, r'd:\Freelaces\Server\htdocs\adk\python')
from docx import Document
from docx.oxml.ns import qn
from utils import DocProcessor

path = r'd:\Freelaces\Server\htdocs\adk\test_files\paket3\cover 1 dimulai dari ii\hasil\Docx 19.docx'
doc  = Document(path)
proc = DocProcessor(doc, 'Times New Roman', 12)

body     = doc.element.body
children = list(body)

print(f"Total body children: {len(children)}")
print(f"Total doc.paragraphs: {len(doc.paragraphs)}")
print()

# Cek apakah ada non-paragraph elements di body
non_para = [(i, c.tag.split('}')[-1]) for i, c in enumerate(children) if not c.tag.endswith('}p')]
print(f"Non-paragraph elements di body: {len(non_para)}")
for i, tag in non_para:
    print(f"  children[{i}] = <{tag}>")
print()

# Fokus area para 181-190 di doc.paragraphs
print("=== doc.paragraphs[181-190] ===")
paras = list(doc.paragraphs)
for i in range(181, min(191, len(paras))):
    p = paras[i]
    # Cari index di children
    try:
        ci = children.index(p._p)
        in_children = True
    except ValueError:
        ci = -1
        in_children = False
    has_sect = DocProcessor._has_sectPr(p._p)
    has_content = bool(p.text.strip())
    print(f"  doc.para[{i}] -> children[{ci}]  content={has_content}  sect={has_sect}  "
          f"text={repr(p.text.strip()[:40])}")
print()

# Simulasi fungsi strip untuk para[187]
print("=== SIMULASI SCAN MUNDUR DARI para[187] ===")
p187 = paras[187]
try:
    i187 = children.index(p187._p)
    print(f"para[187] ada di children[{i187}]")
except ValueError:
    print("para[187] TIDAK ADA di children!")
    sys.exit()

to_remove = set()
j = i187 - 1
steps = 0
while j >= 0 and j not in to_remove and steps < 20:
    steps += 1
    prev = children[j]
    tag = prev.tag.split('}')[-1]
    is_para = prev.tag.endswith('}p')
    if is_para:
        has_c = proc._p_has_content(prev)
        has_s = DocProcessor._has_sectPr(prev)
        txt   = ''.join(t.text or '' for t in prev.iter(qn('w:t'))).strip()[:30]
        print(f"  j={j} <{tag}> content={has_c} sect={has_s} text={repr(txt)}")
        if has_c:
            print(f"    -> BREAK (content)")
            break
        if has_s:
            print(f"    -> BREAK (sect)")
            break
        to_remove.add(j)
        print(f"    -> TAMBAH ke to_remove")
    else:
        print(f"  j={j} <{tag}> (non-para) -> SKIP")
    j -= 1

print(f"\nto_remove: {sorted(to_remove)}")
print()

# Coba hapus
print("=== MENCOBA HAPUS ===")
for idx in sorted(to_remove, reverse=True):
    elem = children[idx]
    parent = elem.getparent()
    print(f"  children[{idx}] parent={parent.tag.split('}')[-1] if parent is not None else 'None'}")
    if parent is not None:
        parent.remove(elem)
        print(f"    -> BERHASIL dihapus")
    else:
        print(f"    -> GAGAL (parent None)")

print(f"\nTotal body children setelah hapus: {len(list(body))}")
