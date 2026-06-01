import sys, os
sys.path.insert(0, os.path.normpath(os.path.join(os.path.dirname(__file__), '..')))
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

W  = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
RI = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id'

OUT = r'D:\Freelaces\Server\htdocs\adk\test_files\paket3\cover 2 dimulai dari iii\hasil\Docx 14_c2.docx'

doc  = Document(OUT)
body = list(doc.element.body)

print('=== Sections ===')
for i, sec in enumerate(doc.sections):
    pn    = sec._sectPr.find(qn('w:pgNumType'))
    fmt   = pn.get(qn('w:fmt'))   if pn is not None else None
    start = pn.get(qn('w:start')) if pn is not None else None
    ftr   = len(sec._sectPr.findall(qn('w:footerReference')))
    hdr   = len(sec._sectPr.findall(qn('w:headerReference')))
    tpg   = sec._sectPr.find(qn('w:titlePg')) is not None
    typ   = sec._sectPr.find(qn('w:type'))
    typ_v = typ.get(qn('w:val')) if typ is not None else 'nextPage'
    print(f'  sec[{i}] fmt={fmt} start={start} type={typ_v} titlePg={tpg} hdr={hdr} ftr={ftr}')
    # footer content
    for ft in ['default', 'first', 'even']:
        try:
            footer = sec.footer if ft == 'default' else (sec.first_page_footer if ft == 'first' else sec.even_page_footer)
            txt = footer.paragraphs[0].text.strip() if footer.paragraphs else ''
            linked = footer.is_linked_to_previous
            print(f'    footer[{ft}]: linked={linked} text={repr(txt[:40])}')
        except Exception as e:
            print(f'    footer[{ft}]: ERROR {e}')

print('\n=== Body [0:50] ===')
for i in range(min(50, len(body))):
    el  = body[i]
    tag = el.tag.split('}')[-1]
    if tag != 'p':
        print(f'  [{i}] <{tag}>')
        continue
    txt   = ''.join(t.text or '' for t in el.iter(f'{{{W}}}t')).strip()[:45]
    ppPr  = el.find(f'{{{W}}}pPr')
    sp    = ppPr.find(f'{{{W}}}sectPr') if ppPr is not None else None
    flags = []
    if sp is not None:
        t2 = sp.find(f'{{{W}}}type')
        flags.append(f'SECT({t2.get(qn("w:val")) if t2 is not None else "nextPage"})')
        pn2 = sp.find(qn('w:pgNumType'))
        if pn2 is not None:
            flags.append(f'pgNum(fmt={pn2.get(qn("w:fmt"))} start={pn2.get(qn("w:start"))})')
    if any(br.get(f'{{{W}}}type') == 'page' for br in el.iter(f'{{{W}}}br')):
        flags.append('pgBr')
    print(f'  [{i}] {repr(txt) if txt else "(empty)"} {flags}')
