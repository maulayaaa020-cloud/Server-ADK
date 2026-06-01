import sys
sys.path.insert(0, r'd:\Freelaces\Server\htdocs\adk\python')
from docx import Document
from docx.oxml.ns import qn

SRC   = r'd:\Freelaces\Server\htdocs\adk\test_files\paket3\cover 1 dimulai dari ii\Docx 19.docx'
HASIL = r'd:\Freelaces\Server\htdocs\adk\test_files\paket3\cover 1 dimulai dari ii\hasil\Docx 19.docx'

src   = Document(SRC)
hasil = Document(HASIL)

def all_text_lines(doc):
    """Kumpulkan semua teks non-kosong dari paragraf + sel tabel."""
    lines = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    lines.append(t)
    return lines

def para_kosong(doc):
    return sum(1 for p in doc.paragraphs if not p.text.strip())

# ── Hitung paragraf kosong vs isi ─────────────────────────────────
print("=== PARAGRAF KOSONG VS ISI ===")
s_kosong = para_kosong(src)
h_kosong = para_kosong(hasil)
s_isi = len(src.paragraphs) - s_kosong
h_isi = len(hasil.paragraphs) - h_kosong
print(f"Asli : total={len(src.paragraphs)}  kosong={s_kosong}  isi={s_isi}")
print(f"Hasil: total={len(hasil.paragraphs)}  kosong={h_kosong}  isi={h_isi}")
print()

# ── Bandingkan teks actual (non-kosong) ───────────────────────────
src_lines   = [p.text.strip() for p in src.paragraphs   if p.text.strip()]
hasil_lines = [p.text.strip() for p in hasil.paragraphs if p.text.strip()]

print("=== PERBANDINGAN TEKS NON-KOSONG ===")
print(f"Asli : {len(src_lines)} baris teks")
print(f"Hasil: {len(hasil_lines)} baris teks")
print()

# Cek apakah semua teks asli ada di hasil
src_set   = set(src_lines)
hasil_set = set(hasil_lines)

hilang = src_set - hasil_set
tambah = hasil_set - src_set

print(f"Teks yang ada di asli tapi HILANG di hasil : {len(hilang)}")
if hilang:
    for t in sorted(hilang)[:10]:
        print(f"  - {repr(t[:70])}")
    if len(hilang) > 10:
        print(f"  ... dan {len(hilang)-10} lagi")

print()
print(f"Teks yang ada di hasil tapi TIDAK di asli  : {len(tambah)}")
if tambah:
    for t in sorted(tambah)[:10]:
        print(f"  + {repr(t[:70])}")
    if len(tambah) > 10:
        print(f"  ... dan {len(tambah)-10} lagi")

print()

# ── Cek urutan teks (apakah urutan berubah) ───────────────────────
print("=== CEK URUTAN TEKS (50 baris pertama) ===")
order_ok = True
j = 0
for i, line in enumerate(src_lines[:50]):
    # cari di hasil mulai posisi j
    found = False
    for k in range(j, min(j + 5, len(hasil_lines))):
        if hasil_lines[k] == line:
            j = k + 1
            found = True
            break
    if not found:
        order_ok = False
        print(f"  Urutan berubah di: {repr(line[:60])}")
        break

if order_ok:
    print(f"  OK  Urutan 50 baris pertama terjaga")

print()

# ── Khusus: apakah teks section break paragraf ikut hilang ────────
print("=== PARAGRAF YANG HILANG (sampel) ===")
# Paragraf kosong di asli yang punya sectPr = section break paragraf
from utils import DocProcessor
sect_break_paras = [p for p in src.paragraphs if DocProcessor._has_sectPr(p._p)]
print(f"Paragraf section break di asli: {len(sect_break_paras)}")
print(f"Paragraf hilang total         : {len(src.paragraphs) - len(hasil.paragraphs)}")
print(f"=> Selisih hampir sama dengan jumlah section break yang dikonsolidasi script")
