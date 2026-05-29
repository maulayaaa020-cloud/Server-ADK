"""Debug v3: investigasi 8 file bermasalah setelah fix batch pertama"""
import os, sys
sys.path.insert(0, os.path.normpath(os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '..')))
from docx import Document
from docx.oxml.ns import qn
from utils import DocProcessor

W    = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
ROOT = os.path.normpath(os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '..'))
SRC  = os.path.join(ROOT, 'test_files', 'paket3', 'cover 2 dimulai dari iii')
HASIL = os.path.join(SRC, 'hasil')

def _has_sectPr(el):
    pPr = el.find("{%s}pPr" % W)
    return pPr is not None and pPr.find("{%s}sectPr" % W) is not None

def _has_pgbr(el):
    return any(br.get("{%s}type" % W) == 'page' for br in el.iter("{%s}br" % W))

def _has_break(el):
    return _has_pgbr(el) or _has_sectPr(el)

def txt(el, n=45):
    return ''.join(t.text or '' for t in el.iter('{%s}t' % W)).strip()[:n]

def mark(el):
    if _has_sectPr(el) and _has_pgbr(el): return '[spr+pgBr]'
    if _has_sectPr(el): return '[sectPr]  '
    if _has_pgbr(el):   return '[pgBr]    '
    return '          '

def show_range(body, lo, hi, rsp_idx=None, new_idx=None, bab_idx=None, extra=None):
    bab_idx  = bab_idx or []
    extra    = extra or {}
    for i in range(max(0, lo), min(len(body), hi)):
        el  = body[i]
        tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
        mk  = ''
        if i == rsp_idx:  mk += ' <--RSP'
        if i == new_idx:  mk += ' <--NEW'
        if i in bab_idx:  mk += ' <--BAB'
        if i in extra:    mk += f' <--{extra[i]}'
        print(f"  [{i:3d}] {tag:<4} {mark(el)} {txt(el)!r}{mk}")

def show_output_sections(fname_base):
    p = os.path.join(HASIL, fname_base + '_v2.docx')
    if not os.path.exists(p):
        print("  (output not found)")
        return
    d = Document(p)
    for i, sec in enumerate(d.sections):
        pn = sec._sectPr.find(qn('w:pgNumType'))
        fmt   = pn.get(qn('w:fmt'), 'decimal') if pn is not None else 'decimal'
        start = pn.get(qn('w:start'), None)   if pn is not None else None
        print(f"  OUTPUT sec[{i}]: {fmt}:{start}")

TARGETS = [6, 9, 14, 15, 16, 17, 18]

for n in TARGETS:
    fname = f"Docx {n}.docx"
    path  = os.path.join(SRC, fname)
    if not os.path.exists(path): continue

    doc      = Document(path)
    proc     = DocProcessor(doc, 'Times New Roman', 12)
    rsp, bab = proc.scan_zones()
    body     = list(doc.element.body)
    rsp_idx  = body.index(rsp)
    bab_idx  = [body.index(p) for p in bab]
    bb       = sum(1 for el in body[:rsp_idx] if _has_break(el))

    # Cek tipe break TERAKHIR sebelum rsp
    last_break_idx  = -1
    last_break_type = None
    for j in range(rsp_idx - 1, -1, -1):
        if _has_break(body[j]):
            last_break_idx = j
            if _has_sectPr(body[j]): last_break_type = 'sectPr'
            elif _has_pgbr(body[j]): last_break_type = 'pgBr'
            break

    new_rsp, use_exact = DocProcessor.advance_roman_start(doc, rsp, 2)
    new_idx = body.index(new_rsp)

    print(f"\n{'='*65}")
    print(f"Docx {n}: bb={bb}, last_break=[{last_break_idx}]={last_break_type}")
    print(f"  rsp=[{rsp_idx}] {txt(rsp)!r}")
    print(f"  new=[{new_idx}] {txt(new_rsp)!r}  exact={use_exact}")
    print(f"  bab: {bab_idx[:3]}")

    # Tampilkan konteks sekitar last_break, rsp, dan new_rsp
    lo = max(0, min(last_break_idx - 1, rsp_idx - 3, new_idx - 3))
    hi = max(rsp_idx, new_idx) + 6
    extra = {}
    if last_break_idx >= 0: extra[last_break_idx] = 'LAST_BREAK'
    show_range(body, lo, hi, rsp_idx, new_idx, bab_idx, extra)
    show_output_sections(f"Docx {n}")

# Docx 10: khusus tampilkan struktur lengkap awal
print(f"\n{'='*65}")
print("Docx 10: structure [0..30]")
doc10 = Document(os.path.join(SRC, 'Docx 10.docx'))
proc10 = DocProcessor(doc10, 'Times New Roman', 12)
rsp10, bab10 = proc10.scan_zones()
body10 = list(doc10.element.body)
rsp10_idx = body10.index(rsp10)
bb10 = sum(1 for el in body10[:rsp10_idx] if _has_break(el))
print(f"  bb={bb10}, rsp=[{rsp10_idx}]={txt(rsp10)!r}")
show_range(body10, 0, 32, rsp10_idx)
show_output_sections("Docx 10")

# Docx 15: tampilkan struktur lebih luas
print(f"\n{'='*65}")
print("Docx 15: structure [0..35]")
doc15 = Document(os.path.join(SRC, 'Docx 15.docx'))
proc15 = DocProcessor(doc15, 'Times New Roman', 12)
rsp15, bab15 = proc15.scan_zones()
body15 = list(doc15.element.body)
rsp15_idx = body15.index(rsp15)
bb15 = sum(1 for el in body15[:rsp15_idx] if _has_break(el))
new15, exact15 = DocProcessor.advance_roman_start(doc15, rsp15, 2)
new15_idx = body15.index(new15)
print(f"  bb={bb15}, rsp=[{rsp15_idx}]={txt(rsp15)!r}  new=[{new15_idx}]={txt(new15)!r}  exact={exact15}")
show_range(body15, 0, 35, rsp15_idx, new15_idx)
show_output_sections("Docx 15")

# Cek isi section output untuk Docx 9 - mana paragraph tepat sebelum dan sesudah new_rsp
print(f"\n{'='*65}")
print("Docx 9 OUTPUT sections detail:")
p9 = os.path.join(HASIL, 'Docx 9_v2.docx')
if os.path.exists(p9):
    d9 = Document(p9)
    b9 = list(d9.element.body)
    for i, sec in enumerate(d9.sections):
        pn = sec._sectPr.find(qn('w:pgNumType'))
        fmt   = pn.get(qn('w:fmt'), 'decimal') if pn is not None else 'decimal'
        start = pn.get(qn('w:start'), None)   if pn is not None else None
        print(f"  sec[{i}]: {fmt}:{start}")
    # Cari paragraph dengan sectPr
    print("  Paragraphs with sectPr:")
    for i, el in enumerate(b9):
        if _has_sectPr(el):
            print(f"    [{i:3d}] {mark(el)} {txt(el)!r}")
        if _has_pgbr(el):
            tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
            if tag == 'p':
                print(f"    [{i:3d}] [pgBr]     {txt(el)!r}")
