import sys
sys.path.insert(0, r'd:\Freelaces\Server\htdocs\adk\python')
from docx import Document
from docx.oxml.ns import qn

path = r'd:\Freelaces\Server\htdocs\adk\test_files\paket3\cover 1 dimulai dari ii\Docx 19.docx'
doc = Document(path)

# Cek 1: column break di doc.paragraphs (body langsung)
body_col = 0
for p in doc.paragraphs:
    for br in p._p.findall('.//' + qn('w:br')):
        if br.get(qn('w:type')) == 'column':
            body_col += 1

# Cek 2: column break di table cells
table_col = 0
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for br in para._p.findall('.//' + qn('w:br')):
                    if br.get(qn('w:type')) == 'column':
                        table_col += 1

# Cek 3: total di seluruh XML body
import zipfile
from lxml import etree
with zipfile.ZipFile(path) as z:
    xml = z.read('word/document.xml')
root = etree.fromstring(xml)
ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
all_col = root.findall('.//w:br[@w:type="column"]', ns)

print(f"Column break di doc.paragraphs (body): {body_col}")
print(f"Column break di table cells          : {table_col}")
print(f"Column break total di XML            : {len(all_col)}")
print()

# Cek apakah mereka di dalam w:tbl atau di luar
in_tbl  = 0
out_tbl = 0
for br in all_col:
    # Naik ke atas sampai ketemu w:tbl atau w:body
    p = br.getparent()
    found_tbl = False
    while p is not None:
        if p.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl':
            found_tbl = True
            break
        if p.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}body':
            break
        p = p.getparent()
    if found_tbl:
        in_tbl += 1
    else:
        out_tbl += 1

print(f"  - Di dalam w:tbl  : {in_tbl}")
print(f"  - Di luar w:tbl   : {out_tbl}")
