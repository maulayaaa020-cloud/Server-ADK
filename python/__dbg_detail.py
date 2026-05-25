"""Detailed debug for problematic files - break counting and section structure"""
import os, sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx import Document
from docx.oxml.ns import qn
from utils import DocProcessor

W    = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
ROOT = os.path.normpath(os.path.join(os.path.dirname(os.path.abspath(__file__)), '..'))
SRC  = os.path.join(ROOT, 'test_files', 'paket3', 'cover 2 dimulai dari iii')

def _has_break(el):
    if any(br.get("{%s}type" % W) == 'page' for br in el.iter("{%s}br" % W)):
        return True
    pPr = el.find("{%s}pPr" % W)
    return pPr is not None and pPr.find("{%s}sectPr" % W) is not None

def txt(el, n=35):
    return ''.join(t.text or '' for t in el.iter('{%s}t' % W)).strip()[:n]

# === Docx 12: break counting detail ===
print("=== DOCX 12: First 4 breaks in document ===")
doc = Document(os.path.join(SRC, 'Docx 12.docx'))
body_els = list(doc.element.body)
count = 0
for i, el in enumerate(body_els):
    if _has_break(el):
        count += 1
        tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
        pPr = el.find("{%s}pPr" % W)
        has_spr  = (pPr is not None and pPr.find("{%s}sectPr" % W) is not None)
        has_pgbr = any(br.get("{%s}type" % W) == 'page' for br in el.iter("{%s}br" % W))
        mark = '[sectPr]' if has_spr else '[pgBr]'
        print(f"  Break {count}: [{i:3d}] {mark} {txt(el)!r}")
        if count == 5:
            print("  (stopping at 5)")
            break

print()
# Trace advance_roman_start manually
proc = DocProcessor(doc, 'Times New Roman', 12)
rsp, bab = proc.scan_zones()
rsp_idx = body_els.index(rsp)
print(f"  rsp=[{rsp_idx}] {txt(rsp)!r}")
target_pg = 3
current_pg = 1
for i, el in enumerate(body_els):
    if current_pg >= target_pg:
        print(f"  candidate=[{i}] {txt(el)!r} (current_pg={current_pg})")
        break
    if _has_break(el):
        current_pg += 1
        tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
        pPr = el.find("{%s}pPr" % W)
        has_spr = (pPr is not None and pPr.find("{%s}sectPr" % W) is not None)
        has_pgbr = any(br.get("{%s}type" % W) == 'page' for br in el.iter("{%s}br" % W))
        mark = '[sectPr]' if has_spr else '[pgBr]'
        print(f"  Break at [{i}] {mark} {txt(el)!r} -> pg={current_pg}")

print()
# === Docx 7: pgBr BEFORE rsp context ===
print("=== DOCX 7: context [22..28] ===")
doc7 = Document(os.path.join(SRC, 'Docx 7.docx'))
body7 = list(doc7.element.body)
proc7 = DocProcessor(doc7, 'Times New Roman', 12)
rsp7, bab7 = proc7.scan_zones()
rsp7_idx = body7.index(rsp7)
for i in range(22, min(29, len(body7))):
    el = body7[i]
    tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
    pPr = el.find("{%s}pPr" % W)
    has_spr  = (pPr is not None and pPr.find("{%s}sectPr" % W) is not None)
    has_pgbr = any(br.get("{%s}type" % W) == 'page' for br in el.iter("{%s}br" % W))
    mark = '[sectPr]' if has_spr else ('[pgBr]' if has_pgbr else '       ')
    mk = ' <--RSP' if i == rsp7_idx else ''
    print(f"  [{i:3d}] {tag:<5} {mark} {txt(el)!r}{mk}")

print()
print("=== DOCX 9: context around rsp ===")
doc9 = Document(os.path.join(SRC, 'Docx 9.docx'))
body9 = list(doc9.element.body)
proc9 = DocProcessor(doc9, 'Times New Roman', 12)
rsp9, bab9 = proc9.scan_zones()
rsp9_idx = body9.index(rsp9)
bb9 = sum(1 for el in body9[:rsp9_idx] if _has_break(el))
print(f"  rsp=[{rsp9_idx}] bb={bb9}")
for i in range(max(0, rsp9_idx-3), min(len(body9), rsp9_idx+5)):
    el = body9[i]
    tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
    pPr = el.find("{%s}pPr" % W)
    has_spr  = (pPr is not None and pPr.find("{%s}sectPr" % W) is not None)
    has_pgbr = any(br.get("{%s}type" % W) == 'page' for br in el.iter("{%s}br" % W))
    mark = '[sectPr]' if has_spr else ('[pgBr]' if has_pgbr else '       ')
    mk = ' <--RSP' if i == rsp9_idx else ''
    print(f"  [{i:3d}] {tag:<5} {mark} {txt(el)!r}{mk}")

print()
print("=== DOCX 17: context [22..30] ===")
doc17 = Document(os.path.join(SRC, 'Docx 17.docx'))
body17 = list(doc17.element.body)
proc17 = DocProcessor(doc17, 'Times New Roman', 12)
rsp17, bab17 = proc17.scan_zones()
rsp17_idx = body17.index(rsp17)
for i in range(22, min(32, len(body17))):
    el = body17[i]
    tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
    pPr = el.find("{%s}pPr" % W)
    has_spr  = (pPr is not None and pPr.find("{%s}sectPr" % W) is not None)
    has_pgbr = any(br.get("{%s}type" % W) == 'page' for br in el.iter("{%s}br" % W))
    mark = '[sectPr]' if has_spr else ('[pgBr]' if has_pgbr else '       ')
    mk = ' <--RSP' if i == rsp17_idx else ''
    print(f"  [{i:3d}] {tag:<5} {mark} {txt(el)!r}{mk}")

print()
print("=== DOCX 14: how _find_roman_page_start works ===")
doc14 = Document(os.path.join(SRC, 'Docx 14.docx'))
body14 = list(doc14.element.body)
proc14 = DocProcessor(doc14, 'Times New Roman', 12)
rsp14, bab14 = proc14.scan_zones()
rsp14_idx = body14.index(rsp14)
bb14 = sum(1 for el in body14[:rsp14_idx] if _has_break(el))
print(f"  rsp=[{rsp14_idx}] {txt(rsp14)!r} bb={bb14}")
# Show the break before rsp (boundary_idx)
boundary_idx = -1
for j in range(rsp14_idx - 1, -1, -1):
    el = body14[j]
    if (el.tag.split('}')[-1] if '}' in el.tag else el.tag) != 'p':
        continue
    if _has_break(el):
        boundary_idx = j
        break
print(f"  boundary_idx={boundary_idx}")
for i in range(max(0, boundary_idx-1), min(len(body14), rsp14_idx+3)):
    el = body14[i]
    tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
    pPr = el.find("{%s}pPr" % W)
    has_spr  = (pPr is not None and pPr.find("{%s}sectPr" % W) is not None)
    has_pgbr = any(br.get("{%s}type" % W) == 'page' for br in el.iter("{%s}br" % W))
    mark = '[sectPr]' if has_spr else ('[pgBr]' if has_pgbr else '       ')
    mk = ' <--RSP' if i == rsp14_idx else ''
    mk += ' <--BOUNDARY' if i == boundary_idx else ''
    print(f"  [{i:3d}] {tag:<5} {mark} {txt(el)!r}{mk}")

print()
print("=== DOCX 10: full structure (first 25 elements) ===")
doc10 = Document(os.path.join(SRC, 'Docx 10.docx'))
body10 = list(doc10.element.body)
proc10 = DocProcessor(doc10, 'Times New Roman', 12)
rsp10, bab10 = proc10.scan_zones()
rsp10_idx = body10.index(rsp10)
for i in range(min(25, len(body10))):
    el = body10[i]
    tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
    pPr = el.find("{%s}pPr" % W)
    has_spr  = (pPr is not None and pPr.find("{%s}sectPr" % W) is not None)
    has_pgbr = any(br.get("{%s}type" % W) == 'page' for br in el.iter("{%s}br" % W))
    mark = '[sectPr]' if has_spr else ('[pgBr]' if has_pgbr else '       ')
    has_content = proc10._p_has_content(el)
    cts = 'C' if has_content else ' '
    mk = ' <--RSP' if i == rsp10_idx else ''
    print(f"  [{i:2d}] {cts}{tag:<5} {mark} {txt(el)!r}{mk}")
