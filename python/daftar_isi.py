


"""
daftar_isi.py — Buat daftar isi otomatis dari dokumen Word.
Usage: python daftar_isi.py <input.docx> <output.docx> <kedalaman> <format_titik>

kedalaman:    H1 | H1+H2 | H1+H2+H3
format_titik: titik | tab
"""
import sys
import os
import json
import re
import zipfile

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from utils import is_toc_entry, _has_toc_field, _fuzzy_match as _fuzzy_str


# ── Konstanta ─────────────────────────────────────────────────────────────────

MAX_LEVEL_MAP = {'H1': 1, 'H1+H2': 2, 'H1+H2+H3': 3}

# Tab stop kanan untuk nomor halaman (twips: 1 cm ≈ 567 twips, ~15 cm = 8505)
TAB_POS_TWIPS = '8505'

_XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'


# ── Validasi file ─────────────────────────────────────────────────────────────

def _fail(code, msg):
    print(json.dumps({"status": "error", "code": code, "message": msg}))
    sys.exit(1)


def validate_file(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == '.doc':
        _fail("FORMAT_NOT_SUPPORTED",
              "Format .doc tidak didukung. Buka di Microsoft Word lalu simpan ulang sebagai .docx.")
    try:
        size = os.path.getsize(path)
    except OSError:
        _fail("FILE_READ_ERROR", "Tidak dapat membaca ukuran file.")
    if size > 30 * 1024 * 1024:
        _fail("FILE_TOO_LARGE", f"Ukuran file ({size // (1024*1024)} MB) melebihi batas 30 MB.")
    try:
        with zipfile.ZipFile(path, 'r') as z:
            names = z.namelist()
            if 'word/document.xml' not in names:
                _fail("INVALID_DOCX", "File bukan dokumen Word (.docx) yang valid.")
            if 'word/vbaProject.bin' in names:
                _fail("MACRO_DETECTED",
                      "File mengandung macro VBA. Simpan ulang sebagai .docx biasa.")
            total = sum(i.file_size for i in z.infolist())
            if total > 200 * 1024 * 1024:
                _fail("FILE_TOO_LARGE", "Konten file melebihi 200 MB.")
    except zipfile.BadZipFile:
        _fail("INVALID_DOCX", "File rusak atau bukan format .docx yang valid.")


# ── Deteksi heading ───────────────────────────────────────────────────────────

def _get_outline_level_from_xml(para):
    """Baca w:outlineLvl dari XML paragraf (0-based), kembalikan 1-based."""
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        return None
    el = pPr.find(qn('w:outlineLvl'))
    if el is None:
        return None
    val = el.get(qn('w:val'))
    if val is None:
        return None
    try:
        lvl = int(val)
        return lvl + 1 if lvl < 9 else None
    except ValueError:
        return None


def _get_outline_level_from_style(para):
    """
    Naik ke rantai parent style dan cek w:outlineLvl di definisi style.
    Berguna untuk style kustom yang menyimpan outlineLvl di sana.
    """
    style = para.style
    while style:
        pPr = style.element.find(qn('w:pPr'))
        if pPr is not None:
            el = pPr.find(qn('w:outlineLvl'))
            if el is not None:
                try:
                    lvl = int(el.get(qn('w:val'), 9))
                    if lvl < 9:
                        return lvl + 1
                except (ValueError, TypeError):
                    pass
        style = style.base_style
    return None


def _get_num_level(para):
    """
    Baca level dari w:numPr (multilevel list Word), kembalikan ilvl 0-based atau None.
    numId=0 berarti penghapusan numbering, diabaikan.
    """
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        return None
    numPr = pPr.find(qn('w:numPr'))
    if numPr is None:
        return None
    numId_el = numPr.find(qn('w:numId'))
    if numId_el is not None:
        try:
            if int(numId_el.get(qn('w:val'), 1)) == 0:
                return None
        except (ValueError, TypeError):
            pass
    ilvl_el = numPr.find(qn('w:ilvl'))
    if ilvl_el is None:
        return None
    try:
        return int(ilvl_el.get(qn('w:val'), 0))
    except (ValueError, TypeError):
        return None


def _get_dominant_font_size(para):
    """Kembalikan ukuran font terbesar (dalam pt) dari run paragraf, atau None."""
    sizes = [r.font.size for r in para.runs if r.font.size]
    if not sizes:
        return None
    try:
        from docx.shared import Pt
        return max(sizes).pt
    except Exception:
        return None


_FRONT_MATTER_RE = re.compile(
    r'^(KATA\s+PENGANTAR|PRAKATA|UCAPAN\s+TERIMA\s+KASIH|SANWACANA|'
    r'ABSTRAK|ABSTRACT|INTISARI|SARI|RINGKASAN|SUMMARY|'
    r'DAFTAR\s+(ISI|GAMBAR|TABEL|LAMPIRAN|SINGKATAN|SIMBOL|NOTASI|'
    r'ARTI\s+LAMBANG|PERSAMAAN)|'
    r'LEMBAR\s+(PERSETUJUAN|PENGESAHAN|PERNYATAAN|ORISINALITAS)(\s+\S+)*|'
    r'HALAMAN\s+(PERSETUJUAN|PENGESAHAN|PERNYATAAN|PERSEMBAHAN|'
    r'JUDUL|COVER|ORISINALITAS)(\s+\S+)*|'
    r'FORMULIR\s+(PERBAIKAN|REVISI|PERSETUJUAN)(\s+\S+)*|'
    r'PENGESAHAN\s+\S+(\s+\S+)*|'
    r'PERNYATAAN\s+(KEASLIAN|ORISINALITAS)(\s+\S+)*|'
    r'PENGAJUAN(\s+\S+)*|'
    r'PERSETUJUAN\s+\S+(\s+\S+)*|'
    r'MOTTO|PERSEMBAHAN|REFERENCES?|BIBLIOGRAPHY|'
    r'LAMPIRAN|APPENDIX|APPENDICES)$',
    re.IGNORECASE
)

_CUSTOM_H1_STYLES = {'H1', 'Heading Bab', 'Judul Bab', 'BAB', 'Bab', 'Judul'}
_CUSTOM_H2_STYLES = {'H2', 'Heading Sub', 'Sub Judul', 'Sub Bab', 'Sub-Bab'}
_CUSTOM_H3_STYLES = {'H3', 'Sub Sub Judul'}

# Label metadata cover/front matter yang TIDAK boleh masuk daftar isi
_META_LABEL_RE = re.compile(
    r'^(DOSEN\s+PENGAMPU|MATA\s+KULIAH|KELOMPOK|KELAS\b|SEMESTER|PROGRAM\s+STUDI|'
    r'JURUSAN|FAKULTAS|UNIVERSITAS|INSTITUTE|NIM\b|NPM\b|NIP\b|NIDN|NAMA\s+MAHASISWA|'
    r'DISUSUN\s+OLEH|DISUSUN|NAMA\s+KELOMPOK|ANGGOTA\s+KELOMPOK|PENGAMPU|'
    r'TAHUN\s+AJARAN|TAHUN\s+AKADEMIK|DOSEN\s+PEMBIMBING|PEMBIMBING)\s*:',
    re.IGNORECASE
)

_STYLE_KW_H1 = re.compile(
    r'(judul\s*(bab|utama|1)?|^bab$|heading\s*1?|chapter|title)$', re.IGNORECASE)
_STYLE_KW_H2 = re.compile(
    r'(sub[\s\-]?(judul|bab|heading)?[\s\-]?1?|heading\s*2|sub\s*title)$', re.IGNORECASE)
_STYLE_KW_H3 = re.compile(
    r'(sub[\s\-]?sub|heading\s*3|sub[\s\-]?(judul|bab)[\s\-]?2)$', re.IGNORECASE)


def get_para_level(para):
    """
    Kembalikan level heading (1-3) atau None.

    Patokan utama (berurutan):
      1. Pola BAB  → H1
      2. Pola numbering: 1.1.1→H3 | 1.1→H2 | 1.→H1
         (termasuk variasi spasi: "3. 10 Judul" → di-normalize dulu)
      3. Pola huruf: A.1→H3 | I.A→H2 | A.→H2
      4. Front matter keyword (KATA PENGANTAR, ABSTRAK, DAFTAR ISI, dst.) → H1

    Sengaja TIDAK mendeteksi berdasarkan Heading style, outlineLvl, atau numPr —
    terlalu banyak false positive dari body text yang salah format.
    Semua heading ber-Heading style yang tidak terdeteksi di sini akan di-exclude
    dari TOC oleh _demote_undetected_headings() di main().
    """
    text = para.text.strip()
    if not text or len(text) > 150:
        return None

    text_norm = re.sub(r'^(\d+)\.\s+(\d)', r'\1.\2', text)

    # 1. Pola BAB → H1
    if re.match(r'^\s*BAB\s+[IVXivx\d]+\b', text, re.IGNORECASE):
        return 1

    # 2. Pola numbering bertitik
    if re.match(r'^\d+\.\d+\.\d+\.?\s+\S', text_norm):
        return 3
    if re.match(r'^\d+\.\d+\.?\s+\S', text_norm):
        return 2
    if re.match(r'^\d+\.\s+\S', text):
        return 1

    # 3. Pola huruf dan kombinasi
    if re.match(r'^[A-Z]\.\d+\.?\s+\S', text):
        return 3
    if re.match(r'^[IVX]+\.[A-Z]\.?\s+\S', text):
        return 2
    if re.match(r'^[A-Z]\.\s+\S', text):
        return 2

    # 4. Front matter keyword (KATA PENGANTAR, ABSTRAK, DAFTAR ISI, dll.)
    if _FRONT_MATTER_RE.match(text):
        return 1

    # 5. Fallback: style chain Heading (untuk list-numbered heading seperti A./1./B.)
    # Teks tidak punya prefix nomor karena nomor dari Word list numbering.
    # Batasi ≤80 char agar body text panjang tidak ikut masuk.
    if len(text) <= 80:
        style = para.style
        while style:
            sname = style.name or ''
            if sname.startswith('Heading '):
                try:
                    return int(sname.split()[-1])
                except ValueError:
                    return 2
            style = style.base_style

    return None


def _dedup_bab_clusters(results, max_gap=10):
    """
    Hapus BAB palsu dari bagian sistematika / daftar isi pengantar.

    Pola sistematika: penulis menyebut "BAB I ..., BAB II ..., BAB III ..."
    berurutan dalam satu paragraf area → jarak antar BAB sangat kecil dan
    tidak ada H2/H3 di antara mereka.

    Algoritma: jika dua entry BAB berurutan berjarak ≤ max_gap paragraf DAN
    tidak ada heading H2/H3 di antara keduanya dalam results, entry pertama
    kemungkinan besar palsu → hapus. Entry TERAKHIR dalam rantai tersebut
    selalu dipertahankan (itu BAB asli yang memulai konten).
    """
    _bab = re.compile(r'^\s*BAB\s+[IVXivx\d]+\b', re.IGNORECASE)
    to_remove = set()

    for k in range(len(results) - 1):
        pi1, lvl1, txt1 = results[k]
        pi2, lvl2, txt2 = results[k + 1]

        if lvl1 != 1 or not _bab.match(txt1):
            continue
        if lvl2 != 1 or not _bab.match(txt2):
            continue

        gap = pi2 - pi1
        if gap > max_gap:
            continue

        # Cek apakah ada H2/H3 di antara k dan k+1 dalam results
        has_sub = any(lvl >= 2 for _, lvl, _ in results[k + 1: k + 1])
        # (selalu False karena slice kosong — benar, kita cek gap bukan results)
        has_sub = any(
            lvl >= 2
            for j in range(k + 1, len(results))
            if results[j][0] < pi2
            for _, lvl, _ in [results[j]]
        )
        if not has_sub:
            to_remove.add(k)  # hapus entry PERTAMA, pertahankan entry berikutnya

    return [entry for i, entry in enumerate(results) if i not in to_remove]


_ALPHA_H2_RE = re.compile(r'^[A-Z]\.\s+\S')  # A. Judul → H2 (model alfabet)
_SINGLE_DIG_RE = re.compile(r'^\d+\.\s+\S')  # 1. Judul → H1 (numeric) atau H3 (alpha)


def detect_headings(doc, max_level):
    """
    Kembalikan list of (para_index, level, text) untuk semua heading ≤ max_level.

    Mendukung dua model penomoran:
      Model 1 (numerik) : 1.1 → H2 | 1.1.1 → H3
      Model 2 (alfabet) : A. → H2 | 1. → H3 | B. → H2 | 1. → H3

    Context tracker: setelah A./B./C. terdeteksi sebagai H2, angka 1./2. berikutnya
    diperlakukan sebagai H3 (bukan H1). BAB baru atau H1 lain mereset konteks.
    """
    paras      = doc.paragraphs
    results    = []
    merged_idx = set()

    first_bab       = _find_first_bab_idx(paras)
    in_alpha_ctx    = False  # True setelah A./B./C. terdeteksi

    for i, para in enumerate(paras):
        if i in merged_idx:
            continue
        text = para.text.strip()
        if is_toc_entry(text):
            continue

        lvl = get_para_level(para)
        if lvl is None:
            continue

        # Zona sebelum BAB I: HANYA keyword front-matter yang lolos.
        if first_bab is not None and i < first_bab:
            if not _FRONT_MATTER_RE.match(text):
                continue

        # Dalam konteks alfabet: 1./2. adalah H3, bukan H1 — cek SEBELUM update context
        if in_alpha_ctx and lvl == 1 and _SINGLE_DIG_RE.match(text):
            lvl = 3

        # Dalam dokumen BAB-structure: 1./2./3. tanpa Heading style → bukan heading.
        # Paragraph Normal yang cocok digit-dot setelah BAB I adalah sub-item/label,
        # bukan judul bab baru (bab baru pakai BAB prefix).
        if (lvl == 1
                and _SINGLE_DIG_RE.match(text)
                and not _FRONT_MATTER_RE.match(text)
                and not re.match(r'^\s*BAB\s+', text, re.IGNORECASE)
                and first_bab is not None
                and i > first_bab):
            _style = para.style
            _has_hdg = False
            while _style:
                _sname = _style.name or ''
                if _sname.startswith('Heading ') or _sname in _ALL_HEADING_STYLES:
                    _has_hdg = True
                    break
                _style = _style.base_style
            if not _has_hdg:
                continue

        # ── Context tracker update ────────────────────────────────────────────
        if lvl == 2 and _ALPHA_H2_RE.match(text):
            in_alpha_ctx = True
        elif lvl == 1:
            in_alpha_ctx = False

        if lvl > max_level:
            continue

        # BAB X tanpa nama → gabung dengan paragraf ALL CAPS berikutnya
        if lvl == 1 and re.match(r'^\s*BAB\s+[IVXivx\d]+\s*$', text, re.IGNORECASE):
            for j in range(i + 1, min(i + 3, len(paras))):
                nt = paras[j].text.strip()
                if nt and nt == nt.upper() and len(nt) > 3:
                    text = text + ' ' + nt
                    merged_idx.add(j)
                    break

        # Paragraf front-matter sebelum BAB I yang berisi dua judul digabung
        # (misal "HALAMAN PENGESAHAN SKRIPSI HALAMAN PERSETUJUAN SKRIPSI"):
        # ambil segmen terakhir saja agar cocok dengan referensi Word.
        if first_bab is not None and i < first_bab:
            _fm_starts = list(re.finditer(r'\b(HALAMAN|LEMBAR|FORMULIR)\b', text, re.IGNORECASE))
            if len(_fm_starts) > 1:
                text = text[_fm_starts[-1].start():]

        results.append((i, lvl, text))

    results = _dedup_bab_clusters(results)

    # Deduplikasi front-matter sebelum BAB I (case-insensitive)
    if first_bab is not None:
        seen_fm = set()
        deduped = []
        for (pidx, lvl, txt) in results:
            if pidx < first_bab:
                key = txt.upper()
                if key in seen_fm:
                    continue
                seen_fm.add(key)
            deduped.append((pidx, lvl, txt))
        results = deduped

    return results


# ── Pra-proses struktur dokumen ───────────────────────────────────────────────

def _fix_bab_line_breaks(doc):
    """Ganti soft return (w:br) dalam paragraf BAB dengan spasi."""
    for para in doc.paragraphs:
        text = para.text.strip()
        if not re.match(r'^\s*BAB\s+[IVXivx\d]+\b', text, re.IGNORECASE):
            continue
        for r_el in para._p.findall(f'.//{qn("w:r")}'):
            for br in r_el.findall(qn('w:br')):
                idx_br = list(r_el).index(br)
                r_el.remove(br)
                t = OxmlElement('w:t')
                t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                t.text = ' '
                r_el.insert(idx_br, t)


def _merge_bab_name_paras(doc):
    """
    Untuk paragraf 'BAB X' yang diikuti nama bab ALL CAPS di paragraf terpisah:
    - Sisipkan soft return (w:br tanpa type = Shift+Enter) + salin runs nama
    - Hapus paragraf nama dari dokumen

    Di halaman: dua baris dalam satu paragraf — identik visual dengan aslinya
                karena alignment (center) dan font runs ikut terbawa.
    Di TOC   : satu entri 'BAB X NAMA' (Word mengabaikan soft return saat
                membuat entri TOC dari outline level).
    """
    import copy as _cp

    paras = doc.paragraphs
    to_remove = []

    for i, para in enumerate(paras):
        text = para.text.strip()
        if not re.match(r'^\s*BAB\s+[IVXivx\d]+\s*$', text, re.IGNORECASE):
            continue
        for j in range(i + 1, min(i + 3, len(paras))):
            nt = paras[j].text.strip()
            if not nt or nt != nt.upper() or len(nt) <= 3:
                continue
            # Soft return (Shift+Enter) — bukan page break, bukan paragraph break
            r_br = OxmlElement('w:r')
            br_el = OxmlElement('w:br')   # tanpa w:type → line break dalam paragraf
            r_br.append(br_el)
            para._p.append(r_br)
            # Salin semua runs nama ke paragraf BAB (font/size terjaga)
            for r_el in paras[j]._p.findall(qn('w:r')):
                para._p.append(_cp.deepcopy(r_el))
            # Tandai paragraf nama untuk dihapus (tidak muncul di halaman lagi)
            to_remove.append(paras[j]._p)
            break

    for p_el in to_remove:
        parent = p_el.getparent()
        if parent is not None:
            parent.remove(p_el)


def _set_outline_excluded(para):
    """Set outlineLvl=9 pada paragraf agar tidak masuk TOC field Word."""
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._p.insert(0, pPr)
    ol = pPr.find(qn('w:outlineLvl'))
    if ol is None:
        ol = OxmlElement('w:outlineLvl')
        pPr.append(ol)
    ol.set(qn('w:val'), '9')


def _get_or_create_cover_style(doc, orig_style_name):
    """
    Buat clone style dari orig_style_name dengan formatting identik,
    tapi tanpa outline level sehingga tidak masuk TOC field.
    Clone disimpan dengan nama '_Cover_<OrigStyleId>'.
    """
    import copy as _cp

    try:
        orig_style = doc.styles[orig_style_name]
    except KeyError:
        return None

    orig_id   = orig_style.element.get(qn('w:styleId'), orig_style_name.replace(' ', ''))
    clone_id  = f'_Cover_{orig_id}'
    clone_name = f'_Cover_{orig_style_name}'

    # Sudah ada dari panggilan sebelumnya?
    try:
        return doc.styles[clone_name]
    except KeyError:
        pass

    # Deep-copy elemen XML style asli
    clone_el = _cp.deepcopy(orig_style.element)

    # Ganti styleId
    clone_el.set(qn('w:styleId'), clone_id)

    # Ganti w:name
    name_el = clone_el.find(qn('w:name'))
    if name_el is not None:
        name_el.set(qn('w:val'), clone_name)
    else:
        n = OxmlElement('w:name')
        n.set(qn('w:val'), clone_name)
        clone_el.insert(0, n)

    # basedOn → Normal (hindari warisan outlineLvl dari Heading)
    based_el = clone_el.find(qn('w:basedOn'))
    if based_el is not None:
        based_el.set(qn('w:val'), 'Normal')

    # Hapus referensi 'next style' dan gallery flags
    for tag in (qn('w:next'), qn('w:qFormat'), qn('w:semiHidden'), qn('w:unhideWhenUsed')):
        el = clone_el.find(tag)
        if el is not None:
            clone_el.remove(el)

    # Pastikan outlineLvl=9 di pPr clone → tidak masuk TOC via \u
    pPr = clone_el.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        clone_el.append(pPr)
    ol = pPr.find(qn('w:outlineLvl'))
    if ol is not None:
        pPr.remove(ol)
    ol = OxmlElement('w:outlineLvl')
    ol.set(qn('w:val'), '9')
    pPr.append(ol)

    # Sisipkan ke styles root dokumen
    doc.styles.element.append(clone_el)

    try:
        return doc.styles[clone_name]
    except KeyError:
        return None


def _demote_heading_to_normal(para, doc):
    """
    Ganti style Heading ke clone style yang identik secara visual tapi
    tidak punya outline level — sehingga tidak masuk TOC field Word.
    outlineLvl=9 saja tidak cukup karena Word scan berdasarkan style definition.
    """
    orig_style_name = para.style.name if para.style else ''

    clone_style = _get_or_create_cover_style(doc, orig_style_name)
    if clone_style is not None:
        para.style = clone_style
    else:
        # Fallback: set outlineLvl=9 saja
        pass

    _set_outline_excluded(para)


def _find_first_bab_idx(paras):
    """Kembalikan index paragraf BAB pertama, atau None."""
    for i, p in enumerate(paras):
        if re.match(r'^\s*BAB\s+[IVXivx\d]+\b', p.text.strip(), re.IGNORECASE):
            return i
    return None


_ALL_HEADING_STYLES = (
    _CUSTOM_H1_STYLES | _CUSTOM_H2_STYLES | _CUSTOM_H3_STYLES
)


def _exclude_prefab_headings(doc):
    """
    Sebelum BAB I: semua paragraf ber-Heading style yang BUKAN judul section
    yang dikenal (KATA PENGANTAR, ABSTRAK, DAFTAR ISI, dll.) dikecualikan dari
    TOC field dengan outlineLvl=9.

    Logika: zona cover/front-matter hanya boleh berisi judul halaman resmi,
    bukan label metadata (Dosen Pengampu, NIM, Program Studi, dll.) — apapun
    style yang dipakai penulis.

    Fallback jika tidak ada BAB: gunakan _META_LABEL_RE (perilaku lama).
    """
    paras = doc.paragraphs
    first_bab = _find_first_bab_idx(paras)

    if first_bab is None:
        # Dokumen tanpa BAB — fallback ke blacklist metadata label
        for para in paras:
            text = para.text.strip()
            if not text or not _META_LABEL_RE.match(text):
                continue
            style_name = para.style.name if para.style else ''
            if style_name.startswith('Heading ') or style_name in _ALL_HEADING_STYLES:
                _demote_heading_to_normal(para, doc)
        return

    for i in range(first_bab):
        para = paras[i]
        style_name = para.style.name if para.style else ''
        is_heading = (
            style_name.startswith('Heading ')
            or style_name in _ALL_HEADING_STYLES
        )
        if not is_heading:
            continue
        text = para.text.strip()
        if not text:
            continue
        # Pertahankan judul section yang valid (KATA PENGANTAR, ABSTRAK, dll.)
        if _FRONT_MATTER_RE.match(text):
            continue
        # Semua heading lain sebelum BAB I → ganti ke Normal (exclude dari TOC)
        _demote_heading_to_normal(para, doc)


# ── Bold normalization untuk heading ─────────────────────────────────────────

def _strip_inline_bold_only(para):
    """Strip w:b/w:bCs dari pPr/rPr dan setiap run. TIDAK menambah w:b val=0."""
    BOLD_TAGS = (qn('w:b'), qn('w:bCs'))
    pPr = para._p.find(qn('w:pPr'))
    if pPr is not None:
        rPr = pPr.find(qn('w:rPr'))
        if rPr is not None:
            for tag in BOLD_TAGS:
                el = rPr.find(tag)
                if el is not None:
                    rPr.remove(el)
    for r_el in para._p.findall(qn('w:r')):
        rPr = r_el.find(qn('w:rPr'))
        if rPr is None:
            continue
        for tag in BOLD_TAGS:
            el = rPr.find(tag)
            if el is not None:
                rPr.remove(el)


def _set_runs_bold(para):
    """Tambahkan w:b/w:bCs ke semua runs (inline bold eksplisit)."""
    for r_el in para._p.findall(qn('w:r')):
        rPr = r_el.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            r_el.insert(0, rPr)
        for tag in (qn('w:b'), qn('w:bCs')):
            el = rPr.find(tag)
            if el is not None:
                rPr.remove(el)
        rPr.append(OxmlElement('w:b'))
        rPr.append(OxmlElement('w:bCs'))


def _normalize_heading_bold(para, level):
    """
    Normalisasi bold heading di body dan TOC.

    H1: tambahkan inline bold → body bold ✓, TOC bold ✓.
    H2/H3 + heading style (standard maupun custom H2/H3):
        strip inline bold saja → bold dari style, TOC tidak bold ✓.
    H2/H3 tanpa heading style: tambahkan inline bold → body bold ✓,
        TOC mungkin juga bold (limitation untuk non-standard style).
    """
    style_name = para.style.name if para.style else ''
    is_heading_style = (
        style_name.startswith('Heading ')
        or style_name in _CUSTOM_H1_STYLES
        or style_name in _CUSTOM_H2_STYLES
        or style_name in _CUSTOM_H3_STYLES
    )
    if level == 1:
        _set_runs_bold(para)
    else:
        _strip_inline_bold_only(para)
        if not is_heading_style:
            _set_runs_bold(para)


# ── Normalisasi list heading → typed ─────────────────────────────────────────

def _roman_to_int(s):
    vals = {'I': 1, 'V': 5, 'X': 10, 'L': 50, 'C': 100, 'D': 500, 'M': 1000}
    result, prev = 0, 0
    for ch in reversed(s.upper()):
        v = vals.get(ch, 0)
        result += v if v >= prev else -v
        prev = v
    return result


def _update_counters(counters, level, text):
    """Update counter array dari teks heading bertipe typed."""
    m3 = re.match(r'^(\d+)\.(\d+)\.(\d+)', text)
    if m3:
        counters[1], counters[2], counters[3] = int(m3.group(1)), int(m3.group(2)), int(m3.group(3))
        return
    m2 = re.match(r'^(\d+)\.(\d+)', text)
    if m2:
        counters[1] = int(m2.group(1))
        counters[2] = int(m2.group(2))
        counters[3] = 0
        return
    mb = re.match(r'^\s*BAB\s+([IVXivx]+)\b', text, re.IGNORECASE)
    if mb:
        counters[1] = _roman_to_int(mb.group(1))
        counters[2] = 0
        counters[3] = 0
        return
    m1 = re.match(r'^(\d+)\.', text)
    if m1 and level == 1:
        counters[1] = int(m1.group(1))
        counters[2] = 0
        counters[3] = 0


def _normalize_list_headings(doc, headings):
    """
    Konversi heading yang pakai multilevel list ke format typed number.

    File campuran (sebagian typed, sebagian list) menyebabkan TOC tidak rata
    karena list heading punya tab+indentasi berbeda dari typed heading.
    Solusi: hapus numPr dari list heading, sisipkan angka sebagai teks biasa.
    Angka dihitung dari konteks heading typed di sekitarnya.
    """
    counters = {1: 0, 2: 0, 3: 0}

    for idx, lvl, txt in headings:
        para    = doc.paragraphs[idx]
        p       = para._p
        pPr     = p.find(qn('w:pPr'))
        numPr   = pPr.find(qn('w:numPr')) if pPr is not None else None

        is_list = False
        if numPr is not None:
            nid_el = numPr.find(qn('w:numId'))
            if nid_el is not None:
                is_list = (nid_el.get(qn('w:val'), '0') != '0')

        if is_list:
            # Hitung nomor dari counter
            if lvl == 1:
                counters[1] += 1
                counters[2]  = 0
                counters[3]  = 0
                prefix = 'BAB %s ' % _int_to_roman(counters[1])
            elif lvl == 2:
                counters[2] += 1
                counters[3]  = 0
                prefix = '%d.%d ' % (counters[1], counters[2])
            else:
                counters[3] += 1
                prefix = '%d.%d.%d ' % (counters[1], counters[2], counters[3])

            # Sisipkan run prefix di awal paragraf
            r_pre = OxmlElement('w:r')
            # Salin rPr dari run pertama agar format konsisten
            first_r = p.find(qn('w:r'))
            if first_r is not None:
                first_rPr = first_r.find(qn('w:rPr'))
                if first_rPr is not None:
                    import copy
                    r_pre.append(copy.deepcopy(first_rPr))
            t_pre = OxmlElement('w:t')
            t_pre.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t_pre.text = prefix
            r_pre.append(t_pre)
            if first_r is not None:
                first_r.addprevious(r_pre)
            else:
                p.append(r_pre)

            # Hapus numPr dan nonaktifkan list dari style
            pPr.remove(numPr)
            # Tambah numId=0 agar list inherited dari style tidak aktif
            numPr_disable = OxmlElement('w:numPr')
            ilvl_el = OxmlElement('w:ilvl')
            ilvl_el.set(qn('w:val'), '0')
            numPr_disable.append(ilvl_el)
            numId_el = OxmlElement('w:numId')
            numId_el.set(qn('w:val'), '0')
            numPr_disable.append(numId_el)
            pPr.append(numPr_disable)

        else:
            _update_counters(counters, lvl, txt)


def _int_to_roman(n):
    vals = [(1000,'M'),(900,'CM'),(500,'D'),(400,'CD'),(100,'C'),(90,'XC'),
            (50,'L'),(40,'XL'),(10,'X'),(9,'IX'),(5,'V'),(4,'IV'),(1,'I')]
    result = ''
    for v, s in vals:
        while n >= v:
            result += s
            n -= v
    return result


# ── Terapkan outline level ─────────────────────────────────────────────────────

def apply_outline_level(para, level):
    """
    Tambahkan w:outlineLvl ke paragraf tanpa mengubah tampilan visual.
    level: 1-based. w:outlineLvl adalah 0-based.
    """
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._p.insert(0, pPr)

    existing = pPr.find(qn('w:outlineLvl'))
    if existing is not None:
        pPr.remove(existing)

    el = OxmlElement('w:outlineLvl')
    el.set(qn('w:val'), str(level - 1))
    pPr.append(el)


# ── TOC styles ────────────────────────────────────────────────────────────────

def _apply_rPr_font_to_style_element(style_el, font, size_pt, bold, is_char_style=False):
    """
    Terapkan font/size/bold ke elemen style XML secara menyeluruh.
    Menghapus semua theme-based font references (asciiTheme, hAnsiTheme, cstheme)
    yang bisa override font eksplisit saat Word update TOC field.

    is_char_style=True: hapus juga sz/color/kern agar tidak ada residual
    dari Heading X Char yang menimpa TOC entry saat dirender.
    """
    rPr = style_el.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        style_el.append(rPr)

    # Hapus tag yang bisa konflik
    remove_tags = [qn('w:rFonts'), qn('w:sz'), qn('w:szCs'),
                   qn('w:b'), qn('w:bCs'), qn('w:i'), qn('w:iCs')]
    if is_char_style:
        # Untuk char style, hapus juga color/kern/ligatures agar tidak
        # ada efek visual dari Heading Char yang ikut masuk ke TOC entry
        remove_tags += [qn('w:color'), qn('w:kern'), qn('w:lang')]
        # Hapus themeColor/themeShade attribute dari color jika ada
    for tag in remove_tags:
        old = rPr.find(tag)
        if old is not None:
            rPr.remove(old)

    # Font: set ascii + hAnsi + cs secara eksplisit (TANPA theme reference)
    # Theme reference seperti asciiTheme="majorHAnsi" akan di-resolve oleh Word
    # menjadi Calibri Light / Cambria, BUKAN Times New Roman.
    fonts_el = OxmlElement('w:rFonts')
    fonts_el.set(qn('w:ascii'),  font)
    fonts_el.set(qn('w:hAnsi'),  font)
    fonts_el.set(qn('w:cs'),     font)
    fonts_el.set(qn('w:eastAsia'), font)
    rPr.insert(0, fonts_el)

    # Ukuran font
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(size_pt * 2))
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(size_pt * 2))
    rPr.append(szCs)

    # Bold
    b = OxmlElement('w:b')
    if not bold:
        b.set(qn('w:val'), '0')
    rPr.append(b)
    bCs = OxmlElement('w:bCs')
    if not bold:
        bCs.set(qn('w:val'), '0')
    rPr.append(bCs)
    # Italic: selalu nonaktifkan di TOC / heading char style
    i_el = OxmlElement('w:i')
    i_el.set(qn('w:val'), '0')
    rPr.append(i_el)
    i_cs = OxmlElement('w:iCs')
    i_cs.set(qn('w:val'), '0')
    rPr.append(i_cs)


def _fix_heading_char_styles(doc, font, size_pt):
    """
    Override Heading X Char styles agar font konsisten dan bold terkontrol.

    - Heading 1 Char: bold=True (TOC H1 harus bold)
    - Heading 2/3+ Char: bold=False (TOC H2/H3 tidak bold)

    Saat Word update TOC field, ia menaruh w:rStyle="HeadingXChar" di setiap run
    TOC entry. Style ini bisa override font TOC style jika tidak di-fix.
    """
    styles_root = doc.styles.element

    heading_char_ids = {
        'Heading1Char', 'Heading2Char', 'Heading3Char',
        'Heading4Char', 'Heading5Char', 'Heading6Char',
        'Heading7Char', 'Heading8Char', 'Heading9Char',
        'H1Char', 'H2Char', 'H3Char',
    }
    h1_char_ids = {'Heading1Char', 'H1Char'}

    for style_el in styles_root.findall(qn('w:style')):
        stype = style_el.get(qn('w:type'), '')
        if stype != 'character':
            continue
        style_id = style_el.get(qn('w:styleId'), '')
        sname_el = style_el.find(qn('w:name'))
        sname = sname_el.get(qn('w:val'), '') if sname_el is not None else ''

        is_heading_char = (
            style_id in heading_char_ids
            or ('Heading' in sname and 'Char' in sname)
            or (style_id.startswith('H') and style_id.endswith('Char')
                and style_id[1:-4].isdigit())
        )
        if not is_heading_char:
            continue

        # H1 Char: bold=True; semua lainnya: bold=False
        is_h1_char = (
            style_id in h1_char_ids
            or sname.lower() in ('heading 1 char', 'heading1char')
        )
        _apply_rPr_font_to_style_element(
            style_el, font=font, size_pt=size_pt, bold=is_h1_char, is_char_style=True
        )


def _get_text_width_twips(doc):
    """
    Baca lebar area teks (page_width - left_margin - right_margin) dalam twips
    langsung dari XML sectPr — lebih robust dari python-docx Section API yang
    bisa return None untuk dokumen dengan margin non-standar.
    Fallback: 8505 twips (~15cm, A4 dengan margin 3cm kiri-kanan).
    """
    try:
        body = doc.element.body
        # sectPr bisa di body langsung atau di paragraf terakhir
        sect = body.find(qn('w:sectPr'))
        if sect is None:
            for child in reversed(list(body)):
                pPr = child.find(qn('w:pPr'))
                if pPr is not None:
                    sect = pPr.find(qn('w:sectPr'))
                    if sect is not None:
                        break
        if sect is None:
            return 8505
        pgSz  = sect.find(qn('w:pgSz'))
        pgMar = sect.find(qn('w:pgMar'))
        if pgSz is None or pgMar is None:
            return 8505
        w     = int(pgSz .get(qn('w:w'),    '12240'))
        left  = int(pgMar.get(qn('w:left'),  '1800'))
        right = int(pgMar.get(qn('w:right'), '1800'))
        return max(w - left - right, 4000)  # minimal 4000 twips sebagai guard
    except Exception:
        return 8505


def _ensure_toc_style(doc, level, use_dots, font='Times New Roman', size_pt=12, line_spacing=1.0, h3_left_tab=None):
    """
    Pastikan style 'TOC X' ada dengan font/size/spasi seragam + tab stop kanan.

    Menangani DUA versi style yang mungkin ada di dokumen Word Indonesia:
      - 'TOC 1' / 'TOC 2' / 'TOC 3'  (Pascal case — dipakai saat Word update field)
      - 'toc 1' / 'toc 2' / 'toc 3'  (lowercase — style bawaan Word Indonesia)

    Kedua versi harus punya font yang sama. Jika 'toc X' lowercase tidak punya
    rPr (seperti yang ditemukan di dokumen sumber), Word akan fallback ke Normal
    style yang fontnya Calibri — menyebabkan inkonsistensi.

    Menggunakan XML langsung agar semua script font (ascii/hAnsi/cs/eastAsia)
    ter-override secara eksplisit tanpa theme reference.
    """
    from docx.enum.style import WD_STYLE_TYPE

    # Tangani kedua versi nama style (Pascal dan lowercase)
    style_names_to_fix = [f'TOC {level}', f'toc {level}']

    for style_name in style_names_to_fix:
        try:
            style = doc.styles[style_name]
        except KeyError:
            if style_name == f'TOC {level}':
                # Hanya buat style TOC jika belum ada (versi Pascal case adalah utama)
                style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                try:
                    style.base_style = doc.styles['Normal']
                except Exception:
                    pass
            else:
                # Versi lowercase opsional — skip jika tidak ada
                continue

        # ── Paragraph format ─────────────────────────────────────────────────
        pf = style.paragraph_format
        pf.space_after  = Pt(0)
        pf.line_spacing = line_spacing

        # Justify: rata kanan-kiri sesuai standar dokumen akademik Indonesia
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Indentasi per level — TOC3 left_tab dinamis berdasarkan panjang angka H3
        # TOC2: left=709t, hang=425t -> nomor di 284t, teks di 709t
        # TOC3: left=h3_left_tab, hang=(h3_left_tab-709) -> nomor di 709t, teks di h3_left_tab
        #   993 untuk angka pendek, 1320 untuk multilevel (1.4.1)
        from docx.shared import Twips as _Twips
        _h3_tab = h3_left_tab if (level == 3 and h3_left_tab is not None) else 1320
        IND = {
            1: (0,    0),
            2: (709, 425),
            3: (_h3_tab, _h3_tab - 709),
        }
        left_t, hang_t = IND.get(level, (284 * (level - 1), 0))
        pf.left_indent       = _Twips(left_t)
        pf.first_line_indent = _Twips(-hang_t) if hang_t else None

        # ── Run properties via XML ────────────────────────────────────────────
        _apply_rPr_font_to_style_element(
            style.element,
            font=font,
            size_pt=size_pt,
            bold=(level == 1),
            is_char_style=False,
        )

        # ── Tab stops ─────────────────────────────────────────────────────────
        # Hitung posisi tab kanan (nomor halaman) dari margin dokumen.
        # Baca langsung dari XML sectPr agar tidak gagal karena python-docx
        # kadang kembalikan None untuk properti margin tertentu.
        right_tab = _get_text_width_twips(doc)

        # TOC2: DUA tab stop — tab kiri 709t + tab kanan (halaman)
        # TOC3: DUA tab stop — tab kiri _h3_tab (dinamis) + tab kanan (halaman)
        # TOC1: hanya tab kanan
        LEFT_TAB_TWIPS = {2: 709, 3: _h3_tab}

        pPr = style.element.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            style.element.append(pPr)

        old_tabs = pPr.find(qn('w:tabs'))
        if old_tabs is not None:
            pPr.remove(old_tabs)

        tabs_el = OxmlElement('w:tabs')

        if level in LEFT_TAB_TWIPS:
            left_tab = OxmlElement('w:tab')
            left_tab.set(qn('w:val'),    'left')
            left_tab.set(qn('w:pos'),    str(LEFT_TAB_TWIPS[level]))
            left_tab.set(qn('w:leader'), 'none')
            tabs_el.append(left_tab)

        right_tab_el = OxmlElement('w:tab')
        right_tab_el.set(qn('w:val'),    'right')
        right_tab_el.set(qn('w:pos'),    str(right_tab))
        right_tab_el.set(qn('w:leader'), 'dot' if use_dots else 'none')
        tabs_el.append(right_tab_el)

        pPr.append(tabs_el)


# ── Helper run ────────────────────────────────────────────────────────────────

def _make_run(text, bold=False, font='Times New Roman', size_pt=12):
    """Buat elemen <w:r> dengan formatting inline (semua script font di-set eksplisit)."""
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    fonts = OxmlElement('w:rFonts')
    fonts.set(qn('w:ascii'),   font)
    fonts.set(qn('w:hAnsi'),   font)
    fonts.set(qn('w:cs'),      font)
    fonts.set(qn('w:eastAsia'), font)
    rPr.append(fonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(size_pt * 2))
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(size_pt * 2))
    rPr.append(szCs)
    if bold:
        rPr.append(OxmlElement('w:b'))
    i_el = OxmlElement('w:i')
    i_el.set(qn('w:val'), '0')
    rPr.append(i_el)
    i_cs = OxmlElement('w:iCs')
    i_cs.set(qn('w:val'), '0')
    rPr.append(i_cs)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    return r


def _make_entry_pPr(level, use_dots):
    """Buat pPr untuk satu entri TOC."""
    pPr = OxmlElement('w:pPr')
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:after'), '0')
    pPr.append(spacing)
    if level > 1:
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), str((level - 1) * 360))
        pPr.append(ind)
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), TAB_POS_TWIPS)
    if use_dots:
        tab.set(qn('w:leader'), 'dot')
    tabs.append(tab)
    pPr.append(tabs)
    return pPr


def _make_toc_field_para(max_level):
    """
    Buat TOC field struktur standar Word (2 paragraf):
      Para 1: fldChar begin + instrText + fldChar separate
      Para 2: fldChar end  (paragraf terpisah, placeholder sebelum Word update)
    Struktur 2-para mencegah Word memindahkan fldChar end ke tempat yang salah
    (misalnya setelah page break) saat TOC di-update.
    """
    # Para 1: begin + instrText + separate
    p1 = OxmlElement('w:p')

    r_begin = OxmlElement('w:r')
    fc = OxmlElement('w:fldChar')
    fc.set(qn('w:fldCharType'), 'begin')
    fc.set(qn('w:dirty'), 'true')
    r_begin.append(fc)
    p1.append(r_begin)

    r_instr = OxmlElement('w:r')
    instr = OxmlElement('w:instrText')
    instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    instr.text = f' TOC \\o "1-{max_level}" \\h \\z \\u '
    r_instr.append(instr)
    p1.append(r_instr)

    r_sep = OxmlElement('w:r')
    fc_sep = OxmlElement('w:fldChar')
    fc_sep.set(qn('w:fldCharType'), 'separate')
    r_sep.append(fc_sep)
    p1.append(r_sep)

    # Para 2: end (paragraf terpisah)
    p2 = OxmlElement('w:p')
    r_end = OxmlElement('w:r')
    fc_end = OxmlElement('w:fldChar')
    fc_end.set(qn('w:fldCharType'), 'end')
    r_end.append(fc_end)
    p2.append(r_end)

    return [p1, p2]


# ── Static TOC (tanpa field) ──────────────────────────────────────────────────

def _add_paragraph_bookmark(para, bk_id, bk_name):
    """Sisipkan bookmark start/end di paragraf heading untuk dirujuk PAGEREF."""
    bkStart = OxmlElement('w:bookmarkStart')
    bkStart.set(qn('w:id'),   str(bk_id))
    bkStart.set(qn('w:name'), bk_name)
    bkEnd = OxmlElement('w:bookmarkEnd')
    bkEnd.set(qn('w:id'), str(bk_id))
    # Sisipkan setelah pPr (bukan di posisi 0 yang akan mendahului pPr)
    pPr = para._p.find(qn('w:pPr'))
    if pPr is not None:
        pPr.addnext(bkStart)
    else:
        para._p.insert(0, bkStart)
    para._p.append(bkEnd)


def _build_run_rPr(font, size_pt, bold):
    """Buat w:rPr dengan font/size/bold eksplisit — tanpa theme reference."""
    rPr = OxmlElement('w:rPr')
    fonts_el = OxmlElement('w:rFonts')
    fonts_el.set(qn('w:ascii'),    font)
    fonts_el.set(qn('w:hAnsi'),    font)
    fonts_el.set(qn('w:cs'),       font)
    fonts_el.set(qn('w:eastAsia'), font)
    rPr.append(fonts_el)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(size_pt * 2))
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(size_pt * 2))
    rPr.append(szCs)
    b = OxmlElement('w:b')
    if not bold:
        b.set(qn('w:val'), '0')
    rPr.append(b)
    bCs = OxmlElement('w:bCs')
    if not bold:
        bCs.set(qn('w:val'), '0')
    rPr.append(bCs)
    i_el = OxmlElement('w:i')
    i_el.set(qn('w:val'), '0')
    rPr.append(i_el)
    i_cs = OxmlElement('w:iCs')
    i_cs.set(qn('w:val'), '0')
    rPr.append(i_cs)
    return rPr


# ── Numbering prefix helpers ──────────────────────────────────────────────────

def _read_heading_num_info(docx_path):
    """
    Baca numbering.xml dan styles.xml untuk membangun map:
      style_id → (ilvl, fmt, template)

    Hanya untuk style yang punya numPr di definisi style (bukan di paragraf).
    Dipakai untuk menghitung prefix "A.", "1.", dll. saat membuat entri TOC.
    """
    try:
        from lxml import etree as _et
    except ImportError:
        return {}

    ns    = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    W     = f'{{{ns}}}'
    ns_map = {'w': ns}

    try:
        with zipfile.ZipFile(docx_path, 'r') as z:
            if 'word/numbering.xml' not in z.namelist():
                return {}
            with z.open('word/numbering.xml') as f:
                num_tree = _et.fromstring(f.read())
            with z.open('word/styles.xml') as f:
                styles_tree = _et.fromstring(f.read())
    except Exception:
        return {}

    # abstractNum map: absId → {ilvl: (fmt, template)}
    abs_map = {}
    for abs_el in num_tree.findall('w:abstractNum', ns_map):
        absId = abs_el.get(f'{W}abstractNumId', '')
        abs_map[absId] = {}
        for lvl_el in abs_el.findall('w:lvl', ns_map):
            ilvl    = int(lvl_el.get(f'{W}ilvl', 0))
            numFmt  = lvl_el.find('w:numFmt',  ns_map)
            lvlText = lvl_el.find('w:lvlText', ns_map)
            fmt  = numFmt.get(f'{W}val', '')  if numFmt  is not None else ''
            tmpl = lvlText.get(f'{W}val', '') if lvlText is not None else ''
            abs_map[absId][ilvl] = (fmt, tmpl)

    # numId → abstractNumId
    num_to_abs = {}
    for num_el in num_tree.findall('w:num', ns_map):
        numId  = num_el.get(f'{W}numId', '')
        absRef = num_el.find('w:abstractNumId', ns_map)
        if absRef is not None:
            num_to_abs[numId] = absRef.get(f'{W}val', '')

    # style_id → (ilvl, fmt, template)
    result = {}
    for style_el in styles_tree.findall('.//w:style', ns_map):
        sid   = style_el.get(f'{W}styleId', '')
        numPr = style_el.find(f'.//{W}numPr')
        if numPr is None:
            continue
        numId_el = numPr.find(f'{W}numId')
        ilvl_el  = numPr.find(f'{W}ilvl')
        if numId_el is None:
            continue
        numId = numId_el.get(f'{W}val', '0')
        if numId == '0':
            continue
        ilvl  = int(ilvl_el.get(f'{W}val', 0)) if ilvl_el is not None else 0
        absId = num_to_abs.get(numId, '')
        if not absId or absId not in abs_map:
            continue
        if ilvl not in abs_map[absId]:
            continue
        fmt, tmpl = abs_map[absId][ilvl]
        result[sid] = (ilvl, fmt, tmpl)

    return result


def _format_num(count, fmt):
    """Format count sesuai numFmt Word (upperLetter, decimal, upperRoman, dst.)."""
    if fmt == 'upperLetter':
        if count <= 26:
            return chr(64 + count)
        q, r = divmod(count - 1, 26)
        return chr(64 + q) + chr(65 + r)
    if fmt == 'lowerLetter':
        if count <= 26:
            return chr(96 + count)
        q, r = divmod(count - 1, 26)
        return chr(96 + q) + chr(97 + r)
    if fmt == 'decimal':
        return str(count)
    if fmt in ('upperRoman', 'lowerRoman'):
        val, result = count, ''
        for n, r in [(1000,'M'),(900,'CM'),(500,'D'),(400,'CD'),(100,'C'),
                     (90,'XC'),(50,'L'),(40,'XL'),(10,'X'),(9,'IX'),
                     (5,'V'),(4,'IV'),(1,'I')]:
            while val >= n:
                result += r
                val   -= n
        return result if fmt == 'upperRoman' else result.lower()
    return str(count)


def _get_num_prefix(style_id, level, num_info, counters):
    """
    Hitung prefix dari style-level numPr.
    Menginkremen counter level saat ini; counters adalah dict mutable.
    """
    if style_id not in num_info:
        return ''
    _ilvl, fmt, tmpl = num_info[style_id]
    counters[level] = counters.get(level, 0) + 1
    count  = counters[level]
    prefix = re.sub(r'%\d+', _format_num(count, fmt), tmpl, count=1)
    return prefix


def _read_num_def_map(docx_path):
    """
    Baca numbering.xml, kembalikan dict:
      (numId_str, ilvl_int) → (fmt_str, tmpl_str)

    Berguna untuk lookup format berdasarkan para-level numPr yang tidak ikut
    terdapat di style definition (misalnya Style3 dengan numPr di paragraf).
    """
    try:
        from lxml import etree as _et
    except ImportError:
        return {}

    ns    = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    W     = f'{{{ns}}}'
    ns_map = {'w': ns}

    try:
        with zipfile.ZipFile(docx_path, 'r') as z:
            if 'word/numbering.xml' not in z.namelist():
                return {}
            with z.open('word/numbering.xml') as f:
                num_tree = _et.fromstring(f.read())
    except Exception:
        return {}

    # absId → {ilvl: (fmt, tmpl)}
    abs_map = {}
    for abs_el in num_tree.findall('w:abstractNum', ns_map):
        absId = abs_el.get(f'{W}abstractNumId', '')
        abs_map[absId] = {}
        for lvl_el in abs_el.findall('w:lvl', ns_map):
            ilvl    = int(lvl_el.get(f'{W}ilvl', 0))
            numFmt  = lvl_el.find('w:numFmt',  ns_map)
            lvlText = lvl_el.find('w:lvlText', ns_map)
            fmt  = numFmt.get(f'{W}val', '')  if numFmt  is not None else ''
            tmpl = lvlText.get(f'{W}val', '') if lvlText is not None else ''
            abs_map[absId][ilvl] = (fmt, tmpl)

    result = {}
    for num_el in num_tree.findall('w:num', ns_map):
        numId  = num_el.get(f'{W}numId', '')
        absRef = num_el.find('w:abstractNumId', ns_map)
        if absRef is None:
            continue
        absId = absRef.get(f'{W}val', '')
        if absId not in abs_map:
            continue
        for ilvl, (fmt, tmpl) in abs_map[absId].items():
            result[(numId, ilvl)] = (fmt, tmpl)

    return result


def _get_effective_num_prefix(para, num_def_map, style_num_info, level, counters):
    """
    Hitung prefix penomoran untuk satu paragraf heading.
    Cek para-level numPr dulu, lalu style-level.
    Menginkremen counters[level].
    """
    fmt, tmpl = '', ''

    # 1. Para-level numPr (lebih spesifik dari style)
    pPr = para._p.find(qn('w:pPr'))
    if pPr is not None:
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            numId_el = numPr.find(qn('w:numId'))
            ilvl_el  = numPr.find(qn('w:ilvl'))
            if numId_el is not None:
                nid  = numId_el.get(qn('w:val'), '0')
                ilvl = int(ilvl_el.get(qn('w:val'), 0)) if ilvl_el is not None else 0
                if nid != '0' and (nid, ilvl) in num_def_map:
                    fmt, tmpl = num_def_map[(nid, ilvl)]

    # 2. Style-level numPr sebagai fallback
    if not fmt:
        style_id = ''
        if para.style and para.style.element is not None:
            style_id = para.style.element.get(qn('w:styleId'), '')
        if style_id in style_num_info:
            _ilvl, fmt, tmpl = style_num_info[style_id]

    if not fmt:
        return ''

    counters[level] = counters.get(level, 0) + 1
    count  = counters[level]
    prefix = re.sub(r'%\d+', _format_num(count, fmt), tmpl, count=1)
    return prefix


def _load_numbering_full(docx_path):
    """
    Returns (num_to_abs, abs_starts, abs_fmttmpl, override_starts):
      num_to_abs:      {numId_str: absId_str}
      abs_starts:      {absId_str: {ilvl_int: start_int}}
      abs_fmttmpl:     {absId_str: {ilvl_int: (fmt_str, tmpl_str)}}
      override_starts: {numId_str: {ilvl_int: start_int}}  # dari lvlOverride/startOverride
    """
    try:
        from lxml import etree as _et
    except ImportError:
        return {}, {}, {}, {}

    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    W  = f'{{{ns}}}'

    try:
        with zipfile.ZipFile(docx_path, 'r') as z:
            if 'word/numbering.xml' not in z.namelist():
                return {}, {}, {}, {}
            with z.open('word/numbering.xml') as f:
                num_tree = _et.fromstring(f.read())
    except Exception:
        return {}, {}, {}, {}

    abs_starts  = {}
    abs_fmttmpl = {}
    for abs_el in num_tree.findall(f'{W}abstractNum'):
        absId = abs_el.get(f'{W}abstractNumId', '')
        abs_starts[absId]  = {}
        abs_fmttmpl[absId] = {}
        for lvl_el in abs_el.findall(f'{W}lvl'):
            ilvl    = int(lvl_el.get(f'{W}ilvl', 0))
            numFmt  = lvl_el.find(f'{W}numFmt')
            lvlText = lvl_el.find(f'{W}lvlText')
            start   = lvl_el.find(f'{W}start')
            fmt  = numFmt.get(f'{W}val', '')  if numFmt  is not None else ''
            tmpl = lvlText.get(f'{W}val', '') if lvlText is not None else ''
            sv   = int(start.get(f'{W}val', 1)) if start is not None else 1
            abs_starts[absId][ilvl]  = sv
            abs_fmttmpl[absId][ilvl] = (fmt, tmpl)

    num_to_abs      = {}
    override_starts = {}
    for num_el in num_tree.findall(f'{W}num'):
        numId  = num_el.get(f'{W}numId', '')
        absRef = num_el.find(f'{W}abstractNumId')
        if absRef is None:
            continue
        num_to_abs[numId] = absRef.get(f'{W}val', '')
        ovr = {}
        for ov in num_el.findall(f'{W}lvlOverride'):
            ol = int(ov.get(f'{W}ilvl', 0))
            so = ov.find(f'{W}startOverride')
            if so is not None:
                ovr[ol] = int(so.get(f'{W}val', 1))
        if ovr:
            override_starts[numId] = ovr

    return num_to_abs, abs_starts, abs_fmttmpl, override_starts


def _compute_list_prefixes(headings, doc, docx_path):
    """
    Scan semua paragraf (termasuk di dalam tabel) sambil melacak counter list
    per abstractNum. Mengembalikan {para_idx: prefix_str} HANYA untuk heading
    yang menggunakan template multilevel (mengandung 2+ placeholder %N).
    Template single-level tetap ditangani oleh _get_effective_num_prefix.
    """
    import re as _re

    num_to_abs, abs_starts, abs_fmttmpl, override_starts = _load_numbering_full(docx_path)
    if not num_to_abs:
        return {}

    # Peta id(para._p) → para_idx untuk heading paragraf.
    # Simpan referensi ke elemen agar proxy lxml tidak di-GC sebelum iterasi,
    # sehingga id() tetap konsisten antara tahap registrasi dan pencarian.
    heading_elems  = {}
    _held_elements = []  # cegah GC proxy lxml
    for (para_idx, _level, _text) in headings:
        p_el = doc.paragraphs[para_idx]._p
        _held_elements.append(p_el)
        heading_elems[id(p_el)] = para_idx

    BULLET_FMTS = {'bullet', 'none', ''}

    def _eff_start(absId, ilvl, numId):
        ovr = override_starts.get(numId, {})
        if ilvl in ovr:
            return ovr[ilvl]
        return abs_starts.get(absId, {}).get(ilvl, 1)

    abs_counters = {}  # {absId: {ilvl: int or None}}

    def _increment(absId, ilvl, numId):
        if absId not in abs_counters:
            abs_counters[absId] = {}
        cur   = abs_counters[absId].get(ilvl)
        start = _eff_start(absId, ilvl, numId)
        abs_counters[absId][ilvl] = start if cur is None else cur + 1
        # Reset semua sub-level ke start-1 (akan naik ke start saat next increment)
        for sub in list(abs_counters[absId].keys()):
            if sub > ilvl:
                abs_counters[absId][sub] = _eff_start(absId, sub, numId) - 1

    def _read(absId, ilvl, numId):
        val = abs_counters.get(absId, {}).get(ilvl)
        return _eff_start(absId, ilvl, numId) if val is None else val

    def _resolve(absId, ilvl, numId):
        fmt, tmpl = abs_fmttmpl.get(absId, {}).get(ilvl, ('', ''))
        if fmt in BULLET_FMTS:
            return ''

        def _repl(m):
            ref_ilvl = int(m.group(1)) - 1          # %1 → ilvl=0
            ref_val  = _read(absId, ref_ilvl, numId)
            ref_fmt  = abs_fmttmpl.get(absId, {}).get(ref_ilvl, ('decimal', ''))[0]
            return _format_num(ref_val, ref_fmt)

        return _re.sub(r'%(\d+)', _repl, tmpl)

    result = {}
    for para_xml in doc.element.body.iter(qn('w:p')):
        pPr   = para_xml.find(qn('w:pPr'))
        if pPr is None:
            continue
        numPr = pPr.find(qn('w:numPr'))
        if numPr is None:
            continue
        nid_el = numPr.find(qn('w:numId'))
        ilv_el = numPr.find(qn('w:ilvl'))
        if nid_el is None:
            continue
        nid  = nid_el.get(qn('w:val'), '0')
        if nid == '0':
            continue
        ilvl = int(ilv_el.get(qn('w:val'), 0)) if ilv_el is not None else 0

        absId = num_to_abs.get(nid, '')
        if not absId:
            continue

        _increment(absId, ilvl, nid)

        elem_id = id(para_xml)
        if elem_id in heading_elems:
            fmt, tmpl = abs_fmttmpl.get(absId, {}).get(ilvl, ('', ''))
            if len(_re.findall(r'%\d+', tmpl)) >= 2:
                prefix = _resolve(absId, ilvl, nid)
                if prefix:
                    result[heading_elems[elem_id]] = prefix

    return result


def _find_max_bookmark_id(doc):
    """Kembalikan ID terbesar dari semua bookmarkStart di dokumen."""
    max_id = 0
    for bks in doc.element.body.iter(qn('w:bookmarkStart')):
        try:
            max_id = max(max_id, int(bks.get(qn('w:id'), '0')))
        except (ValueError, TypeError):
            pass
    return max_id


# ── Builder TOC content pre-populated ────────────────────────────────────────

def _prepend_toc_field_begin(p_elem, max_level):
    """
    Prepend fldChar(begin) + instrText("TOC ...") + fldChar(separate)
    ke elemen paragraf — menjadi awal dari outer TOC field yang mencakup
    semua entry paragraf hingga fldChar(end) di paragraf terakhir.
    """
    r_begin = OxmlElement('w:r')
    fc = OxmlElement('w:fldChar')
    fc.set(qn('w:fldCharType'), 'begin')
    fc.set(qn('w:dirty'), 'true')
    r_begin.append(fc)

    r_instr = OxmlElement('w:r')
    instr = OxmlElement('w:instrText')
    instr.set(_XML_SPACE, 'preserve')
    instr.text = f' TOC \\o "1-{max_level}" \\h \\z \\u '
    r_instr.append(instr)

    r_sep = OxmlElement('w:r')
    fc_sep = OxmlElement('w:fldChar')
    fc_sep.set(qn('w:fldCharType'), 'separate')
    r_sep.append(fc_sep)

    # Sisipkan setelah pPr (jika ada) agar urutan XML valid
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is not None:
        children   = list(p_elem)
        insert_pos = children.index(pPr) + 1
    else:
        insert_pos = 0

    p_elem.insert(insert_pos,     r_begin)
    p_elem.insert(insert_pos + 1, r_instr)
    p_elem.insert(insert_pos + 2, r_sep)


def _create_toc_sdt(toc_paras):
    """
    Buat elemen SDT (Content Control) baru berisi TOCHeading + toc_paras.
    Dipakai ketika dokumen tidak memiliki SDT placeholder setelah DAFTAR ISI.
    """
    sdt = OxmlElement('w:sdt')

    sdtPr   = OxmlElement('w:sdtPr')
    tag_el  = OxmlElement('w:tag')
    tag_el.set(qn('w:val'), 'Contents')
    sdtPr.append(tag_el)
    sdt.append(sdtPr)

    sdtContent = OxmlElement('w:sdtContent')
    sdt.append(sdtContent)

    # TOCHeading paragraph (kosong)
    p_hd    = OxmlElement('w:p')
    pPr_hd  = OxmlElement('w:pPr')
    pSt_hd  = OxmlElement('w:pStyle')
    pSt_hd.set(qn('w:val'), 'TOCHeading')
    pPr_hd.append(pSt_hd)
    p_hd.append(pPr_hd)
    sdtContent.append(p_hd)

    for p in toc_paras:
        sdtContent.append(p)

    return sdt


def _build_toc_content(headings, doc, docx_path, font, size_pt, use_dots,
                       right_tab_pos, max_level):
    """
    Bangun daftar paragraf TOC pre-populated dengan bookmark + PAGEREF.

    Struktur yang dikembalikan (untuk diisi ke sdtContent setelah TOCHeading):
      [0]      : entri pertama, outer TOC field begin+instrText+separate di-prepend
      [1..N-1] : entri-entri berikutnya (TOC1/2/3 dengan PAGEREF per entry)
      [N]      : paragraf penutup berisi fldChar(end) outer TOC field

    Setiap entri mengandung teks heading (dengan prefix "A.", "1." jika ada) +
    tab + PAGEREF field ke bookmark yang disisipkan di paragraf heading asli.
    """
    import copy as _cp

    num_info      = _read_heading_num_info(docx_path)
    num_def_map   = _read_num_def_map(docx_path)
    para_list_pfx = _compute_list_prefixes(headings, doc, docx_path)
    bk_id_base    = _find_max_bookmark_id(doc) + 1
    counters      = {}  # {level: current_count}

    entry_paras = []
    for seq, (para_idx, level, text) in enumerate(headings):
        para = doc.paragraphs[para_idx]

        # Reset counter sub-level saat heading yang lebih tinggi muncul
        if level <= 1:
            counters.clear()
        elif level == 2:
            for k in list(counters.keys()):
                if k > 2:
                    del counters[k]
        elif level == 3:
            for k in list(counters.keys()):
                if k > 3:
                    del counters[k]

        if para_idx in para_list_pfx:
            num_text     = para_list_pfx[para_idx]
            content_text = text.replace('\n', ' ')
        else:
            prefix = _get_effective_num_prefix(para, num_def_map, num_info, level, counters)
            if prefix:
                num_text     = prefix
                content_text = text.replace('\n', ' ')
            else:
                # Tidak ada numPr prefix — cek apakah teks diawali nomor bertitik
                # seperti "1.1 Latar Belakang" (diketik manual di dokumen asli)
                t_clean = text.replace('\n', ' ')
                m = re.match(r'^((?:\d+\.)+\d+) (.+)', t_clean)
                if m:
                    num_text, content_text = m.group(1), m.group(2)
                else:
                    m2 = re.match(r'^([A-Z]\.) (.+)', t_clean)
                    if m2:
                        num_text, content_text = m2.group(1), m2.group(2)
                    else:
                        num_text, content_text = '', t_clean

        bk_id   = bk_id_base + seq
        bk_name = f'_ADKToc{bk_id}'

        _add_paragraph_bookmark(para, bk_id, bk_name)

        p = _make_static_toc_para(
            num_text, content_text, level, bk_name, font, size_pt, use_dots, right_tab_pos
        )
        entry_paras.append(p)

    if not entry_paras:
        # Fallback jika tidak ada heading terdeteksi
        return _make_toc_field_para(max_level)

    # Prepend outer TOC field begin ke entri pertama
    _prepend_toc_field_begin(entry_paras[0], max_level)

    # Paragraf penutup: fldChar(end) outer TOC field
    p_close = OxmlElement('w:p')
    r_end   = OxmlElement('w:r')
    r_end.append(_build_run_rPr(font, size_pt, False))
    fc_end  = OxmlElement('w:fldChar')
    fc_end.set(qn('w:fldCharType'), 'end')
    r_end.append(fc_end)
    p_close.append(r_end)
    entry_paras.append(p_close)

    return entry_paras


def _make_static_toc_para(num_text, content_text, level, bk_name, font, size_pt,
                          use_dots, right_tab_pos):
    """
    Buat satu paragraf TOC statis dengan PAGEREF untuk nomor halaman.

    Keuntungan vs TOC field:
    - Formatting (font, bold) sepenuhnya dikontrol kita via rPr inline
    - Tidak ada warisan bold/font dari paragraf sumber (DAF2 problem)
    - Nomor halaman tetap dinamis via PAGEREF — diupdate Word saat dibuka

    Jika num_text non-kosong, struktur: <t>num_text</t><TAB/><t>content_text</t>
    Jika num_text kosong: <t>content_text</t>
    """
    import copy
    p       = OxmlElement('w:p')
    is_bold = (level == 1)

    # ── pPr ──────────────────────────────────────────────────────────────────
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), f'TOC{level}')
    pPr.append(pStyle)
    p.append(pPr)

    # ── Teks heading ──────────────────────────────────────────────────────────
    if num_text:
        # Run nomor
        r_num = OxmlElement('w:r')
        r_num.append(_build_run_rPr(font, size_pt, is_bold))
        t_num = OxmlElement('w:t')
        t_num.set(_XML_SPACE, 'preserve')
        t_num.text = num_text
        r_num.append(t_num)
        p.append(r_num)
        # Tab pemisah nomor-konten
        r_tab_sep = OxmlElement('w:r')
        r_tab_sep.append(_build_run_rPr(font, size_pt, is_bold))
        r_tab_sep.append(OxmlElement('w:tab'))
        p.append(r_tab_sep)

    r_text = OxmlElement('w:r')
    r_text.append(_build_run_rPr(font, size_pt, is_bold))
    t = OxmlElement('w:t')
    t.set(_XML_SPACE, 'preserve')
    t.text = content_text
    r_text.append(t)
    p.append(r_text)

    # ── Tab ───────────────────────────────────────────────────────────────────
    r_tab = OxmlElement('w:r')
    r_tab.append(_build_run_rPr(font, size_pt, is_bold))
    r_tab.append(OxmlElement('w:tab'))
    p.append(r_tab)

    # ── PAGEREF field ─────────────────────────────────────────────────────────
    base_rPr = _build_run_rPr(font, size_pt, is_bold)

    r_begin = OxmlElement('w:r')
    r_begin.append(copy.deepcopy(base_rPr))
    fc = OxmlElement('w:fldChar')
    fc.set(qn('w:fldCharType'), 'begin')
    r_begin.append(fc)
    p.append(r_begin)

    r_instr = OxmlElement('w:r')
    r_instr.append(copy.deepcopy(base_rPr))
    instr = OxmlElement('w:instrText')
    instr.set(_XML_SPACE, 'preserve')
    instr.text = f' PAGEREF {bk_name} \\h '
    r_instr.append(instr)
    p.append(r_instr)

    r_sep = OxmlElement('w:r')
    r_sep.append(copy.deepcopy(base_rPr))
    fc_sep = OxmlElement('w:fldChar')
    fc_sep.set(qn('w:fldCharType'), 'separate')
    r_sep.append(fc_sep)
    p.append(r_sep)

    r_ph = OxmlElement('w:r')
    r_ph.append(copy.deepcopy(base_rPr))
    t_ph = OxmlElement('w:t')
    t_ph.text = '1'
    r_ph.append(t_ph)
    p.append(r_ph)

    r_end = OxmlElement('w:r')
    r_end.append(copy.deepcopy(base_rPr))
    fc_end = OxmlElement('w:fldChar')
    fc_end.set(qn('w:fldCharType'), 'end')
    r_end.append(fc_end)
    p.append(r_end)

    return p


# ── Helper page break ─────────────────────────────────────────────────────────

def _has_inline_page_break(para):
    for br in para._p.findall('.//' + qn('w:br')):
        if br.get(qn('w:type')) == 'page':
            return True
    return False


def _remove_inline_page_break(para):
    for r_elem in list(para._p.findall(qn('w:r'))):
        br = r_elem.find(qn('w:br'))
        if br is not None and br.get(qn('w:type')) == 'page':
            para._p.remove(r_elem)


def _make_page_break_para():
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    p.append(r)
    return p


def _find_sdt_after_daftar(body, daftar_p_elem):
    """
    Cari SDT yang langsung setelah paragraf DAFTAR ISI di body children.
    Kembalikan (sdt_element, has_toc_content) atau (None, False).
    has_toc_content=True jika SDT sudah berisi TOC field/konten.
    """
    children = list(body)
    daftar_child_idx = None
    for i, child in enumerate(children):
        if child is daftar_p_elem:
            daftar_child_idx = i
            break
    if daftar_child_idx is None:
        return None, False
    next_idx = daftar_child_idx + 1
    if next_idx >= len(children):
        return None, False
    next_child = children[next_idx]
    if next_child.tag != qn('w:sdt'):
        return None, False
    has_toc = bool(next_child.findall(f'.//{qn("w:instrText")}'))
    return next_child, has_toc


def _populate_sdt_with_toc(sdt, toc_paras):
    """
    Isi sdtContent SDT dengan TOCHeading + TOC field paragraphs.
    Konten lama dihapus dulu.
    """
    sdtContent = sdt.find(qn('w:sdtContent'))
    if sdtContent is None:
        sdtContent = OxmlElement('w:sdtContent')
        sdt.append(sdtContent)
    for child in list(sdtContent):
        sdtContent.remove(child)

    # Tambahkan TOCHeading paragraph (header row kosong)
    p_heading = OxmlElement('w:p')
    pPr_h = OxmlElement('w:pPr')
    pStyle_h = OxmlElement('w:pStyle')
    pStyle_h.set(qn('w:val'), 'TOCHeading')
    pPr_h.append(pStyle_h)
    p_heading.append(pPr_h)
    sdtContent.append(p_heading)

    for p in toc_paras:
        sdtContent.append(p)


def _remove_old_toc_block(doc, daftar_idx):
    """Hapus seluruh block TOC field lama (dari fldChar begin hingga fldChar end)."""
    paras = doc.paragraphs
    start_j = None
    for _j in range(daftar_idx + 1, min(daftar_idx + 20, len(paras))):
        if _has_toc_field(paras[_j]._p):
            start_j = _j
            break
    if start_j is None:
        return
    # Kumpulkan semua paragraf dari begin hingga fldChar end
    depth = 0
    end_j = start_j
    for _j in range(start_j, min(start_j + 300, len(paras))):
        p = paras[_j]._p
        for fld in p.iter(qn('w:fldChar')):
            ft = fld.get(qn('w:fldCharType'), '')
            if ft == 'begin':
                depth += 1
            elif ft == 'end':
                depth -= 1
        end_j = _j
        if depth <= 0:
            break
    # Hapus semua paragraf dari end ke start (reverse agar indeks tidak bergeser)
    to_remove = [paras[_j]._p for _j in range(start_j, end_j + 1)]
    parent = to_remove[0].getparent() if to_remove else None
    if parent is not None:
        for p_el in to_remove:
            if p_el.getparent() is not None:
                p_el.getparent().remove(p_el)


def _remove_trailing_empty_paras(after_el):
    """Hapus paragraf kosong dan sisa TOC field berturut-turut setelah after_el."""
    nxt = after_el.getnext()
    while nxt is not None:
        if not nxt.tag.endswith('}p'):
            break
        # Jangan hapus jika ada sectPr (section break)
        pPr = nxt.find(qn('w:pPr'))
        if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
            break
        # TOC field block sisa (ada instrText TOC) → hapus seluruh block begin→end
        if _has_toc_field(nxt):
            depth = 0
            block = []
            cur = nxt
            while cur is not None and cur.tag.endswith('}p'):
                block.append(cur)
                for fld in cur.iter(qn('w:fldChar')):
                    ft = fld.get(qn('w:fldCharType'), '')
                    if ft == 'begin':
                        depth += 1
                    elif ft == 'end':
                        depth -= 1
                if depth <= 0:
                    break
                cur = cur.getnext()
            nxt = block[-1].getnext() if block else nxt.getnext()
            for p_el in block:
                if p_el.getparent() is not None:
                    p_el.getparent().remove(p_el)
            continue
        # Cek apakah ada teks atau konten khusus
        text = ''.join(t.text or '' for t in nxt.findall(f'.//{qn("w:t")}')).strip()
        fld_chars  = nxt.findall(f'.//{qn("w:fldChar")}')
        instr_texts = nxt.findall(f'.//{qn("w:instrText")}')
        all_brs    = nxt.findall(f'.//{qn("w:br")}')
        page_brs   = [b for b in all_brs if b.get(qn('w:type')) == 'page']
        other_brs  = [b for b in all_brs if b.get(qn('w:type')) != 'page']
        has_drawing = bool(nxt.findall(f'.//{qn("w:drawing")}') or nxt.findall(f'.//{qn("w:object")}'))

        # Orphaned fldChar (sisa end/separate tanpa instrText dan tanpa teks) → hapus
        if not text and fld_chars and not instr_texts and not all_brs and not has_drawing:
            to_remove = nxt
            nxt = nxt.getnext()
            to_remove.getparent().remove(to_remove)
            continue
        # Paragraf kosong berisi hanya page break → redundant setelah pb kita, hapus
        if not text and page_brs and not other_brs and not fld_chars and not has_drawing:
            to_remove = nxt
            nxt = nxt.getnext()
            to_remove.getparent().remove(to_remove)
            continue
        if text or other_brs or has_drawing or (fld_chars and instr_texts):
            break
        to_remove = nxt
        nxt = nxt.getnext()
        to_remove.getparent().remove(to_remove)


# ── Cari posisi DAFTAR ISI ────────────────────────────────────────────────────

def find_daftar_isi_idx(doc):
    """Kembalikan index paragraf 'DAFTAR ISI', atau None jika tidak ada.
    Coba exact match dulu, fallback ke fuzzy untuk toleransi typo."""
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if re.match(r'^\s*daftar\s+isi\s*$', text, re.IGNORECASE):
            return i
        if len(text) >= 4 and _fuzzy_str(text.lower(), 'daftar isi', threshold=0.85):
            return i
    return None


# ── Auto-update fields on open ────────────────────────────────────────────────

def enable_auto_update_fields(doc):
    """Tambah w:updateFields=true ke settings.xml agar Word langsung update tanpa prompt."""
    settings = doc.settings.element
    existing = settings.find(qn('w:updateFields'))
    if existing is not None:
        settings.remove(existing)
    el = OxmlElement('w:updateFields')
    el.set(qn('w:val'), 'true')
    settings.append(el)


# ── Sisipkan setelah paragraf ──────────────────────────────────────────────────

def insert_element_after(doc, para_idx, new_elem):
    """Sisipkan elemen XML setelah paragraf pada index para_idx."""
    paras = doc.paragraphs
    ref = paras[para_idx]._element
    ref.addnext(new_elem)


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    if len(sys.argv) < 5:
        _fail("MISSING_ARGUMENTS",
              "Usage: python daftar_isi.py <input> <output> <kedalaman> <format_titik> [font] [size] [space]")

    input_file   = sys.argv[1]
    output_file  = sys.argv[2]
    kedalaman    = sys.argv[3]   # H1 | H1+H2 | H1+H2+H3
    format_titik = sys.argv[4]   # titik | tab | kecualikan_bab
    font         = sys.argv[5] if len(sys.argv) > 5 else 'Times New Roman'
    size_pt      = int(sys.argv[6]) if len(sys.argv) > 6 else 12
    try:
        line_spacing = float(sys.argv[7]) if len(sys.argv) > 7 else 1.5
    except (ValueError, TypeError):
        line_spacing = 1.5

    max_level = MAX_LEVEL_MAP.get(kedalaman, 1)
    use_dots  = format_titik in ('titik', 'kecualikan_bab')

    validate_file(input_file)

    try:
        doc = Document(input_file)
    except Exception as e:
        _fail("FILE_READ_ERROR", f"Gagal membuka file: {e}")


    # ── Seragamkan style TOC ager font/size/spasi konsisten (H1/H2 dulu) ─────
    for lvl in range(1, min(max_level, 2) + 1):
        _ensure_toc_style(doc, lvl, use_dots, font=font, size_pt=size_pt, line_spacing=line_spacing)

    # ── Fix Heading X Char styles (root cause inkonsistensi font) ────────────
    # Saat Word update TOC field, ia menaruh rStyle="Heading2Char" di setiap
    # run paragraf TOC 2/3. Style karakter ini menggunakan theme font
    # (majorHAnsi = Calibri Light) yang OVERRIDE font Times New Roman dari
    # definisi style TOC. Solusi: samakan font Heading X Char dengan TOC font.
    _fix_heading_char_styles(doc, font=font, size_pt=size_pt)

    # ── Pra-proses struktur dokumen ──────────────────────────────────────────
    # Harus sebelum detect_headings agar hasil deteksi sudah bersih.
    _merge_bab_name_paras(doc)  # "BAB II" + "LANDASAN TEORI" → satu paragraf

    # ── Deteksi heading ──────────────────────────────────────────────────────
    headings = detect_headings(doc, max_level)

    if not headings:
        _fail("NO_HEADINGS_FOUND",
              "Tidak ditemukan heading di dokumen. Pastikan dokumen menggunakan "
              "pola penomoran seperti 'BAB I', '1.', '1.1', atau judul front matter "
              "seperti 'KATA PENGANTAR', 'DAFTAR ISI'.")


    # ── TOC3 style: hitung tab dinamis berdasarkan panjang angka H3 ─────────
    if max_level >= 3:
        _h3_deep = any(
            re.match(r'^\d+\.\d+', txt)
            for _, lvl, txt in headings if lvl == 3
        )
        if not _h3_deep:
            # Cek via numbering XML: ada template multilevel (%N.%N.)?
            try:
                _n2a, _, _aft, _ = _load_numbering_full(input_file)
                for pidx, lvl, _ in headings:
                    if lvl != 3: continue
                    _pPr = doc.paragraphs[pidx]._p.find(qn('w:pPr'))
                    if _pPr is None: continue
                    _nPr = _pPr.find(qn('w:numPr'))
                    if _nPr is None: continue
                    _nid_el = _nPr.find(qn('w:numId')); _ilvl_el = _nPr.find(qn('w:ilvl'))
                    if _nid_el is None: continue
                    _nid = _nid_el.get(qn('w:val')); _ilvl = int(_ilvl_el.get(qn('w:val'), '0') if _ilvl_el is not None else '0')
                    _absid = _n2a.get(_nid)
                    if _absid and _absid in _aft:
                        _, _tmpl = _aft[_absid].get(_ilvl, (None, ''))
                        if _tmpl and _tmpl.count('%') >= 2:
                            _h3_deep = True
                            break
            except Exception:
                pass
        _h3_left_tab = 1320 if _h3_deep else 993
        _ensure_toc_style(doc, 3, use_dots, font=font, size_pt=size_pt, line_spacing=line_spacing, h3_left_tab=_h3_left_tab)
    # ── Demote semua Heading-styled paragraf yang tidak terdeteksi ────────────
    # Heading yang tidak lolos deteksi (body text salah format, metadata cover,
    # dll.) di-clone ke style non-Heading agar tidak masuk TOC field Word.
    # Cek seluruh style chain — custom style basedOn Heading X juga harus demote.
    detected_idx = {idx for idx, _, _ in headings}
    for i, para in enumerate(doc.paragraphs):
        if i in detected_idx:
            continue
        style = para.style
        is_heading_family = False
        while style:
            sname = style.name or ''
            if sname.startswith('Heading ') or sname in _ALL_HEADING_STYLES:
                is_heading_family = True
                break
            style = style.base_style
        if is_heading_family:
            _demote_heading_to_normal(para, doc)

    # ── Paksa outlineLvl=9 pada semua paragraf yang tidak terdeteksi ──────────
    # Paragraf non-heading bisa punya outlineLvl < 9 dari style definition
    # (bukan hanya dari paragraph pPr). TOC field \u membaca keduanya.
    # _set_outline_excluded() selalu menyisipkan outlineLvl=9 ke paragraph pPr
    # sehingga override style-level outlineLvl apapun nilainya.
    for i, para in enumerate(doc.paragraphs):
        if i in detected_idx:
            continue
        _set_outline_excluded(para)

    # ── Terapkan outline level pada heading terdeteksi ──────────────────────
    # Hanya set outlineLvl (metadata tak terlihat) — tidak ada perubahan visual
    # pada heading asli user. Bold, alignment, style heading dibiarkan apa adanya.
    for idx, lvl, _txt in headings:
        apply_outline_level(doc.paragraphs[idx], lvl)

    # ── Aktifkan auto-update agar Word langsung hitung nomor halaman ─────────
    enable_auto_update_fields(doc)

    # ── Hitung lebar area teks untuk tab stop nomor halaman ──────────────────
    right_tab_pos = _get_text_width_twips(doc)

    # ── Bangun TOC pre-populated (entri + bookmark + PAGEREF) ────────────────
    toc_paras = _build_toc_content(
        headings, doc, input_file, font, size_pt, use_dots, right_tab_pos, max_level
    )

    # ── Sisipkan setelah "DAFTAR ISI" ────────────────────────────────────────
    daftar_idx = find_daftar_isi_idx(doc)
    if daftar_idx is not None:
        # Hapus seluruh TOC field lama (begin → end) di paragraf langsung jika ada
        _remove_old_toc_block(doc, daftar_idx)

        import copy as _copy
        daftar_para = doc.paragraphs[daftar_idx]
        had_pb = _has_inline_page_break(daftar_para)
        if had_pb:
            _remove_inline_page_break(daftar_para)

        # Jika paragraf DAFTAR ISI menyimpan sectPr, pindahkan ke setelah TOC
        saved_sectPr = None
        daftar_pPr = daftar_para._p.find(qn('w:pPr'))
        if daftar_pPr is not None:
            sect_pr = daftar_pPr.find(qn('w:sectPr'))
            if sect_pr is not None:
                saved_sectPr = _copy.deepcopy(sect_pr)
                daftar_pPr.remove(sect_pr)

        # Cari SDT placeholder langsung setelah DAFTAR ISI heading
        sdt_elem, sdt_has_content = _find_sdt_after_daftar(
            doc.element.body, daftar_para._p
        )

        if sdt_elem is not None:
            # SDT sudah ada — isi ulang kontennya
            _populate_sdt_with_toc(sdt_elem, toc_paras)
            last_inserted = sdt_elem
        else:
            # Tidak ada SDT — buat SDT baru dan sisipkan setelah heading
            sdt_new = _create_toc_sdt(toc_paras)
            daftar_para._p.addnext(sdt_new)
            last_inserted = sdt_new

        # Re-sisipkan sectPr setelah SDT jika perlu
        if saved_sectPr is not None:
            sect_para = OxmlElement('w:p')
            sp_pPr = OxmlElement('w:pPr')
            sp_pPr.append(saved_sectPr)
            sect_para.append(sp_pPr)
            last_inserted.addnext(sect_para)
            last_inserted = sect_para

        # Page break setelah TOC agar konten berikutnya mulai di halaman baru.
        # Lalu hapus page break paragraf yang sudah ada tepat setelahnya
        # agar tidak double break.
        pb_para = _make_page_break_para()
        last_inserted.addnext(pb_para)
        last_inserted = pb_para
        _remove_trailing_empty_paras(pb_para)

        insert_pos_desc = f'setelah paragraf "DAFTAR ISI" (index {daftar_idx})'
    else:
        insert_at = 0
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                insert_at = i
                break
        for new_p in reversed(toc_paras):
            insert_element_after(doc, insert_at, new_p)
        insert_pos_desc = 'di awal dokumen (paragraf DAFTAR ISI tidak ditemukan)'

    # ── Simpan ───────────────────────────────────────────────────────────────
    try:
        os.makedirs(os.path.dirname(output_file), exist_ok=True)
        doc.save(output_file)
    except Exception as e:
        _fail("FILE_SAVE_ERROR", f"Gagal menyimpan file: {e}")

    print(json.dumps({
        "status":           "success",
        "kedalaman":        kedalaman,
        "format_titik":     format_titik,
        "heading_count":    len(headings),
        "insert_position":  insert_pos_desc,
        "headings_preview": [
            {"level": lvl, "text": txt[:60]}
            for _, lvl, txt in headings[:10]
        ],
    }))


if __name__ == '__main__':
    main()
