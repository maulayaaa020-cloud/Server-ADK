"""Debug semua 18 file - analisis advance_roman_start + output section structure"""
import os, sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx import Document
from docx.oxml.ns import qn
from utils import DocProcessor

W    = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
ROOT = os.path.normpath(os.path.join(os.path.dirname(os.path.abspath(__file__)), '..'))
SRC  = os.path.join(ROOT, 'test_files', 'paket3', 'cover 2 dimulai dari iii')
OUT  = os.path.join(SRC, 'hasil')

def _has_break(el):
    if any(br.get("{%s}type" % W) == 'page' for br in el.iter("{%s}br" % W)):
        return True
    pPr = el.find("{%s}pPr" % W)
    return pPr is not None and pPr.find("{%s}sectPr" % W) is not None

def txt(el, n=40):
    return ''.join(t.text or '' for t in el.iter('{%s}t' % W)).strip()[:n]

def get_sections(doc2):
    rows = []
    for i, sec in enumerate(doc2.sections):
        pn    = sec._sectPr.find(qn('w:pgNumType'))
        fmt   = pn.get(qn('w:fmt'), 'decimal') if pn is not None else 'decimal'
        start = pn.get(qn('w:start'), None) if pn is not None else None
        rows.append(f"sec[{i}]:{fmt}:{start}")
    return ', '.join(rows)

TARGET = [3, 6, 7, 9, 10, 12, 14, 15, 16, 17, 18]

for n in range(1, 19):
    fname = f"Docx {n}.docx"
    path  = os.path.join(SRC, fname)
    if not os.path.exists(path):
        continue

    doc      = Document(path)
    proc     = DocProcessor(doc, 'Times New Roman', 12)
    rsp, bab = proc.scan_zones()
    body_els = list(doc.element.body)
    rsp_idx  = body_els.index(rsp)
    bb       = sum(1 for el in body_els[:rsp_idx] if _has_break(el))

    new_rsp, use_exact = DocProcessor.advance_roman_start(doc, rsp, 2)
    new_idx = body_els.index(new_rsp)

    # Hitung page breaks antara rsp dan new_rsp
    breaks_between = sum(1 for el in body_els[rsp_idx:new_idx] if _has_break(el))

    # Output file sections
    out_path = os.path.join(OUT, f"Docx {n}_v2.docx")
    sec_info = ""
    if os.path.exists(out_path):
        doc2     = Document(out_path)
        sec_info = get_sections(doc2)

    flag = " *** " if n in TARGET else "     "
    print(f"{flag}Docx {n:2d}: bb={bb}, rsp=[{rsp_idx}]{txt(rsp)!r:30} -> new=[{new_idx}]{txt(new_rsp)!r:30} exact={use_exact} | {sec_info}")

print()
print("=== DETAIL CONTEXT untuk file bermasalah ===")
for n in TARGET:
    fname = f"Docx {n}.docx"
    path  = os.path.join(SRC, fname)
    if not os.path.exists(path):
        continue

    doc      = Document(path)
    proc     = DocProcessor(doc, 'Times New Roman', 12)
    rsp, bab = proc.scan_zones()
    body_els = list(doc.element.body)
    rsp_idx  = body_els.index(rsp)
    bb       = sum(1 for el in body_els[:rsp_idx] if _has_break(el))
    bab_idx  = [body_els.index(p) for p in bab]

    new_rsp, use_exact = DocProcessor.advance_roman_start(doc, rsp, 2)
    new_idx = body_els.index(new_rsp)

    print(f"\nDocx {n}: bb={bb}, rsp=[{rsp_idx}] new_rsp=[{new_idx}] use_exact={use_exact}")
    print(f"  bab_indices: {bab_idx[:4]}")

    # Show context around rsp and new_rsp
    lo = max(0, min(rsp_idx, new_idx) - 2)
    hi = min(len(body_els), max(rsp_idx, new_idx) + 5)
    for i in range(lo, hi):
        el   = body_els[i]
        tag  = el.tag.split('}')[-1] if '}' in el.tag else el.tag
        pPr  = el.find("{%s}pPr" % W)
        has_spr  = (pPr is not None and pPr.find("{%s}sectPr" % W) is not None)
        has_pgbr = any(br.get("{%s}type" % W) == 'page' for br in el.iter("{%s}br" % W))
        mark = '[sectPr]' if has_spr else ('[pgBr]' if has_pgbr else '       ')
        mk   = ' <--RSP' if i == rsp_idx else ''
        mk  += ' <--NEW' if i == new_idx else ''
        mk  += ' <--BAB' if i in bab_idx else ''
        print(f"    [{i:3d}] {tag:<5} {mark} {txt(el)!r}{mk}")
