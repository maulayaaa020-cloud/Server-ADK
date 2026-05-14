from docx import Document
import json
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt, Cm
from copy import deepcopy
import sys
import re

if len(sys.argv) < 3:
    print(json.dumps({
        "status": "error",
        "code": "MISSING_ARGUMENTS",
        "message": "Usage: python proses.py <input_file> <output_file>"
    }))
    sys.exit(1)

input_file   = sys.argv[1]
output_file  = sys.argv[2]
paket        = sys.argv[3] if len(sys.argv) > 3 else 'paket3'
font_arg     = sys.argv[4] if len(sys.argv) > 4 else 'Times New Roman'
size_arg     = sys.argv[5] if len(sys.argv) > 5 else '12 pt'
hidden_cov   = sys.argv[6] if len(sys.argv) > 6 else 'Ya'
posisi       = sys.argv[7] if len(sys.argv) > 7 else 'Tengah Bawah'
_m           = re.search(r'\d+', size_arg)
font_size_pt = int(_m.group()) if _m else 12

try:
    doc = Document(input_file)
except Exception as e:
    print(json.dumps({
        "status": "error",
        "code": "FILE_READ_ERROR",
        "message": f"Gagal membaca file: {str(e)}"
    }))
    sys.exit(1)


PAGE_NUMBER_FONT_NAME = font_arg
PAGE_NUMBER_FONT_SIZE = font_size_pt


# =========================================================
# HELPERS: HEADER / FOOTER / PAGE NUMBER
# =========================================================

def add_page_number(paragraph):
    run = paragraph.add_run()
    run.font.name = PAGE_NUMBER_FONT_NAME
    run.font.size = Pt(PAGE_NUMBER_FONT_SIZE)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = " PAGE "
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def _first_para(part):
    """Return first paragraph of a header/footer, inserting one if the part is empty."""
    if not part.paragraphs:
        part._element.append(OxmlElement('w:p'))
    return part.paragraphs[0]


def clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def _set_pn_spacing(p):
    """Set single spacing + no space before/after on a page number paragraph."""
    pf = p.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE


# =========================================================
# HELPERS: POSITION-BASED (paket1 & paket2)
# =========================================================

def _get_h_align(pos):
    if 'Tengah' in pos: return WD_ALIGN_PARAGRAPH.CENTER
    if 'Kanan'  in pos: return WD_ALIGN_PARAGRAPH.RIGHT
    return WD_ALIGN_PARAGRAPH.LEFT

def _is_top(pos):
    return 'Atas' in pos

def _place_num_in_part(part, align):
    """Write page number into a header or footer part."""
    part.is_linked_to_previous = False
    for p in part.paragraphs:
        clear_paragraph(p)
    p = _first_para(part)
    clear_paragraph(p)
    p.alignment = align
    add_page_number(p)
    _set_pn_spacing(p)


def purge_all_headers_footers():
    """Remove all content from every header/footer in the document.
    Called once before reformatting so original PAGE field codes can't double up.
    Also removes <w:sdt> Content Control elements and collapses multi-paragraph
    headers/footers to a single paragraph."""
    for section in doc.sections:
        for part in [
            section.header, section.footer,
            section.first_page_header, section.first_page_footer,
        ]:
            part.is_linked_to_previous = False
            elem = part._element
            for sdt in list(elem.findall(qn('w:sdt'))):
                elem.remove(sdt)
            paras = list(elem.findall(qn('w:p')))
            if paras:
                # Clear first paragraph, remove all extra ones
                for child in list(paras[0]):
                    paras[0].remove(child)
                for extra in paras[1:]:
                    elem.remove(extra)
            else:
                elem.append(OxmlElement('w:p'))


def clear_header(section):
    h = section.header
    h.is_linked_to_previous = False
    for p in h.paragraphs:
        clear_paragraph(p)


def clear_footer(section):
    f = section.footer
    f.is_linked_to_previous = False
    for p in f.paragraphs:
        clear_paragraph(p)


def set_page_number_format(section, fmt='decimal', start=None):
    sectPr = section._sectPr
    for old in sectPr.findall(qn('w:pgNumType')):
        sectPr.remove(old)
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), fmt)
    if start is not None:
        pgNumType.set(qn('w:start'), str(start))
    sectPr.append(pgNumType)


# =========================================================
# HELPERS: XML-LEVEL SECTION BREAK (tabel-aware)
# =========================================================

def _p_text(p_elem):
    """Return plain text of a raw <w:p> XML element."""
    return ''.join(t.text or '' for t in p_elem.findall('.//' + qn('w:t'))).strip()


def _p_has_content(p_elem):
    """True if paragraph has text OR embedded objects (images, drawings)."""
    if _p_text(p_elem):
        return True
    if p_elem.find('.//' + qn('w:drawing')) is not None:
        return True
    if p_elem.find('.//' + qn('w:pict')) is not None:
        return True
    if p_elem.find('.//' + qn('w:fldChar')) is not None:
        return True
    return False


def _has_sectPr(p_elem):
    """True if raw <w:p> element already contains a sectPr."""
    pPr = p_elem.find(qn('w:pPr'))
    return pPr is not None and pPr.find(qn('w:sectPr')) is not None


def _make_sectPr():
    """Build a next-page sectPr cloned from the document's body section."""
    new_sectPr = deepcopy(doc.sections[-1]._sectPr)
    for child in list(new_sectPr):
        tag = child.tag
        if tag.endswith('headerReference') or tag.endswith('footerReference'):
            new_sectPr.remove(child)
    type_elem = new_sectPr.find(qn('w:type'))
    if type_elem is None:
        type_elem = OxmlElement('w:type')
        new_sectPr.append(type_elem)
    type_elem.set(qn('w:val'), 'nextPage')
    return new_sectPr


def _attach_sectPr(p_elem):
    """Attach a next-page sectPr to a raw <w:p> element."""
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p_elem.insert(0, pPr)
    pPr.append(_make_sectPr())


def has_section_before_xml(target_p):
    """
    Scan body children backward from target_p (a raw <w:p> element).
    Return True if a sectPr already exists immediately before target_p
    (only empty paragraphs are skipped; a table stops the scan → False).
    """
    body = doc.element.body
    children = list(body)
    idx = children.index(target_p)
    j = idx - 1
    while j >= 0:
        elem = children[j]
        tag  = elem.tag
        if tag.endswith('}p'):
            if _has_sectPr(elem):
                pPr = elem.find(qn('w:pPr'))
                sectPr = pPr.find(qn('w:sectPr'))
                if sectPr.find(qn('w:pgSz')) is None:
                    # sectPr tidak lengkap (tanpa pgSz) — hapus dan lanjut scan
                    pPr.remove(sectPr)
                    if _p_has_content(elem):
                        return False
                    j -= 1
                    continue
                return True
            if _p_has_content(elem):   # non-empty paragraph (text or image) without sectPr
                return False
            # empty paragraph → keep scanning
        elif tag.endswith('}tbl'):
            return False               # table is here, no break between table and BAB
        elif tag.endswith('}sdt'):
            return False               # Content Control (e.g. TOC) — no break before it
        j -= 1
    return False


def _purge_continuous_sectPr_in_bab_zone(first_bab_p):
    """Hapus sectPr continuous non-landscape dari BAB zone (mulai dari BAB I heading).
    sectPr continuous portrait di area BAB tidak menciptakan halaman baru dan hanya
    menambah section ekstra yang tidak berguna. Landscape sectPr dipertahankan.
    Cover/roman area (sebelum BAB I) tidak disentuh.
    """
    body = doc.element.body
    children = list(body)
    try:
        start_idx = children.index(first_bab_p)
    except ValueError:
        return
    for elem in children[start_idx:]:
        if not elem.tag.endswith('}p'):
            continue
        pPr = elem.find(qn('w:pPr'))
        if pPr is None:
            continue
        sectPr = pPr.find(qn('w:sectPr'))
        if sectPr is None:
            continue
        type_e = sectPr.find(qn('w:type'))
        if type_e is None or type_e.get(qn('w:val')) != 'continuous':
            continue
        pgsz = sectPr.find(qn('w:pgSz'))
        if pgsz is not None and pgsz.get(qn('w:orient')) == 'landscape':
            continue
        pPr.remove(sectPr)


def _strip_empty_paras_before_bab(target_p):
    """Hapus paragraf kosong tepat sebelum target_p (heading BAB).
    Dokumen asli sering memakai Enter berulang untuk mendorong BAB ke halaman baru.
    Setelah section break disisipkan, Enter-Enter itu jadi baris kosong di atas BAB.
    Berhenti di: elemen non-paragraf, paragraf berkonten, atau paragraf ber-sectPr."""
    body = doc.element.body
    children = list(body)
    try:
        tgt_idx = children.index(target_p)
    except ValueError:
        return
    to_remove = []
    for j in range(tgt_idx - 1, -1, -1):
        elem = children[j]
        if not elem.tag.endswith('}p'):
            break
        pPr = elem.find(qn('w:pPr'))
        if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
            break
        if _p_has_content(elem):
            break
        to_remove.append(elem)
    for elem in to_remove:
        body.remove(elem)


def _remove_page_breaks_before(target_p):
    """Remove manual page breaks from target_p and consecutive empty paragraphs before it.
    Berhenti di paragraf berkonten atau elemen non-paragraf — tidak menyentuh
    page break yang memisahkan section lain (misal DAFTAR ISI dari KATA PENGANTAR)."""
    body = doc.element.body
    children = list(body)
    idx = children.index(target_p)
    # Hapus page break dari target_p sendiri
    for br in target_p.findall('.//' + qn('w:br')):
        if br.get(qn('w:type')) == 'page':
            parent = br.getparent()
            if parent is not None:
                parent.remove(br)
    # Hanya proses empty paragraphs yang langsung sebelum target
    for j in range(idx - 1, -1, -1):
        elem = children[j]
        if not elem.tag.endswith('}p'):
            break
        if _p_has_content(elem):
            break
        for br in elem.findall('.//' + qn('w:br')):
            if br.get(qn('w:type')) == 'page':
                parent = br.getparent()
                if parent is not None:
                    parent.remove(br)


def _remove_defective_sectPr_before(target_p):
    """
    Hapus sectPr continuous/tanpa-type sebelum target_p agar path insert normal bisa berjalan.
    sectPr seperti ini tidak menghasilkan page break dan harus diganti dengan sectPr baru.
    sectPr yang sudah nextPage eksplisit dibiarkan (sudah benar).
    """
    body = doc.element.body
    children = list(body)
    idx = children.index(target_p)
    for j in range(idx - 1, -1, -1):
        elem = children[j]
        if not elem.tag.endswith('}p'):
            continue
        pPr = elem.find(qn('w:pPr'))
        if pPr is None:
            if _p_has_content(elem):
                break
            continue
        sectPr = pPr.find(qn('w:sectPr'))
        if sectPr is None:
            if _p_has_content(elem):
                break
            continue
        type_elem = sectPr.find(qn('w:type'))
        val = type_elem.get(qn('w:val')) if type_elem is not None else None
        has_pgsz = sectPr.find(qn('w:pgSz')) is not None
        # Hapus sectPr yang tidak punya pgSz — ini sectPr tidak lengkap yang tidak menghasilkan
        # page break yang benar. sectPr legitimate (landscape, dll.) selalu punya pgSz.
        if not has_pgsz:
            pPr.remove(sectPr)
        return  # stop setelah sectPr pertama ditemukan


def insert_break_before_xml(target_p):
    """
    Insert a next-page section break immediately before target_p.
    Single-pass backward scan: sectPr tanpa pgSz (tidak lengkap) dihapus saat ditemukan
    dan scan dilanjutkan. sectPr dengan pgSz (legitimate) dianggap sudah benar.
    """
    _remove_page_breaks_before(target_p)

    body     = doc.element.body
    children = list(body)
    tgt_idx  = children.index(target_p)

    j = tgt_idx - 1
    while j >= 0:
        elem = children[j]
        tag  = elem.tag
        if tag.endswith('}p'):
            pPr = elem.find(qn('w:pPr'))
            if pPr is not None:
                sectPr = pPr.find(qn('w:sectPr'))
                if sectPr is not None:
                    if sectPr.find(qn('w:pgSz')) is not None:
                        # sectPr lengkap (punya pgSz) — sudah ada section break yang valid
                        return
                    else:
                        # sectPr tidak lengkap — hapus dan lanjut scan
                        pPr.remove(sectPr)
                        if _p_has_content(elem):
                            _attach_sectPr(elem)
                            return
                        j -= 1
                        continue
            if _p_has_content(elem):
                _attach_sectPr(elem)
                return
            # Empty paragraph → skip
        elif tag.endswith('}tbl') or tag.endswith('}sdt'):
            # Table atau Content Control (e.g. TOC) sebelum BAB →
            # sisipkan paragraf kosong dengan sectPr sebelum target_p
            new_p   = OxmlElement('w:p')
            new_pPr = OxmlElement('w:pPr')
            new_p.append(new_pPr)
            new_pPr.append(_make_sectPr())
            body.insert(tgt_idx, new_p)   # inserted at tgt_idx, target_p shifts to tgt_idx+1
            return
        j -= 1


# =========================================================
# SECTION FORMATTING
# =========================================================

def fmt_cover(section, first_cover=False, show_pos=None):
    """Cover: set lowerRoman format. show_pos=(align, top) shows roman numeral on cover page."""
    clear_header(section)
    clear_footer(section)
    if first_cover:
        set_page_number_format(section, 'lowerRoman', 1)   # cover = page i
    else:
        set_page_number_format(section, 'lowerRoman')       # extra cover sections continue
    if show_pos and first_cover:
        align, top = show_pos
        if top:
            _place_num_in_part(section.header, align)
        else:
            _place_num_in_part(section.footer, align)


def fmt_roman(section):
    """Roman numeral pages: center-bottom footer. Continues from cover counter."""
    section.different_first_page_header_footer = False
    section.footer_distance = Cm(1.25)
    clear_header(section)
    f = section.footer
    f.is_linked_to_previous = False
    p = _first_para(f)
    clear_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(p)
    _set_pn_spacing(p)
    # No start= → continues naturally from whatever page the cover ended on
    set_page_number_format(section, 'lowerRoman')


def fmt_bab_first(section, reset_to_1=False):
    """First section of a BAB zone.
    Page 1 of section → center-bottom footer (the BAB heading page).
    Page 2+ of section → top-right header."""
    section.different_first_page_header_footer = True
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    clear_header(section)
    clear_footer(section)

    # First page: clear header, center-bottom footer
    fph = section.first_page_header
    fph.is_linked_to_previous = False
    for p in fph.paragraphs:
        clear_paragraph(p)

    fpf = section.first_page_footer
    fpf.is_linked_to_previous = False
    for p in fpf.paragraphs:
        clear_paragraph(p)
    p = _first_para(fpf)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(p)
    _set_pn_spacing(p)

    # Other pages: top-right header, empty footer
    h = section.header
    h.is_linked_to_previous = False
    p = _first_para(h)
    clear_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(p)
    _set_pn_spacing(p)

    if reset_to_1:
        set_page_number_format(section, 'decimal', 1)
    else:
        set_page_number_format(section, 'decimal')


def fmt_bab_continuation(section):
    """Continuation section inside a BAB zone (e.g. landscape page, inline section break).
    All pages → top-right header. No different-first-page."""
    section.different_first_page_header_footer = False
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    clear_header(section)
    clear_footer(section)
    h = section.header
    h.is_linked_to_previous = False
    p = _first_para(h)
    clear_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(p)
    _set_pn_spacing(p)
    set_page_number_format(section, 'decimal')


# =========================================================
# FORMAT FUNCTIONS: PAKET 1 & 2
# =========================================================

def _fmt_uniform_section(section, align, top, fmt='decimal', start=None):
    """Apply uniform page number (header or footer) to all pages in section."""
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    section.different_first_page_header_footer = False
    clear_header(section)
    clear_footer(section)
    if top:
        _place_num_in_part(section.header, align)
    else:
        _place_num_in_part(section.footer, align)
    set_page_number_format(section, fmt, start)


def _format_paket1():
    """Paket 1: Full Arabic numeral at user-selected position, all sections."""
    align = _get_h_align(posisi)
    top   = _is_top(posisi)
    for i, section in enumerate(doc.sections):
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(1.25)
        if i == 0 and hidden_cov == 'Ya':
            section.different_first_page_header_footer = True
            for part in [section.first_page_header, section.first_page_footer]:
                part.is_linked_to_previous = False
                for p in part.paragraphs:
                    clear_paragraph(p)
            set_page_number_format(section, 'decimal', 1)
        else:
            section.different_first_page_header_footer = False
            set_page_number_format(section, 'decimal', 1 if i == 0 else None)
        if top:
            clear_footer(section)
            _place_num_in_part(section.header, align)
        else:
            clear_header(section)
            _place_num_in_part(section.footer, align)


# =========================================================
# DETECTION
# =========================================================

ROMAN_START_KEYWORDS = [
    # Indonesia
    "kata pengantar", "prakata", "ucapan terima kasih",
    "abstrak",
    "lembar pengesahan", "lembar persetujuan", "halaman pengesahan",
    "pengesahan", "persetujuan",
    "lembar pernyataan", "pernyataan keaslian",
    "halaman persembahan", "persembahan", "motto",
    "ringkasan",
    "daftar isi",
    # Inggris
    "abstract", "summary", "executive summary",
    "preface", "foreword", "acknowledgment", "acknowledgements", "acknowledgements",
    "approval page", "approval sheet", "declaration", "originality statement",
    "dedication",
    "table of contents", "list of contents", "contents",
    "list of tables", "list of figures", "list of appendices",
]

_ROMAN_PAT = r'm{0,4}(cm|cd|d?c{0,3})(xc|xl|l?x{0,3})(ix|iv|v?i{0,3})'
BAB_HEAD_RE = re.compile(
    rf'^\s*(bab|chapter)\s+({_ROMAN_PAT}|\d+)\b(.*?)$',
    re.IGNORECASE | re.DOTALL
)


def is_roman_start(text):
    lower = text.strip().lower()
    return any(lower == k or lower.startswith(k) for k in ROMAN_START_KEYWORDS)


def is_bab_heading(text):
    text = text.strip()
    if not text:
        return False
    if BAB_HEAD_RE.match(text):
        return True
    if re.match(r'^\s*(daftar\s+pustaka|referensi|references?|reference\s+list|bibliography|bibliographies|works?\s+cited|literature\s+cited)\s*$', text, re.IGNORECASE):
        return True
    if re.match(r'^\s*(lampiran|appendix|appendices|attachment)(\s+.*)?$', text, re.IGNORECASE):
        return True
    return False


def is_false_bab(para):
    """Return True when a BAB-matching paragraph is body text, not a real heading."""
    text  = para.text.strip()
    style = para.style.name.lower() if para.style else ""
    m = BAB_HEAD_RE.match(text)
    if m:
        sisa_raw = m.group(6) or ""
        sisa = sisa_raw.strip()
        # Subtitle pada baris baru (soft return) = heading dua baris asli,
        # bukan body text — skip pengecekan panjang/jumlah kata
        if not sisa_raw.startswith('\n'):
            if len(sisa) > 60:
                return True
            if len(text.split()) > 8:
                return True
        non_heading = ('body', 'list', 'quote', 'caption', 'web')
        if sisa and any(s in style for s in non_heading):
            return True
    for pattern in [
        r'^\s*(daftar\s+pustaka|referensi|references?|reference\s+list|bibliography|bibliographies|works?\s+cited|literature\s+cited)\s+\S',
        r'^\s*(lampiran|appendix|appendices|attachment)\s+[^\n]{60,}',
    ]:
        if re.match(pattern, text, re.IGNORECASE):
            return True
    # Entri daftar lampiran: "Lampiran 1: Judul\t81" — ada nomor halaman di akhir
    if re.match(r'^\s*(lampiran|appendix|appendices|attachment)', text, re.IGNORECASE):
        if (re.search(r'\t\s*\d+\s*$', text) or
                re.search(r'\.{3,}.*\d+\s*$', text) or
                re.search(r'\s{4,}\d+\s*$', text)):
            return True
    return False


def is_toc_heading(text):
    lower = text.strip().lower()
    return any(k in lower for k in [
        # Indonesia
        "daftar isi", "daftar tabel", "daftar gambar", "daftar lampiran",
        # Inggris
        "table of contents", "list of contents", "list of tables",
        "list of figures", "list of appendices",
    ])


def is_toc_entry(text):
    if '\n' in text:
        return True
    if re.search(r'\t\s*[\divxlcdmIVXLCDM]+\s*$', text):
        return True
    if re.search(r'\.{3,}.*\d+\s*$', text):
        return True
    if re.search(r'\s{4,}\d+\s*$', text):
        return True
    return False


def _find_section_start(target_p):
    """
    Return the first <w:p> of the section that contains target_p.
    Walks backward in body children to find the nearest preceding sectPr,
    then returns the first <w:p> after that break.
    If no prior sectPr found (target is in section 0), returns target_p unchanged.
    """
    body     = doc.element.body
    children = list(body)
    try:
        idx = children.index(target_p)
    except ValueError:
        return target_p
    for j in range(idx - 1, -1, -1):
        elem = children[j]
        if elem.tag.endswith('}p') and _has_sectPr(elem):
            # Section break at children[j] — section starts at j+1
            for k in range(j + 1, idx + 1):
                c = children[k]
                if c.tag.endswith('}p'):
                    return c
            break
    return target_p   # no prior section break → already in section 0


def _has_page_break_before(paras, from_idx, to_idx):
    """True if any paragraph between from_idx and to_idx has a sectPr or manual page break."""
    for para in paras[from_idx:to_idx]:
        if _has_sectPr(para._p):
            return True
        for br in para._p.findall('.//' + qn('w:br')):
            if br.get(qn('w:type')) == 'page':
                return True
    return False


def _para_has_page_break_before(para):
    """True if paragraph has 'Page Break Before' paragraph formatting set."""
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        return False
    pb = pPr.find(qn('w:pageBreakBefore'))
    if pb is None:
        return False
    return pb.get(qn('w:val'), 'true') not in ('false', '0')


# =========================================================
# MAIN PROCESSING
# =========================================================

try:
    # ----------------------------------------------------------
    # PRE-PROCESS: Purge all existing headers/footers
    # ----------------------------------------------------------
    purge_all_headers_footers()

    # ----------------------------------------------------------
    # PHASE 1: Scan paragraphs
    # Store XML element references (_p) — stable even after insertions
    # ----------------------------------------------------------
    roman_start_p     = None   # <w:p> element of first roman-zone paragraph
    bab_p_list        = []     # ordered list of <w:p> elements for BAB headings
    lampiran_found    = False  # only first lampiran heading becomes a zone boundary
    last_bab_para_idx = None   # index of last accepted BAB in all_paras
    found_numbered_bab = False  # sudah menemukan BAB bernomor/roman (BAB I, II, ...)

    inside_toc = False
    all_paras  = list(doc.paragraphs)

    for para_idx, para in enumerate(all_paras):
        text  = para.text.strip()
        lower = text.lower()

        if is_toc_heading(lower):
            inside_toc = True
            if roman_start_p is None:   # DAFTAR ISI dll. juga merupakan roman start
                roman_start_p = para._p
            continue

        if inside_toc:
            if is_toc_entry(text):
                continue
            else:
                inside_toc = False   # TOC ended; fall through

        if not text:
            continue

        if roman_start_p is None and is_roman_start(text):
            roman_start_p = para._p

        if is_bab_heading(text) and not is_false_bab(para):
            is_numbered_bab = BAB_HEAD_RE.match(text) is not None
            # Lampiran/Daftar Pustaka sebelum BAB I pertama → abaikan, masuk roman zone
            if not is_numbered_bab and not found_numbered_bab:
                continue
            if re.match(r'^\s*lampiran', text, re.IGNORECASE):
                if lampiran_found:
                    continue   # Lampiran II, III, dst. → continuation, bukan zone baru
                lampiran_found = True
            if is_numbered_bab:
                found_numbered_bab = True
            # Require a page/section break between consecutive BAB headings.
            # Sistematika Penulisan lists consecutive BABs with no page break → skip.
            if last_bab_para_idx is not None:
                has_break = _has_page_break_before(all_paras, last_bab_para_idx + 1, para_idx)
                if not has_break:
                    has_break = _para_has_page_break_before(para)  # "Page Break Before" formatting
                if not has_break:
                    # Fallback: cek konten setelah kandidat BAB sebelum BAB berikutnya.
                    # Entri Sistematika Penulisan punya sedikit paragraf ke depan (<10);
                    # BAB asli tanpa page break punya banyak konten.
                    forward_count = 0
                    for _k in range(para_idx + 1, len(all_paras)):
                        _nt = all_paras[_k].text.strip()
                        if is_bab_heading(_nt) and not is_false_bab(all_paras[_k]):
                            break
                        if _nt:
                            forward_count += 1
                    if forward_count < 10:
                        continue
            bab_p_list.append(para._p)
            last_bab_para_idx = para_idx

    # ----------------------------------------------------------
    # PHASE 1.5: Hapus sectPr continuous dari BAB zone
    # Continuous sectPr portrait di dalam BAB tidak berguna dan terlihat di Word
    # ----------------------------------------------------------
    if bab_p_list:
        _purge_continuous_sectPr_in_bab_zone(bab_p_list[0])

    # ----------------------------------------------------------
    # PHASE 2: Add section breaks at zone boundaries (XML-aware)
    # ----------------------------------------------------------
    if roman_start_p is not None:
        roman_start_p = _find_section_start(roman_start_p)
        insert_break_before_xml(roman_start_p)

    for bab_p in bab_p_list:
        insert_break_before_xml(bab_p)
        _strip_empty_paras_before_bab(bab_p)

    # ----------------------------------------------------------
    # PHASE 3: Rebuild section boundary map
    # (doc.paragraphs now includes any new paragraphs inserted in Phase 2)
    # ----------------------------------------------------------
    breaks = []
    for para in doc.paragraphs:
        if _has_sectPr(para._p):
            breaks.append(para._p)   # store _p references, not indices

    # Build index-based boundaries for section lookup
    para_list   = list(doc.paragraphs)
    para_p_list = [p._p for p in para_list]

    # boundaries[s] = index into para_p_list where section s starts
    break_indices = [para_p_list.index(bp) for bp in breaks]
    boundaries    = [0] + [bi + 1 for bi in break_indices]
    n_sections    = len(boundaries)

    def para_to_section_xml(p_elem):
        """Return the section index that currently contains p_elem."""
        try:
            para_idx = para_p_list.index(p_elem)
        except ValueError:
            return n_sections - 1
        for s in range(n_sections):
            start = boundaries[s]
            end   = boundaries[s + 1] if s + 1 < n_sections else len(para_p_list)
            if start <= para_idx < end:
                return s
        return n_sections - 1

    roman_sec    = para_to_section_xml(roman_start_p) if roman_start_p is not None else None
    bab_sec_list = [para_to_section_xml(p) for p in bab_p_list]
    first_bab_sec = bab_sec_list[0] if bab_sec_list else None

    # ----------------------------------------------------------
    # PHASE 4: Apply header/footer format to every section
    # ----------------------------------------------------------
    if paket == 'paket1':
        _format_paket1()
    else:
        _p2_align   = _get_h_align(posisi)
        _p2_top     = _is_top(posisi)
        _cov_show_2 = None if hidden_cov == 'Ya' else (_p2_align, _p2_top)
        _cov_show_3 = None if hidden_cov == 'Ya' else (WD_ALIGN_PARAGRAPH.CENTER, False)

        for i, section in enumerate(doc.sections):

            # ---- Cover zone: all sections before roman_sec ----
            if roman_sec is not None and i < roman_sec:
                show = _cov_show_2 if paket == 'paket2' else _cov_show_3
                fmt_cover(section, first_cover=(i == 0), show_pos=show)
                continue
            if roman_sec is None and i == 0:
                show = _cov_show_2 if paket == 'paket2' else _cov_show_3
                fmt_cover(section, first_cover=True, show_pos=show)
                continue

            # ---- Roman zone ----
            if first_bab_sec is None or i < first_bab_sec:
                if paket == 'paket2':
                    _fmt_uniform_section(section, _p2_align, _p2_top, fmt='lowerRoman')
                else:
                    fmt_roman(section)
                continue

            # ---- BAB zones ----
            zone_idx      = -1
            is_zone_first = False
            for k, bab_sec in enumerate(bab_sec_list):
                next_sec = bab_sec_list[k + 1] if k + 1 < len(bab_sec_list) else n_sections
                if bab_sec <= i < next_sec:
                    zone_idx      = k
                    is_zone_first = (i == bab_sec)
                    break

            if paket == 'paket2':
                _fmt_uniform_section(
                    section, _p2_align, _p2_top, fmt='decimal',
                    start=1 if (is_zone_first and zone_idx == 0) else None
                )
            elif is_zone_first:
                fmt_bab_first(section, reset_to_1=(zone_idx == 0))
            else:
                fmt_bab_continuation(section)

except Exception as e:
    print(json.dumps({
        "status": "error",
        "code": "PROCESSING_ERROR",
        "message": f"Gagal memproses dokumen: {str(e)}"
    }))
    sys.exit(1)


# =========================================================
# SAVE
# =========================================================
try:
    doc.save(output_file)
except Exception as e:
    print(json.dumps({
        "status": "error",
        "code": "FILE_SAVE_ERROR",
        "message": f"Gagal menyimpan file: {str(e)}"
    }))
    sys.exit(1)


# =========================================================
# OUTPUT SUMMARY
# =========================================================
out_breaks = []
out_paras  = list(doc.paragraphs)
for para in out_paras:
    if _has_sectPr(para._p):
        out_breaks.append(out_paras.index(para))

out_bounds    = [0] + [bi + 1 for bi in out_breaks]
sections_info = []

for sec_idx, start in enumerate(out_bounds):
    end = out_bounds[sec_idx + 1] if sec_idx + 1 < len(out_bounds) else len(out_paras)
    first_text = ""
    for i in range(start, min(start + 10, end)):
        t = out_paras[i].text.strip()
        if t:
            first_text = t[:50]
            break
    sections_info.append({"index": sec_idx, "first_content": first_text})

print(json.dumps({
    "status": "success",
    "total_sections": len(doc.sections),
    "sections": sections_info
}))
